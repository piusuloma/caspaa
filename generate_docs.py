from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colour palette ──────────────────────────────────────────────────────────
DARK_NAVY   = RGBColor(0x0a, 0x25, 0x40)
EMERALD     = RGBColor(0x0d, 0x6e, 0x3f)
LIGHT_MINT  = RGBColor(0xd1, 0xfa, 0xe5)
RED         = RGBColor(0xdc, 0x26, 0x26)
LIGHT_RED   = RGBColor(0xfe, 0xe2, 0xe2)
AMBER       = RGBColor(0xd9, 0x77, 0x06)
LIGHT_AMBER = RGBColor(0xfe, 0xf3, 0xc7)
BLUE        = RGBColor(0x1d, 0x4e, 0xd8)
LIGHT_BLUE  = RGBColor(0xdb, 0xea, 0xfe)
GRAY_BG     = RGBColor(0xf3, 0xf4, 0xf6)
GRAY_TEXT   = RGBColor(0x6b, 0x72, 0x80)
WHITE       = RGBColor(0xff, 0xff, 0xff)
MID_GRAY    = RGBColor(0x37, 0x41, 0x51)


# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    hex_color = str(rgb)  # RGBColor.__str__ returns 6-char hex e.g. '0A2540'
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    # Remove any existing shd
    for existing in tcPr.findall(qn('w:shd')):
        tcPr.remove(existing)
    tcPr.append(shd)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val is not None:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'),  val.get('val',  'single'))
            el.set(qn('w:sz'),   val.get('sz',   '4'))
            el.set(qn('w:space'),'0')
            el.set(qn('w:color'),val.get('color','000000'))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def set_row_height(row, twips):
    trPr = row._tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(twips))
    trPr.append(trHeight)


def add_paragraph(doc, text='', bold=False, size=11, color=DARK_NAVY,
                  align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6,
                  italic=False, highlight_color=None):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size  = Pt(size)
        run.font.color.rgb = color
    return p


def add_heading(doc, text, level=1, color=DARK_NAVY, space_before=14, space_after=4):
    sizes = {1: 22, 2: 16, 3: 13, 4: 12}
    p = add_paragraph(doc, text, bold=True, size=sizes.get(level, 12),
                      color=color, space_before=space_before, space_after=space_after)
    return p


def add_label(doc, text, color=EMERALD):
    """Small all-caps section label."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = color
    return p


def add_bullet(doc, text, indent=0, bold_prefix=None, color=MID_GRAY, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.3 + indent * 0.2)
    if bold_prefix:
        rb = p.add_run(bold_prefix + ' ')
        rb.bold = True
        rb.font.size = Pt(size)
        rb.font.color.rgb = DARK_NAVY
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def add_page_break(doc):
    doc.add_page_break()


def set_doc_margins(doc, top=1.0, bottom=1.0, left=1.1, right=1.1):
    for section in doc.sections:
        section.top_margin    = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin   = Inches(left)
        section.right_margin  = Inches(right)


def add_colored_table_row(table, cells_data, bg_color, text_color=WHITE, bold=True, size=10):
    """Add a header row to a table with background color."""
    row = table.add_row()
    for i, (cell, data) in enumerate(zip(row.cells, cells_data)):
        set_cell_bg(cell, bg_color)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(str(data))
        run.bold = bold
        run.font.size = Pt(size)
        run.font.color.rgb = text_color
    return row


def add_data_row(table, cells_data, bg_color=WHITE, text_color=MID_GRAY, size=10, bold_first=True):
    row = table.add_row()
    for i, (cell, data) in enumerate(zip(row.cells, cells_data)):
        set_cell_bg(cell, bg_color)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(str(data))
        run.bold = (i == 0 and bold_first)
        run.font.size = Pt(size)
        run.font.color.rgb = DARK_NAVY if (i == 0 and bold_first) else text_color
    return row


def set_col_widths(table, widths):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths):
                cell.width = widths[i]


# ════════════════════════════════════════════════════════════════════════════
# DOCUMENT 1: Feature Gap Analysis
# ════════════════════════════════════════════════════════════════════════════

def build_gap_analysis():
    doc = Document()
    set_doc_margins(doc)

    # ── Default style ────────────────────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ── COVER ────────────────────────────────────────────────────────────────
    # Title block table (no borders, just colour)
    cover = doc.add_table(rows=1, cols=1)
    cover.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = cover.rows[0].cells[0]
    set_cell_bg(c, DARK_NAVY)
    set_row_height(cover.rows[0], 8000)

    # Logo line
    p = c.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    r = p.add_run('CASPAA')
    r.bold = True
    r.font.size = Pt(36)
    r.font.color.rgb = WHITE

    # Badge
    p2 = c.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(24)
    r2 = p2.add_run('STRATEGIC DOCUMENT  ·  CONFIDENTIAL')
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(0x34, 0xd3, 0x99)

    # Main title
    p3 = c.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run('Feature Gap Analysis &\nCompetitive Positioning')
    r3.bold = True
    r3.font.size = Pt(28)
    r3.font.color.rgb = WHITE

    # Subtitle
    p4 = c.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_before = Pt(16)
    r4 = p4.add_run(
        'A structured assessment of CASPAA\'s current capabilities versus the EDVES benchmark,\n'
        'with strategic recommendations for the next product phase.'
    )
    r4.font.size = Pt(11)
    r4.font.color.rgb = RGBColor(0xb0, 0xc4, 0xde)

    # Meta row
    p5 = c.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p5.paragraph_format.space_before = Pt(40)
    r5 = p5.add_run('Prepared By: Product & Engineering     |     Date: June 2026     |     Version: 1.0     |     Audience: Stakeholders')
    r5.font.size = Pt(9)
    r5.font.color.rgb = RGBColor(0x90, 0xa8, 0xc0)

    add_page_break(doc)

    # ── 01 EXECUTIVE SUMMARY ─────────────────────────────────────────────────
    add_label(doc, '01 — Overview')
    add_heading(doc, 'Executive Summary', 1)
    add_paragraph(doc,
        'This document benchmarks CASPAA against EDVES — one of the leading school management '
        'platforms in Africa with 1,800+ schools across 11 countries. The goal is to identify where '
        'CASPAA leads, where it lags, and what to build next to win the market.',
        size=11, color=GRAY_TEXT, space_after=12)

    # KPI table
    kpi = doc.add_table(rows=2, cols=3)
    kpi.alignment = WD_TABLE_ALIGNMENT.CENTER
    kpi_data = [
        ('7', 'Unique Advantages', 'Features CASPAA has that EDVES does not offer — genuine competitive moats.', EMERALD, LIGHT_MINT),
        ('14', 'Critical Gaps', 'Features EDVES has that CASPAA is fully missing — must close to compete at parity.', AMBER, LIGHT_AMBER),
        ('5', 'Partial / Stub', 'Features CASPAA has started but not completed — quick wins with focused effort.', BLUE, LIGHT_BLUE),
    ]
    for i, (num, label, desc, col, bg) in enumerate(kpi_data):
        cell = kpi.rows[0].cells[i]
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(10)
        r = p.add_run(num)
        r.bold = True; r.font.size = Pt(32); r.font.color.rgb = col

        p2 = cell.add_paragraph(label)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.runs[0].bold = True; p2.runs[0].font.size = Pt(11); p2.runs[0].font.color.rgb = DARK_NAVY

        p3 = cell.add_paragraph(desc)
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_after = Pt(10)
        p3.runs[0].font.size = Pt(9); p3.runs[0].font.color.rgb = GRAY_TEXT

    set_col_widths(kpi, [Inches(2.1)] * 3)
    doc.add_paragraph()

    # Strategic Positioning box
    pos = doc.add_table(rows=1, cols=1)
    pc = pos.rows[0].cells[0]
    set_cell_bg(pc, DARK_NAVY)
    ph = pc.add_paragraph('Strategic Positioning Statement')
    ph.paragraph_format.space_before = Pt(12)
    ph.runs[0].bold = True; ph.runs[0].font.size = Pt(13); ph.runs[0].font.color.rgb = RGBColor(0x34, 0xd3, 0x99)

    pb = pc.add_paragraph(
        'CASPAA is not simply a school management system — it is a School Operating System with a '
        'built-in financial infrastructure layer that no African EdTech platform currently offers at the '
        'same depth. Our fee lending, reconciliation engine, SaaS operator dashboard, and digital consent '
        'workflows are genuine differentiators. The priority now is to close the learning content gap '
        '(video lessons, class notes, simulations) and complete the operational modules (transport, '
        'hostel, health) so that CASPAA is a complete platform — not just a strong one.'
    )
    pb.paragraph_format.space_after = Pt(12)
    pb.runs[0].font.size = Pt(10.5); pb.runs[0].font.color.rgb = RGBColor(0xd0, 0xe0, 0xf0)

    doc.add_paragraph()

    # ── 02 WHERE WE LEAD vs LAG ──────────────────────────────────────────────
    add_label(doc, '02 — Competitive Landscape')
    add_heading(doc, 'Where We Lead vs. Where We Lag')

    add_heading(doc, 'CASPAA Advantages — Features EDVES Does Not Have', 2, color=EMERALD, space_before=8)
    advantages = [
        ('Fee Lending & Financing Module', 'Parents apply for fee loans; full repayment schedule tracking. No African EdTech competitor has this.'),
        ('Platform Super Admin Layer', 'MRR, ARR, school health scores, remittance management — a true SaaS operator view.'),
        ('Role-Based Chat System', 'Parents can only message their children\'s teachers (not the whole school). Admin/Principal has separate "Chat with Teacher" and "Chat with Parent" buttons. Includes file/image attachments and WhatsApp handoff. EDVES messaging is generic; CASPAA\'s is context-aware.'),
        ('Digital Consent Forms', 'Paperless activity & trip approvals with parent tracking. EDVES does not offer this.'),
        ('Student Gamification', 'Stars, points, achievement badges driving student engagement — more sophisticated than EDVES\'s house points.'),
        ('Substitute Teacher Coverage Workflow', 'Structured coverage requests and acceptance; no equivalent in EDVES.'),
        ('Role Impersonation', 'Admins can view the system exactly as any user role sees it — powerful for support and demos.'),
    ]
    for title, desc in advantages:
        add_bullet(doc, desc, bold_prefix=title + ':', color=MID_GRAY)

    add_heading(doc, 'Critical Gaps — Features EDVES Has That We Are Missing', 2, color=RED, space_before=12)
    gaps = [
        ('Video Lessons & E-Videos', 'Student and teacher video content library — core daily student usage.'),
        ('Virtual Classroom', 'Live online classes; currently only a stub in CASPAA.'),
        ('Class Notes for Students', 'Teacher-posted notes accessible in the student portal.'),
        ('Maths & Science Simulations (PhET)', 'Interactive science labs — a key EDVES selling point, free to implement.'),
        ('School Bus & Transport Management', 'Route tracking, driver assignment, parent pickup tracking.'),
        ('Hostel Management', 'Boarding student records and room assignments.'),
        ('House Point System', 'Dedicated pastoral care module — currently buried in discipline.'),
        ('Parent Feedback Tool', 'Structured parent-to-school feedback forms, distinct from consent forms.'),
        ('Secure Pickup (Parent Portal)', 'Approved pickup person management per child.'),
        ('Event Calendar', 'School-wide events visible across all roles — Admin, Teacher, Parent, Student.'),
        ('Teacher Payslip View', 'Teachers viewing their own salary statements; payroll data exists but view is missing.'),
        ('Communication Diary', 'Structured per-student daily notes (Homework/Behaviour/Health categories) with read receipts — distinct from the free-form chat system.'),
        ('Formative Assessment', 'Distinct in-class ongoing assessment, not just end-of-term CBT.'),
        ('Bulk SMS & Email', 'Critical in Nigerian market. Currently marked TBD in codebase.'),
    ]
    for title, desc in gaps:
        add_bullet(doc, desc, bold_prefix=title + ':', color=MID_GRAY)

    add_page_break(doc)

    # ── 03 FULL FEATURE TABLE ────────────────────────────────────────────────
    add_label(doc, '03 — Detailed Comparison')
    add_heading(doc, 'Full Feature-by-Feature Breakdown')
    add_paragraph(doc,
        'Every major feature category compared across both platforms, with priority classification for the CASPAA roadmap.',
        size=10.5, color=GRAY_TEXT, space_after=10)

    table_data = [
        # (Feature, EDVES, CASPAA, Priority, Notes)
        # Section headers marked with a special flag
        ('__SECTION__', 'Financial Management', '', '', ''),
        ('Fee Structure Management',        '✓ Full',    '✓ Full',    '—',       'Both complete. CASPAA adds activity-based fee layering.'),
        ('Invoice Generation',              '✓ Full',    '✓ Full',    '—',       'CASPAA includes line-item detail per activity.'),
        ('Online Payment Integration',      '✓ Full',    '⚠ Partial', 'P1',     'Gateway UI exists; actual payment processing not wired.'),
        ('Payment Reconciliation',          '⚠ Basic',   '✓ Full',    'LEAD',   'CASPAA advantage: dedicated reconciliation workflow.'),
        ('Expense Tracking',                '✓ Full',    '✓ Full',    '—',       'Parity.'),
        ('Payroll Processing',              '✓ Full',    '✓ Full',    '—',       'Parity.'),
        ('Teacher Payslip (self-view)',     '✓ Full',    '✗ Missing', 'P2',     'Teachers cannot see their own payslip; data exists, view missing.'),
        ('Financial Reports',               '✓ Full',    '✓ Full',    '—',       'Parity.'),
        ('Fee Lending / Financing',         '✗ None',    '✓ Full',    'LEAD',   'Unique CASPAA differentiator. Parents borrow school fees.'),
        ('__SECTION__', 'Academic & Learning', '', '', ''),
        ('Result Management',               '✓ Full',    '✓ Full',    '—',       'Parity.'),
        ('Attendance Tracking',             '✓ Full',    '✓ Full',    '—',       'CASPAA adds staff clock-in/out.'),
        ('Online CBT / Exams',              '✓ Full',    '✓ Full',    '—',       'Parity.'),
        ('Offline CBT / Exams',             '✓ Full',    '✗ Missing', 'P3',     'EDVES supports offline exam-taking.'),
        ('Assignments',                     '✓ Full',    '✓ Full',    '—',       'Parity.'),
        ('Lesson Plans (Teacher)',          '✓ Full',    '✓ Full',    '—',       'Parity.'),
        ('Class Notes (Student access)',    '✓ Full',    '✗ Missing', 'P1',     'Teachers post notes; students read them in portal.'),
        ('Video Lessons / E-Videos',       '✓ Full',    '⚠ Stub',   'P1',     'Learning hub exists but has no video content functionality.'),
        ('Virtual Classroom (Live)',        '✓ Full',    '⚠ Stub',   'P1',     'Can integrate Zoom/Google Meet or Jitsi.'),
        ('Maths & Science Simulations',    '✓ PhET',    '✗ Missing', 'P1',     'EDVES key differentiator. PhET is free and embeddable.'),
        ('Formative Assessment',            '✓ Full',    '✗ Missing', 'P2',     'Ongoing in-class assessment, distinct from term-end CBT.'),
        ('Timetable',                       '✓ Full',    '✓ Full',    '—',       'Parity.'),
        ('Curriculum Mapping',             '✓ Multi',   '⚠ Partial', 'P3',    'EDVES supports NERDC, Cambridge, IGCSE, IB.'),
        ('__SECTION__', 'Student Management', '', '', ''),
        ('Student Enrollment',             '✓ Full',    '✓ Full',    '—',       'Parity.'),
        ('Admissions Workflow',            '✓ Full',    '⚠ Partial', 'P2',    'Full application form workflow incomplete.'),
        ('Behaviour / Discipline',         '✓ Full',    '✓ Full',    '—',       'Parity.'),
        ('House Point System',             '✓ Dedicated','⚠ Buried', 'P2',    'CASPAA has behavior tracking but no dedicated house point module.'),
        ('Student Gamification',           '⚠ Basic',   '✓ Full',    'LEAD',   'CASPAA stars/points/badges system is more advanced.'),
        ('Alumni Tracking',                '⚠ Basic',   '✓ Full',    'LEAD',   'CASPAA has career tracking and alumni network support.'),
        ('__SECTION__', 'Staff Management', '', '', ''),
        ('Staff Directory',                '✓ Full',    '✓ Full',    '—',       'Parity.'),
        ('Leave Management',               '✓ Full',    '✓ Full',    '—',       'Parity.'),
        ('Performance Appraisal',          '✓ Full',    '✓ Full',    '—',       'Parity. CASPAA has multi-stage approval.'),
        ('Substitute Coverage',            '✗ None',    '✓ Full',    'LEAD',   'Unique CASPAA feature.'),
        ('Teacher Analytics',              '✓ Full',    '⚠ Partial', 'P2',    'Cross-teacher vs. student analytics missing.'),
        ('__SECTION__', 'Communication', '', '', ''),
        ('Messaging / Chat System',        '✓ Full',    '✓ Full',    'LEAD',    'CASPAA has role-based routing: parents message only their children\'s teachers; admin/principal can message both parents and teachers separately. Includes file/image attachments and WhatsApp handoff.'),
        ('Announcements / Broadcast',      '✓ Full',    '✓ Full',    '—',       'Parity. CASPAA supports audience targeting (all / parents / teachers).'),
        ('Bulk SMS',                        '✓ Full',    '✗ Missing', 'P1',     'Critical for Nigerian schools. Currently TBD in CASPAA.'),
        ('Email Notifications',            '✓ Full',    '✗ Missing', 'P1',     'Currently TBD in CASPAA.'),
        ('Communication Diary',            '✓ Full',    '✗ Missing', 'P2',     'Structured per-student daily notes with categories (Homework/Behaviour/Health). Different from the chat system — this is record-keeping, not conversation.'),
        ('Parent Feedback Tool',           '✓ Full',    '✗ Missing', 'P2',     'Formal school feedback surveys.'),
        ('Event Calendar',                 '✓ Full',    '✗ Missing', 'P1',     'School-wide event calendar visible to all roles.'),
        ('Digital Consent Forms',          '✗ None',    '✓ Full',    'LEAD',   'Unique CASPAA feature.'),
        ('__SECTION__', 'Operations', '', '', ''),
        ('Inventory Management',           '✓ Full',    '⚠ Stub',   'P2',     'Menu exists; logic not built.'),
        ('Health / Sickbay',               '✓ Full',    '⚠ Stub',   'P2',     'Menu exists; logic not built.'),
        ('Transport / School Bus',         '✓ Full',    '✗ Missing', 'P2',     'Route management, driver assignment, parent tracking.'),
        ('Secure Pickup (Parent)',         '✓ Full',    '✗ Missing', 'P2',     'Approved pickup persons per child.'),
        ('Hostel Management',              '✓ Full',    '✗ Missing', 'P3',     'Boarding student records, room assignments.'),
        ('__SECTION__', 'AI & Intelligence', '', '', ''),
        ('AI Grading Assistance',          '✓ Full',    '✗ Missing', 'P3',     'AI-assisted marking and feedback generation.'),
        ('Student Performance Prediction', '✓ Full',    '✗ Missing', 'P3',     'AI flags at-risk students based on trends.'),
        ('AI Question Generation',         '✓ Full',    '✗ Missing', 'P3',     'Auto-generate CBT questions from lesson notes.'),
        ('AI Chatbot',                     '✓ Full',    '✗ Missing', 'P3',     'Student and teacher assistant chatbot.'),
        ('__SECTION__', 'Platform & Business', '', '', ''),
        ('Multi-Curriculum Support',       '✓ Full',    '✗ Missing', 'P3',     'NERDC, Cambridge, IGCSE, IB.'),
        ('SaaS Operator Dashboard',        '✗ None',    '✓ Full',    'LEAD',   'MRR, ARR, school health — unique to CASPAA.'),
        ('Role Impersonation',             '✗ None',    '✓ Full',    'LEAD',   'Admin can view as any user for support/debugging.'),
        ('Tiered Subscription Plans',      '✓ Full',    '⚠ Partial', 'P2',    'Plans defined; feature gating not fully enforced.'),
    ]

    tbl = doc.add_table(rows=0, cols=5)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'

    # Header row
    hdr = add_colored_table_row(tbl,
        ['Feature', 'EDVES', 'CASPAA Now', 'Priority', 'Notes'],
        DARK_NAVY, WHITE, bold=True, size=9)

    alt = False
    for row_data in table_data:
        if row_data[0] == '__SECTION__':
            r = add_colored_table_row(tbl,
                [row_data[1], '', '', '', ''],
                RGBColor(0x1a, 0x4a, 0x7a), WHITE, bold=True, size=9)
            # Merge section header cells
            r.cells[0].merge(r.cells[1])
            r.cells[0].merge(r.cells[2])
            r.cells[0].merge(r.cells[3])
            r.cells[0].merge(r.cells[4])
            alt = False
        else:
            feature, edves, caspaa, priority, notes = row_data
            bg = GRAY_BG if alt else WHITE

            # Colour-code status fields
            def status_color(val):
                if '✓' in val: return (EMERALD, LIGHT_MINT)
                if '✗' in val: return (RED, LIGHT_RED)
                if '⚠' in val: return (AMBER, LIGHT_AMBER)
                if val in ('LEAD', '—'): return (EMERALD, LIGHT_MINT)
                return (DARK_NAVY, bg)

            priority_colors = {
                'P1': (RED, LIGHT_RED), 'P2': (AMBER, LIGHT_AMBER),
                'P3': (BLUE, LIGHT_BLUE), 'LEAD': (EMERALD, LIGHT_MINT), '—': (GRAY_TEXT, bg)
            }

            dr = tbl.add_row()
            cells = dr.cells
            for ci, (val, col_bg) in enumerate([
                (feature, bg),
                (edves,   status_color(edves)[1]),
                (caspaa,  status_color(caspaa)[1]),
                (priority, priority_colors.get(priority, (GRAY_TEXT, bg))[1]),
                (notes,   bg),
            ]):
                set_cell_bg(cells[ci], col_bg)
                cells[ci].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p = cells[ci].paragraphs[0]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after  = Pt(2)
                run = p.add_run(val)
                run.bold = (ci == 0)
                run.font.size = Pt(9)
                if ci == 0:
                    run.font.color.rgb = DARK_NAVY
                elif ci in (1, 2):
                    run.font.color.rgb = status_color(val)[0]
                elif ci == 3:
                    run.font.color.rgb = priority_colors.get(val, (GRAY_TEXT, bg))[0]
                else:
                    run.font.color.rgb = GRAY_TEXT
            alt = not alt

    set_col_widths(tbl, [Inches(1.8), Inches(0.85), Inches(0.85), Inches(0.65), Inches(2.35)])
    doc.add_paragraph()

    add_page_break(doc)

    # ── 04 PRIORITY ROADMAP ───────────────────────────────────────────────────
    add_label(doc, '04 — Roadmap')
    add_heading(doc, 'Priority Breakdown for Next Phase')
    add_paragraph(doc,
        'Prioritized by stakeholder visibility, competitive parity, and implementation effort.',
        size=10.5, color=GRAY_TEXT, space_after=10)

    phases = [
        ('P1 — Critical (Must Build Now)', RED, [
            ('1', 'Event Calendar', 'All roles need a shared school calendar. High visibility to parents & teachers.'),
            ('2', 'Bulk SMS & Email', 'Critical in Nigerian market. Parents expect SMS alerts for fees, results, events.'),
            ('3', 'Class Notes (Student Portal)', 'Teachers post notes; students read in Learning hub. Core to student daily usage.'),
            ('4', 'Video Lessons', 'Embed video content per subject. Can start with YouTube embeds + upload links.'),
            ('5', 'Science Simulations (PhET)', 'PhET is free & embeddable. Major EDVES selling point — easy to replicate.'),
            ('6', 'Online Payment Gateway', 'UI exists; needs to connect to Paystack or Flutterwave to process real payments.'),
        ]),
        ('P2 — Important (Build Next)', AMBER, [
            ('7',  'House Point System', 'Dedicated pastoral care module surfaced to students, teachers, and parents.'),
            ('8',  'Teacher Payslip View', 'Data already in payroll; just needs a teacher-facing view. Quick win.'),
            ('9',  'Transport Management', 'School bus routes, driver assignment, parent tracking alerts and secure pickup.'),
            ('10', 'Parent Feedback Tool', 'Structured survey/rating from parents to school. Builds trust and NPS.'),
            ('11', 'Communication Diary', 'Structured daily teacher ↔ parent notes with read receipts.'),
            ('12', 'Formative Assessment', 'In-class ongoing quizzes and scoring, distinct from end-of-term CBT.'),
            ('13', 'Health / Sickbay', 'Complete the stub: student health records, incidents, parent notification.'),
            ('14', 'Inventory Management', 'Complete the stub: items, stock levels, procurement, issue tracking.'),
        ]),
        ('P3 — Strategic (Future Phase)', BLUE, [
            ('15', 'Virtual Classroom (Live)', 'Zoom/Jitsi/Google Meet integration for live online classes.'),
            ('16', 'AI Grading & Assistance', 'AI-powered marking, question generation from lesson notes, chatbot.'),
            ('17', 'Student Performance AI', 'Predictive analytics flagging at-risk students based on attendance and results.'),
            ('18', 'Hostel Management', 'Boarding students: rooms, meals, check-in/out, parent communication.'),
            ('19', 'Multi-Curriculum Support', 'NERDC, Cambridge, IGCSE, IB curriculum tagging on lessons and assessments.'),
            ('20', 'Offline CBT', 'Exam-taking when internet is unavailable; syncs when reconnected.'),
            ('21', 'Reseller / Partner Network', 'Commission-based partner portal for city-level resellers.'),
        ]),
    ]

    for phase_title, color, items in phases:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(6)
        r = p.add_run(phase_title)
        r.bold = True; r.font.size = Pt(12); r.font.color.rgb = color

        pt = doc.add_table(rows=0, cols=3)
        pt.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_colored_table_row(pt, ['#', 'Module', 'Why It Matters'], color, WHITE, size=9)
        for num, name, why in items:
            dr = pt.add_row()
            for ci, val in enumerate([num, name, why]):
                dr.cells[ci].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                pp = dr.cells[ci].paragraphs[0]
                pp.paragraph_format.space_before = Pt(3)
                pp.paragraph_format.space_after  = Pt(3)
                run = pp.add_run(val)
                run.bold = (ci == 1)
                run.font.size = Pt(9.5)
                run.font.color.rgb = color if ci == 0 else (DARK_NAVY if ci == 1 else GRAY_TEXT)
        set_col_widths(pt, [Inches(0.35), Inches(1.8), Inches(4.35)])
        doc.add_paragraph()

    add_page_break(doc)

    # ── 05 CONCLUSION ─────────────────────────────────────────────────────────
    add_label(doc, '05 — Conclusion')
    add_heading(doc, 'Recommendation')
    add_paragraph(doc,
        'CASPAA enters the market with a stronger financial and platform infrastructure than EDVES. '
        'The single most important action is to close the learning content gap — because that is what '
        'students and parents interact with daily, and it is the primary reason schools choose EDVES today.',
        size=11, color=GRAY_TEXT, space_after=12)

    add_heading(doc, 'Immediate Focus (Next 60 Days)', 2, color=EMERALD, space_before=8)
    for item in [
        'Wire up online payment gateway (Paystack/Flutterwave)',
        'Build Event Calendar visible across all roles',
        'Implement Class Notes in the student learning hub',
        'Add Video Lesson embedding per subject',
        'Integrate PhET Simulations for Maths & Science',
        'Activate Bulk SMS via Africa\'s Talking or Termii',
    ]:
        add_bullet(doc, item, color=MID_GRAY)

    add_heading(doc, 'Medium-Term Focus (60–120 Days)', 2, color=BLUE, space_before=10)
    for item in [
        'Dedicate House Point System as its own module',
        'Build Transport Management + Secure Pickup',
        'Add Parent Feedback Tool',
        'Complete Inventory & Health/Sickbay stubs',
        'Add Communication Diary for teacher-parent notes',
        'Begin AI roadmap: question generation from lesson notes',
    ]:
        add_bullet(doc, item, color=MID_GRAY)

    doc.save(r'C:\Users\USER\Desktop\CASPAA\CASPAA_Feature_Gap_Analysis.docx')
    print('Done: Gap Analysis document saved.')


# ════════════════════════════════════════════════════════════════════════════
# DOCUMENT 2: Implementation Roadmap
# ════════════════════════════════════════════════════════════════════════════

def add_module_block(doc, num, title, subtitle, effort_label, effort_color, roles,
                     data_steps, ui_steps, tech_notes=None):
    """Render a full module section."""
    # Module header bar
    hdr_tbl = doc.add_table(rows=1, cols=1)
    hdr_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hc = hdr_tbl.rows[0].cells[0]
    set_cell_bg(hc, DARK_NAVY)
    hp = hc.paragraphs[0]
    hp.paragraph_format.space_before = Pt(5)
    hp.paragraph_format.space_after  = Pt(5)
    hr1 = hp.add_run(f'Module {num}: {title}')
    hr1.bold = True; hr1.font.size = Pt(12); hr1.font.color.rgb = WHITE
    hr2 = hp.add_run(f'  |  {subtitle}')
    hr2.font.size = Pt(10); hr2.font.color.rgb = RGBColor(0xb0, 0xc4, 0xde)

    # Metadata row
    meta = doc.add_table(rows=1, cols=3)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_bg(meta.rows[0].cells[0], effort_color)
    set_cell_bg(meta.rows[0].cells[1], LIGHT_BLUE)
    set_cell_bg(meta.rows[0].cells[2], GRAY_BG)
    for ci, (label, val, tc) in enumerate([
        ('EFFORT', effort_label, DARK_NAVY),
        ('ROLES', ' · '.join(roles), DARK_NAVY),
        ('PHASE', f'Module {num}', GRAY_TEXT),
    ]):
        cell = meta.rows[0].cells[ci]
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        rl = p.add_run(label + ': ')
        rl.bold = True; rl.font.size = Pt(8); rl.font.color.rgb = GRAY_TEXT
        rv = p.add_run(val)
        rv.font.size = Pt(9); rv.font.color.rgb = tc
    set_col_widths(meta, [Inches(1.5), Inches(3.2), Inches(1.8)])

    # Steps table — 2 columns
    steps_tbl = doc.add_table(rows=1, cols=2)
    steps_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Column headers
    for ci, header in enumerate(['Data & Backend Steps', 'UI & View Steps']):
        hh = steps_tbl.rows[0].cells[ci]
        set_cell_bg(hh, GRAY_BG)
        p = hh.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        r = p.add_run(header)
        r.bold = True; r.font.size = Pt(9); r.font.color.rgb = GRAY_TEXT

    max_steps = max(len(data_steps), len(ui_steps))
    for i in range(max_steps):
        row = steps_tbl.add_row()
        for ci, steps in enumerate([data_steps, ui_steps]):
            cell = row.cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if i < len(steps):
                step_num, step_title, step_desc = steps[i]
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(5)
                p.paragraph_format.space_after  = Pt(2)
                p.paragraph_format.left_indent  = Inches(0.2)
                rn = p.add_run(f'Step {step_num}:  ')
                rn.bold = True; rn.font.size = Pt(9); rn.font.color.rgb = EMERALD
                rt = p.add_run(step_title + '\n')
                rt.bold = True; rt.font.size = Pt(9.5); rt.font.color.rgb = DARK_NAVY
                p2 = cell.add_paragraph(step_desc)
                p2.paragraph_format.left_indent  = Inches(0.2)
                p2.paragraph_format.space_after  = Pt(6)
                p2.runs[0].font.size = Pt(9); p2.runs[0].font.color.rgb = GRAY_TEXT

    set_col_widths(steps_tbl, [Inches(3.25), Inches(3.25)])

    if tech_notes:
        tn_tbl = doc.add_table(rows=1, cols=1)
        tn_c = tn_tbl.rows[0].cells[0]
        set_cell_bg(tn_c, RGBColor(0xf0, 0xf9, 0xf4))
        tnp = tn_c.paragraphs[0]
        tnp.paragraph_format.space_before = Pt(5)
        rtnl = tnp.add_run('TECHNICAL NOTES: ')
        rtnl.bold = True; rtnl.font.size = Pt(8.5); rtnl.font.color.rgb = EMERALD
        for note in tech_notes:
            np = tn_c.add_paragraph(f'→  {note}')
            np.paragraph_format.left_indent = Inches(0.15)
            np.paragraph_format.space_after = Pt(3)
            np.runs[0].font.size = Pt(9); np.runs[0].font.color.rgb = MID_GRAY

    doc.add_paragraph()


def build_roadmap():
    doc = Document()
    set_doc_margins(doc)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ── COVER ────────────────────────────────────────────────────────────────
    cover = doc.add_table(rows=1, cols=1)
    cover.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = cover.rows[0].cells[0]
    set_cell_bg(c, DARK_NAVY)
    set_row_height(cover.rows[0], 8000)

    p = c.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    r = p.add_run('CASPAA')
    r.bold = True; r.font.size = Pt(36); r.font.color.rgb = WHITE

    p2 = c.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(20)
    r2 = p2.add_run('ENGINEERING ROADMAP  ·  INTERNAL')
    r2.font.size = Pt(9); r2.font.color.rgb = RGBColor(0x34, 0xd3, 0x99)

    p3 = c.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run('New Modules\nImplementation Roadmap')
    r3.bold = True; r3.font.size = Pt(28); r3.font.color.rgb = WHITE

    p4 = c.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_before = Pt(16)
    r4 = p4.add_run(
        'Step-by-step flow for implementing every new module identified in the Feature Gap Analysis.\n'
        'Covers data schema, UI flow, roles affected, and technical notes.'
    )
    r4.font.size = Pt(11); r4.font.color.rgb = RGBColor(0xb0, 0xc4, 0xde)

    p5 = c.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p5.paragraph_format.space_before = Pt(36)
    r5 = p5.add_run('Phase 1: Weeks 1–4     |     Phase 2: Weeks 5–10     |     Phase 3: Weeks 11–16+     |     14 New Modules')
    r5.font.size = Pt(9); r5.font.color.rgb = RGBColor(0x90, 0xa8, 0xc0)

    add_page_break(doc)

    # ── OVERVIEW & MODULE INDEX TABLE ─────────────────────────────────────────
    add_label(doc, '00 — Overview')
    add_heading(doc, 'Module Index — All 14 Modules at a Glance')
    add_paragraph(doc,
        '14 new modules across 3 phases. Phase 1 unlocks daily student/parent usage. '
        'Phase 2 completes operational coverage. Phase 3 delivers strategic differentiation.',
        size=10.5, color=GRAY_TEXT, space_after=10)

    idx = doc.add_table(rows=0, cols=6)
    idx.style = 'Table Grid'
    idx.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_colored_table_row(idx, ['#', 'Module', 'Phase', 'Effort', 'Primary File(s)', 'Roles'], DARK_NAVY, WHITE, size=9)
    rows_data = [
        ('1', 'Event Calendar',       'P1', 'Medium', 'admin.js, shared.js',   'All roles'),
        ('2', 'Bulk SMS & Email',     'P1', 'Medium', 'admin.js + backend',    'Admin, Finance'),
        ('3', 'Class Notes',          'P1', 'Low',    'teacher.js, student.js','Teacher, Student'),
        ('4', 'Video Lessons',        'P1', 'Low',    'teacher.js, student.js','Teacher, Student'),
        ('5', 'PhET Simulations',     'P1', 'Low',    'student.js, shared.js', 'Student, Teacher'),
        ('6', 'Payment Gateway',      'P1', 'High',   'finance.js, parent.js', 'Finance, Parent'),
        ('7', 'House Point System',   'P2', 'Low',    'admin.js, teacher.js',  'All roles'),
        ('8', 'Teacher Payslip',      'P2', 'Low',    'teacher.js',            'Teacher, Finance'),
        ('9', 'Transport Management', 'P2', 'High',   'admin.js, parent.js',   'Admin, Parent'),
        ('10','Parent Feedback',      'P2', 'Low',    'parent.js, admin.js',   'Parent, Admin'),
        ('11','Communication Diary',  'P2', 'Medium', 'teacher.js, parent.js', 'Teacher, Parent'),
        ('12','Formative Assessment', 'P2', 'Medium', 'teacher.js, student.js','Teacher, Student'),
        ('13','Health / Sickbay',     'P2', 'Medium', 'admin.js',              'Admin, Parent'),
        ('14','Inventory Management', 'P2', 'Medium', 'admin.js',              'Admin'),
    ]
    phase_colors = {'P1': (LIGHT_RED, RED), 'P2': (LIGHT_AMBER, AMBER), 'P3': (LIGHT_BLUE, BLUE)}
    effort_colors = {'Low': (LIGHT_MINT, EMERALD), 'Medium': (LIGHT_AMBER, AMBER), 'High': (LIGHT_RED, RED)}

    for i, (num, name, phase, effort, files, roles) in enumerate(rows_data):
        bg = GRAY_BG if i % 2 else WHITE
        dr = idx.add_row()
        for ci, val in enumerate([num, name, phase, effort, files, roles]):
            cell = dr.cells[ci]
            if ci == 2:
                pb, pt_ = phase_colors.get(phase, (bg, DARK_NAVY))
                set_cell_bg(cell, pb)
            elif ci == 3:
                eb, et_ = effort_colors.get(effort, (bg, DARK_NAVY))
                set_cell_bg(cell, eb)
            else:
                set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(val)
            run.bold = (ci == 1)
            run.font.size = Pt(9)
            if ci == 2:
                run.font.color.rgb = phase_colors.get(phase, (bg, DARK_NAVY))[1]
            elif ci == 3:
                run.font.color.rgb = effort_colors.get(effort, (bg, DARK_NAVY))[1]
            else:
                run.font.color.rgb = DARK_NAVY if ci == 1 else GRAY_TEXT

    set_col_widths(idx, [Inches(0.3), Inches(1.5), Inches(0.55), Inches(0.65), Inches(1.5), Inches(2.0)])
    doc.add_paragraph()

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════════════════
    # PHASE 1 MODULES
    # ════════════════════════════════════════════════════════════════════════

    add_label(doc, 'Phase 1 — Weeks 1–4')
    add_heading(doc, 'Core Content & Communication Layer')
    add_paragraph(doc,
        'These six modules directly impact daily usage by students, parents, and teachers. '
        'They unlock the content and communication layer that is currently CASPAA\'s biggest gap versus EDVES.',
        size=10.5, color=GRAY_TEXT, space_after=12)

    # Module 1 — Event Calendar
    add_module_block(doc, 1, 'Event Calendar', 'School-wide events visible across all roles',
        'Medium Effort', LIGHT_AMBER, ['Admin (Create)', 'Teacher (View)', 'Parent (View)', 'Student (View)'],
        data_steps=[
            (1, 'Add Events table to DB',
             'Fields: id, schoolId, title, description, date, endDate, type (holiday/exam/sports/meeting/other), targetAudience, createdBy.'),
            (2, 'Seed sample events',
             'Mid-term break, PTA meeting, Sports Day, End-of-term exams, Prize Giving Day.'),
            (3, 'COMPUTE.upcomingEvents(userId)',
             'Returns next 5 events relevant to the user\'s role for dashboard widget.'),
        ],
        ui_steps=[
            (4, 'Admin: Create & manage events',
             'New view_adm_calendar — month grid + list toggle. Add/Edit/Delete with audience targeting and color-coded types.'),
            (5, 'Teacher & Parent: Read-only calendar',
             'Add "Calendar" to teacher and parent menus. Month view with clickable event detail panel.'),
            (6, 'Dashboard widgets',
             'Add "Upcoming Events" card to all role dashboards showing next 3 events.'),
        ],
        tech_notes=[
            'Month grid: pure CSS 7-column grid populated by JS — no external calendar library needed.',
            'Color-code event types: holiday = red, exam = amber, sports = green, meeting = blue.',
            'Re-use existing modal pattern for Create/Edit event form.',
        ])

    # Module 2 — Bulk SMS & Email
    add_module_block(doc, 2, 'Bulk SMS & Email Notifications', 'Critical communication channel for Nigerian schools',
        'Medium Effort', LIGHT_AMBER, ['Admin (Send)', 'Finance (Fee alerts)'],
        data_steps=[
            (1, 'Integrate SMS provider',
             'Integrate Africa\'s Talking or Termii API. Store API key in school settings. Track SMS units per school.'),
            (2, 'Integrate Email provider',
             'Use Resend or SendGrid. Each school gets a sending identity (e.g. noreply@schoolname.caspaa.ng).'),
            (3, 'Add SMSLog & EmailLog tables',
             'Fields: id, schoolId, sentBy, recipients (count), message, channel, status, sentAt.'),
        ],
        ui_steps=[
            (4, 'Admin: SMS Compose view',
             'New view_adm_sms — choose audience (All Parents / All Teachers / Specific Class), compose, see unit cost, send.'),
            (5, 'SMS Log history',
             'Table of past sends: audience, message preview, units used, delivery rate, timestamp.'),
            (6, 'Auto-triggers',
             'Trigger SMS on: invoice generated, payment received, new announcement, exam results published, event reminder.'),
        ],
        tech_notes=[
            'Numbers on DND (Do Not Disturb) consume more SMS units — warn admin in the compose UI.',
            'Store remaining SMS balance in school settings; alert when balance < 50 units.',
            'For demo mode: simulate API call with 1.5s delay and mock success response.',
        ])

    # Module 3 — Class Notes
    add_module_block(doc, 3, 'Class Notes', 'Teachers post notes; students access them in the learning portal',
        'Low Effort', LIGHT_MINT, ['Teacher (Create)', 'Student (Read)', 'Admin (View all)'],
        data_steps=[
            (1, 'Add ClassNotes table',
             'Fields: id, schoolId, classId, subjectId, teacherId, title, content, attachmentUrl, term, week, createdAt.'),
            (2, 'Add NoteViews table',
             'Track which students have read each note: noteId, studentId, viewedAt. Used for teacher analytics.'),
        ],
        ui_steps=[
            (3, 'Teacher: Post & manage notes',
             'Add "Notes" tab to tch_lessons. Create note form: title, subject, class, content, optional file attach.'),
            (4, 'Student: Browse notes',
             'In stu_learning, add "Class Notes" tab. Filter by subject. Card list with "New" badge if unread. Mark as read on open.'),
            (5, 'Dashboard nudge',
             'Student dashboard shows "X new notes this week". Teacher dashboard shows note read rate per class.'),
        ])

    # Module 4 — Video Lessons
    add_module_block(doc, 4, 'Video Lessons', 'Embed video content per subject for student self-study',
        'Low Effort', LIGHT_MINT, ['Teacher (Post)', 'Student (Watch)', 'Parent (View list)'],
        data_steps=[
            (1, 'Add Videos table',
             'Fields: id, schoolId, classId, subjectId, teacherId, title, videoUrl (YouTube/Vimeo/upload), thumbnailUrl, duration.'),
            (2, 'URL parser utility',
             'Helper that detects YouTube/Vimeo links and auto-generates embed URLs and thumbnails from the video ID.'),
        ],
        ui_steps=[
            (3, 'Teacher: Upload/link videos',
             'Add "Videos" tab to tch_lessons. Paste YouTube/Vimeo URL or upload file. Preview thumbnail auto-loads.'),
            (4, 'Student: Video library',
             'In stu_learning, add "Video Lessons" tab. Grid of video cards. Click opens inline iframe player.'),
            (5, 'Watch progress tracking',
             'On video open, log VideoView record. Student dashboard shows "Continue Watching" row for incomplete videos.'),
        ])

    # Module 5 — PhET Simulations
    add_module_block(doc, 5, 'Maths & Science Simulations (PhET)', 'Embed free PhET simulations — a key EDVES selling point',
        'Low Effort', LIGHT_MINT, ['Student (Launch)', 'Teacher (Assign)'],
        data_steps=[
            (1, 'Create SimulationCatalog',
             'Static JS array of 30–50 curated PhET simulations: { id, name, subject, topic, gradeLevel, embedUrl, thumbnailUrl }.'),
            (2, 'Cover Maths & Sciences',
             'Include: Build an Atom, Wave on a String, Projectile Motion, Ohm\'s Law, Fraction Matcher, Balancing Chemical Equations, etc.'),
        ],
        ui_steps=[
            (3, 'Student: Simulations library',
             'New tab "Simulations" in stu_learning. Filter by subject. Card grid with thumbnail and "Launch" button.'),
            (4, 'Simulation player',
             'Full-screen modal with PhET simulation in an iframe. PhET provides embed-ready URLs — no API key required.'),
            (5, 'Teacher: Assign simulations',
             'Teacher can "assign" a simulation to a class — appears in student dashboard as "Simulation this week".'),
        ],
        tech_notes=[
            'PhET simulations at phet.colorado.edu are free, open source, and embeddable without authentication.',
            'Use direct simulation HTML embed URLs — no dependency on external API.',
            'Only catalogue HTML5 sims (not Java) — they work on tablets and mobile.',
        ])

    # Module 6 — Payment Gateway
    add_module_block(doc, 6, 'Online Payment Gateway', 'Connect fee payments to Paystack or Flutterwave',
        'High Effort', LIGHT_RED, ['Parent (Pay)', 'Finance (Receive)', 'Admin (Configure)'],
        data_steps=[
            (1, 'Payment provider config in School Settings',
             'Admin enters Paystack/Flutterwave public and secret key. Test mode toggle for sandbox testing.'),
            (2, 'Payment verification backend',
             'On callback, verify transaction reference with Paystack API. Only mark as successful after server-side verification.'),
            (3, 'Auto-create Transaction record',
             'On verified success: create Transaction with method=Card, amount, invoiceId, reference, status=Completed. Update invoice balance.'),
        ],
        ui_steps=[
            (4, 'Parent: Pay Now flow',
             'Parent clicks "Pay Now" on invoice → Paystack inline popup opens with amount pre-filled → on success, callback fires.'),
            (5, 'Send receipt SMS & email',
             'Auto-trigger SMS + email to parent with payment confirmation and receipt number. CC finance officer.'),
            (6, 'Finance dashboard update',
             'Online payments appear in fin_payments with a "Card" badge. Auto-marked as reconciled — no manual step.'),
        ],
        tech_notes=[
            'Use Paystack\'s inline JS for popup mode — no redirect, better UX.',
            'Paystack charges 1.5% per transaction (capped at ₦2,000) — show this to parents before payment.',
            'Store transaction reference in Transactions table for reconciliation and dispute resolution.',
        ])

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════════════════
    # PHASE 2 MODULES
    # ════════════════════════════════════════════════════════════════════════

    add_label(doc, 'Phase 2 — Weeks 5–10')
    add_heading(doc, 'Operational & Engagement Modules')
    add_paragraph(doc,
        'Eight modules that complete operational coverage and deepen parent/teacher engagement. '
        'Several are quick wins built on data that already exists in the system.',
        size=10.5, color=GRAY_TEXT, space_after=12)

    # Module 7 — House Points
    add_module_block(doc, 7, 'House Point System', 'Dedicated pastoral care module with school-wide leaderboard',
        'Low Effort', LIGHT_MINT, ['Admin', 'Teacher', 'Student', 'Parent'],
        data_steps=[
            (1, 'Add Houses table',
             'Fields: id, schoolId, name (e.g. "Red House"), color, emblem. Admin creates and assigns students to houses.'),
            (2, 'Add HousePoints table',
             'Fields: id, houseId, studentId, points, reason, awardedBy, awardedAt.'),
            (3, 'COMPUTE.houseLeaderboard(schoolId)',
             'Aggregate house points by house, return sorted leaderboard with house name, total points, leading student.'),
        ],
        ui_steps=[
            (4, 'Admin: Manage houses & award points',
             'New view_adm_houses. Create/edit houses, assign students, award bulk points for school events.'),
            (5, 'Teacher: Award individual points',
             'In tch_classes, "Award House Point" button per student. Select reason from dropdown.'),
            (6, 'Student & Parent: Leaderboard',
             'House leaderboard widget on student dashboard. Parent sees child\'s house, points, and rank this term.'),
        ])

    # Module 8 — Teacher Payslip
    add_module_block(doc, 8, 'Teacher Payslip (Self-View)', 'Teachers access their own monthly payslip — data already exists',
        'Low Effort — Quick Win', LIGHT_MINT, ['Teacher (Own only)', 'Finance (All staff)'],
        data_steps=[
            (1, 'Pull from existing Payroll data',
             'No new table needed. Query current teacher\'s records from existing Salary/Payroll data filtered by session.userId.'),
            (2, 'Privacy: own data only',
             'Query must filter strictly by session.userId — teacher cannot see other staff payslips.'),
        ],
        ui_steps=[
            (3, 'Add "Payslip" to teacher menu',
             'New view_tch_payslip. Term/month selector. Show: Gross, Deductions, Net Pay, Bank Account, Payment Date, Status.'),
            (4, 'Salary advance status',
             'If teacher has a pending advance, show it on payslip with status (Pending/Approved/Deducted).'),
            (5, 'PDF download',
             'Print/PDF button using window.print() with a payslip-specific print stylesheet. No external library needed.'),
            (6, 'Finance: Payslip history',
             'Add "View Payslip" action per row in fin_payroll. Finance officer can view/download any staff payslip.'),
        ])

    # Module 9 — Transport
    add_module_block(doc, 9, 'Transport Management & Secure Pickup', 'Bus routes, driver records, and approved pickup persons',
        'High Effort', LIGHT_RED, ['Admin (Manage)', 'Parent (Register pickup / Track)'],
        data_steps=[
            (1, 'Add Buses table',
             'Fields: id, schoolId, plateNumber, capacity, driverName, driverPhone, route, status.'),
            (2, 'Add Routes table',
             'Fields: id, schoolId, name, stops (JSON array of {name, time}), busId.'),
            (3, 'Add PickupPersons table',
             'Fields: id, studentId, name, relationship, phone, photo, isDefault. Max 3 per child.'),
        ],
        ui_steps=[
            (4, 'Admin: Manage buses & routes',
             'Complete view_adm_transport. Add buses, define routes and stops, assign students, view daily boarding manifest.'),
            (5, 'Parent: Register pickup persons',
             'In par_children, "Authorized Pickup" tab. Add persons with photo, name, relationship. Admin can verify.'),
            (6, 'Parent: Bus tracker',
             'View child\'s bus route and expected arrival time at each stop. "Bus En Route" status alert (manual trigger by admin).'),
        ])

    # Module 10 — Parent Feedback
    add_module_block(doc, 10, 'Parent Feedback Tool', 'Structured school satisfaction surveys and teacher ratings',
        'Low Effort', LIGHT_MINT, ['Parent (Submit)', 'Admin (Create / Analyze)'],
        data_steps=[
            (1, 'Add FeedbackForms table',
             'Fields: id, schoolId, title, questions (JSON: {question, type: rating/text/yesno}), targetTerm, status.'),
            (2, 'Add FeedbackResponses table',
             'Fields: id, formId, parentId, answers (JSON), submittedAt. Anonymous to teachers; visible to admin.'),
        ],
        ui_steps=[
            (3, 'Admin: Create & publish forms',
             'New view_adm_feedback. Build form with questions. Publish with deadline. View aggregate results (charts).'),
            (4, 'Parent: Complete feedback',
             'New "Feedback" item in parent menu. Star ratings, text fields. Submit once per form per child.'),
            (5, 'Parent dashboard nudge',
             'Alert card: "Term feedback form is open" with link. Disappears after submission.'),
        ])

    # Module 11 — Communication Diary
    add_module_block(doc, 11, 'Communication Diary', 'Structured daily teacher–parent notes with read receipts',
        'Medium Effort', LIGHT_AMBER, ['Teacher (Write)', 'Parent (Read & Reply)', 'Admin (View all)'],
        data_steps=[
            (1, 'Add DiaryEntries table',
             'Fields: id, studentId, teacherId, date, note, category (Homework/Behaviour/Health/General), readByParent, parentReply, replyAt.'),
            (2, 'Distinguish from Messaging',
             'Diary = structured per-student per-day entries. Messaging = free-form chat. Both coexist.'),
        ],
        ui_steps=[
            (3, 'Teacher: Write diary entries',
             'In tch_classes, "Diary" tab per class. Select student, date, category, write note. See unread count badge per student.'),
            (4, 'Parent: Read & reply',
             'New "Diary" section in parent portal. Per-child feed, newest first. Tap to expand, reply inline. Read receipt on open.'),
            (5, 'SMS alert on new entry',
             '"New diary note from Mrs. X for [Child] — log in to read and reply." Uses Bulk SMS module (Module 2).'),
        ])

    # Module 12 — Formative Assessment
    add_module_block(doc, 12, 'Formative Assessment', 'In-class ongoing quizzes, distinct from end-of-term CBT',
        'Medium Effort', LIGHT_AMBER, ['Teacher (Create & Score)', 'Student (View)', 'Admin (Overview)'],
        data_steps=[
            (1, 'Add FormativeAssessments table',
             'Fields: id, classId, subjectId, teacherId, title, type (ClassTest/Quiz/Oral/Project), maxScore, term, week, date.'),
            (2, 'Add FormativeScores table',
             'Fields: id, assessmentId, studentId, score, remarks, gradedAt. Separate from CBTSubmissions (computer-graded).'),
        ],
        ui_steps=[
            (3, 'Teacher: Create & score',
             'New "Formative" tab in tch_results. Create assessment (type, max score). Enter scores per student in a row view.'),
            (4, 'Student: View formative scores',
             'In stu_results, add "Assessments" tab showing all formative scores by week with class average for comparison.'),
            (5, 'Include in terminal report',
             'Formative scores contribute to cumulative CA score visible on the end-of-term report card.'),
        ])

    # Module 13 — Health / Sickbay
    add_module_block(doc, 13, 'Health Management / Sickbay', 'Complete the stub: student health records and incident tracking',
        'Medium Effort', LIGHT_AMBER, ['Admin (Log visits)', 'Parent (View + Update profile)'],
        data_steps=[
            (1, 'Add HealthProfiles table',
             'Fields: id, studentId, bloodGroup, allergies (array), conditions (array), emergencyContact, doctorName, vaccinations.'),
            (2, 'Add SickbayVisits table',
             'Fields: id, studentId, dateIn, dateOut, complaint, treatment, medications, referredOut, parentNotified.'),
        ],
        ui_steps=[
            (3, 'Admin: Sickbay log',
             'Complete view_adm_sickbay stub. Check student in, record complaint and treatment, check out, view history.'),
            (4, 'Parent notification on visit',
             'On check-in, auto-send SMS: "Your child [Name] visited the sickbay at [time]. Complaint: [brief]."'),
            (5, 'Parent: View health record',
             'In par_children, add "Health" tab showing health profile and sickbay visit history. Parent can update profile.'),
        ])

    # Module 14 — Inventory
    add_module_block(doc, 14, 'Inventory Management', 'Complete the stub: items, stock levels, and issue tracking',
        'Medium Effort', LIGHT_AMBER, ['Admin (Full access)'],
        data_steps=[
            (1, 'Add InventoryItems table',
             'Fields: id, schoolId, name, category (Furniture/Electronics/Stationery/Equipment/Uniforms), quantity, reorderLevel, location.'),
            (2, 'Add InventoryTransactions table',
             'Fields: id, itemId, type (Stock-In/Issue/Return/Write-Off), quantity, issuedTo, notes, recordedBy, date.'),
        ],
        ui_steps=[
            (3, 'Admin: Item catalogue',
             'Complete view_adm_inventory. Table of all items. Color-coded: red if below reorder level. Quick-add restock.'),
            (4, 'Issue & return flow',
             'Issue item to staff/student: deducts from stock, logs transaction with recipient. Return adds back. All audit-logged.'),
            (5, 'Low stock alerts',
             '"3 items below reorder level" dashboard alert linking to inventory view.'),
        ])

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════════════════
    # PHASE 3 + DEPENDENCY MAP
    # ════════════════════════════════════════════════════════════════════════

    add_label(doc, 'Phase 3 — Weeks 11–16+')
    add_heading(doc, 'Strategic & AI Differentiation Modules')
    add_paragraph(doc,
        'Phase 3 modules require deeper engineering investment or third-party integrations. '
        'Plan now, build after Phase 1 & 2 are stable.',
        size=10.5, color=GRAY_TEXT, space_after=10)

    p3_data = [
        ('Virtual Classroom', 'Embed Jitsi Meet or Zoom SDK for live classes. Schedule sessions linked to timetable.', 'Needed for hybrid/remote school models growing post-COVID.', '4–6 weeks'),
        ('AI Grading & Question Generation', 'Integrate Claude API. Teacher selects lesson note → AI suggests 10 CBT questions. AI marks short-answer questions.', 'Reduces teacher workload by ~40%. Clear differentiator vs. EDVES.', '6–8 weeks'),
        ('Student Performance Prediction', 'Rule-based or ML model flagging students at risk: declining attendance, falling grades, missed assignments.', 'Early intervention tool. Schools pay premium for this insight.', '6–8 weeks'),
        ('Hostel Management', 'Boarding schools only. Rooms, bedspace assignment, check-in/out, meals schedule, parent visit authorization.', 'Unlocks boarding school market. Currently zero boarding support in CASPAA.', '4–5 weeks'),
        ('Multi-Curriculum Support', 'Tag lessons, notes, and assessments with curriculum standard (NERDC / Cambridge / IGCSE / IB).', 'Required for international schools. EDVES Cambridge/IGCSE support is a key differentiator in Lagos Island.', '3–4 weeks'),
        ('Offline CBT', 'Service Worker + IndexedDB to cache exam questions. Students take exam offline; answers sync on reconnect.', 'Critical for schools with unreliable internet. Blocks adoption in rural and peri-urban areas.', '6–8 weeks'),
    ]

    p3_tbl = doc.add_table(rows=0, cols=4)
    p3_tbl.style = 'Table Grid'
    p3_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_colored_table_row(p3_tbl, ['Module', 'What It Involves', 'Why It Matters', 'Effort'], DARK_NAVY, WHITE, size=9)
    for i, (name, involves, why, effort) in enumerate(p3_data):
        bg = GRAY_BG if i % 2 else WHITE
        dr = p3_tbl.add_row()
        for ci, (val, col) in enumerate([
            (name, DARK_NAVY), (involves, GRAY_TEXT), (why, GRAY_TEXT), (effort, AMBER)
        ]):
            set_cell_bg(dr.cells[ci], bg)
            dr.cells[ci].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = dr.cells[ci].paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            run = p.add_run(val)
            run.bold = (ci == 0)
            run.font.size = Pt(9)
            run.font.color.rgb = col
    set_col_widths(p3_tbl, [Inches(1.4), Inches(2.3), Inches(2.0), Inches(0.8)])

    doc.add_paragraph()
    add_label(doc, 'Appendix — Build Order')
    add_heading(doc, 'Dependency Map')
    add_paragraph(doc, 'Some modules depend on others. Follow this order to avoid rework.', size=10.5, color=GRAY_TEXT, space_after=8)

    deps = [
        ('P1', 'Bulk SMS & Email must come FIRST',
         'Communication Diary (11), Health/Sickbay (13), Transport (9), and Payment Gateway (6) all send SMS/email alerts. Build Module 2 before those modules.'),
        ('P1', 'Class Notes before Video Lessons',
         'Both live in stu_learning. Build the Notes tab structure first; Video Lessons tab slots in alongside it. Shared learning hub UX is established once.'),
        ('P2', 'Formative Assessment before Multi-Curriculum (P3)',
         'Formative Assessment data model (week-based scoring) becomes the foundation for curriculum-tagged assessment in Phase 3.'),
        ('P2', 'House Point System is fully independent',
         'No dependencies on other new modules. Safe to start any time in Phase 2 in parallel with other work.'),
        ('P3', 'AI features depend on Class Notes and Results data',
         'AI question generation reads from ClassNotes content. Performance prediction reads from Results + Attendance — both must be populated with real data first.'),
    ]
    for phase, title, desc in deps:
        pc = RED if phase == 'P1' else (AMBER if phase == 'P2' else BLUE)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(2)
        rp = p.add_run(f'[{phase}]  ')
        rp.bold = True; rp.font.size = Pt(9.5); rp.font.color.rgb = pc
        rt = p.add_run(title)
        rt.bold = True; rt.font.size = Pt(10); rt.font.color.rgb = DARK_NAVY
        pd = doc.add_paragraph(desc)
        pd.paragraph_format.left_indent = Inches(0.3)
        pd.paragraph_format.space_after = Pt(6)
        pd.runs[0].font.size = Pt(9.5); pd.runs[0].font.color.rgb = GRAY_TEXT

    doc.save(r'C:\Users\USER\Desktop\CASPAA\CASPAA_Implementation_Roadmap.docx')
    print('Done: Implementation Roadmap document saved.')


if __name__ == '__main__':
    build_gap_analysis()
    build_roadmap()
    print('Both Word documents generated successfully.')
