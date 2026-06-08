# -*- coding: utf-8 -*-
"""Generate the refreshed CASPAA PRD (v1.1) as a .docx."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BRAND = RGBColor(0x1F, 0x4E, 0x79)      # deep blue
ACCENT = RGBColor(0x2E, 0x75, 0xB6)     # mid blue
HDR_FILL = "1F4E79"
SUBHDR_FILL = "DCE6F1"
GREY = RGBColor(0x59, 0x59, 0x59)

doc = Document()

# ---- base styles ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.12

for lvl, sz, col in [("Heading 1", 16, BRAND), ("Heading 2", 13, BRAND), ("Heading 3", 11.5, ACCENT)]:
    st = doc.styles[lvl]
    st.font.name = "Calibri"
    st.font.size = Pt(sz)
    st.font.color.rgb = col
    st.font.bold = True
    st.paragraph_format.space_before = Pt(12)
    st.paragraph_format.space_after = Pt(4)


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), fill)
    tcPr.append(sh)


def set_cell_text(cell, text, bold=False, white=False, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = "Calibri"
    if white:
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def add_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    t.autofit = True
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, white=True)
        shade(hdr[i], HDR_FILL)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val)
    if widths:
        for r_ in t.rows:
            for i, w in enumerate(widths):
                r_.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


def para(text, italic=False, color=None, size=None, bold=False, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    if color:
        r.font.color.rgb = color
    if size:
        r.font.size = Pt(size)
    return p


def bullets(items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(it)


def hr():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "BFBFBF")
    pbdr.append(bottom)
    pPr.append(pbdr)


# ============================ COVER ============================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(60)
r = title.add_run("PRODUCT REQUIREMENTS DOCUMENT")
r.bold = True
r.font.size = Pt(24)
r.font.color.rgb = BRAND

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("CASPAA — School Operating System (MVP)")
r.bold = True
r.font.size = Pt(15)
r.font.color.rgb = ACCENT

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("EdTech × FinTech for African Schools")
r.italic = True
r.font.size = Pt(11)
r.font.color.rgb = GREY

for label, value in [("Version", "1.2"), ("Status", "MVP — Working Prototype Validated + Pilot Addendum"),
                     ("Date", "7 June 2026"), ("Prepared by", "CASPAA / Afrisprings")]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    rl = p.add_run(f"{label}:  ")
    rl.bold = True
    rl.font.size = Pt(10)
    rv = p.add_run(value)
    rv.font.size = Pt(10)

doc.add_page_break()

# ============================ EXECUTIVE SUMMARY ============================
doc.add_heading("1. Executive Summary", level=1)

doc.add_heading("Product Overview", level=2)
para("CASPAA is a mobile-first School ERP and embedded-finance platform built for nursery, primary, "
     "and secondary schools across Africa, beginning with Lagos, Nigeria. It unifies school management, "
     "parent–teacher communication, payment infrastructure, school-fee financing, financial analytics, "
     "and AI-powered education workflows into a single operating system for schools.")

doc.add_heading("Product Vision", level=2)
para("To become the default operating system powering education finance and school operations across Africa.")

doc.add_heading("Implementation Status", level=2)
para("A fully clickable, end-to-end working prototype now exists, covering all 10 MVP modules across 5 user roles. "
     "It validates the product logic and user experience ahead of the production build. The prototype runs entirely "
     "in the browser with a LocalStorage-backed mock data layer; the remaining work to reach production is integrating "
     "real services (Paystack, WhatsApp Business API, SMS gateway, and a managed PostgreSQL backend). "
     "Module-by-module status is summarised in Section 5.")

# ============================ PROBLEM & OBJECTIVES ============================
doc.add_heading("2. Problem Statement", level=1)
para("Most schools in Africa still run on fragmented, manual systems. The recurring problems include:")
bullets([
    "Fragmented and manual operations across admissions, academics, and finance",
    "Weak parent-engagement infrastructure",
    "Cash-based fee collection",
    "Poor financial visibility and reconciliation gaps",
    "Unstable revenue caused by school-fee payment gaps",
    "Financial leakage and poor fee tracking",
    "Outdated or unreliable financial statements",
    "No integrated EdTech–FinTech infrastructure — existing solutions are siloed",
])
para("These problems compound into real operational pain for schools:")
bullets([
    "Revenue leakage",
    "Administrative inefficiency",
    "Poor parent engagement",
    "Delayed fee payments",
    "Lack of access to education financing",
])

doc.add_heading("3. MVP Objectives", level=1)
para("The objective is to launch quickly, validate market demand, and establish the financial-infrastructure "
     "layer before expanding into more advanced fintech products.")
add_table(
    ["Objective", "Description"],
    [
        ["Digitize school operations", "Attendance, academics, and finance in one place"],
        ["Improve fee payment & collection", "Online payments and automated reminders"],
        ["Increase parent engagement", "Messaging and progress tracking, with WhatsApp adaptability"],
        ["Enable school-fee financing", "Embedded lending"],
        ["Offline capability", "Synchronisation once internet connectivity is restored"],
        ["Create recurring revenue", "SaaS plus fintech monetisation"],
        ["Validate product-market fit", "Lagos pilot schools"],
    ],
    widths=[2.4, 4.0],
)

# ============================ TARGET USERS ============================
doc.add_heading("4. Target Users", level=1)
add_table(
    ["User Type", "Description"],
    [
        ["School Admin", "Proprietors and principals"],
        ["Finance / Admin Officer", "Accounts and billing"],
        ["Teachers", "Classroom operations"],
        ["Parents / Guardians", "Payments and child monitoring"],
        ["Students", "Assignments and results"],
        ["Super Admin", "Platform operator"],
    ],
    widths=[2.4, 4.0],
)

# ============================ MVP MODULES ============================
doc.add_heading("5. MVP Modules", level=1)
para("All 10 modules below are implemented end-to-end in the working prototype. Each module lists its "
     "objective, feature set, and current build status.", italic=True, color=GREY)

def module(num, name, objective, feature_groups, status="Implemented (prototype)"):
    doc.add_heading(f"Module {num} — {name}", level=2)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Objective: ")
    r.bold = True
    p.add_run(objective)
    s = doc.add_paragraph()
    s.paragraph_format.space_after = Pt(4)
    sr = s.add_run("Status: ")
    sr.bold = True
    sv = s.add_run(status)
    sv.font.color.rgb = RGBColor(0x1E, 0x7A, 0x34)
    sv.bold = True
    for group, feats in feature_groups:
        if group:
            doc.add_heading(group, level=3)
        bullets(feats)

module(1, "School ERP (Core)", "Centralise all school-operations management.", [
    ("Student Information System (SIS)", ["Student registration and profile management", "Admission workflows",
        "Class assignment", "Student promotion", "Transfer management", "Alumni records",
        "Bulk student upload (Excel/CSV)"]),
    ("Staff Management", ["Staff records", "Teacher assignment", "Payroll records",
        "Staff attendance", "Leave management", "Staff performance tracking"]),
    ("Timetable Management", ["Class timetable", "Exam timetable", "Teacher timetable",
        "Timetable conflict detection"]),
    ("Result & Academic Management", ["CA and exam score entry", "Automated grading",
        "Broadsheet generation", "Result approval workflow", "Transcript generation", "Academic analytics"]),
    ("Attendance Management", ["Daily and bulk attendance", "Teacher attendance",
        "Late-resumption tracking", "Attendance reports", "Parent notifications"]),
    ("Discipline & Behaviour", ["Commendations and reward points", "Discipline history",
        "Misconduct reports", "Parent alerts"]),
    ("Inventory Management", ["Books, stationery, and uniforms", "Stock alerts",
        "Asset tracking"]),
])

module(2, "Parent App (Web + Mobile)", "Improve parent engagement, fee payments, and loan financing.", [
    ("Child Dashboard", ["Attendance summary", "Academic performance", "Upcoming events",
        "Notifications", "Multi-child overview", "Announcements"]),
    ("Fees & Billing", ["View school fees", "Pay online in the app", "Download receipts",
        "Instalment tracking", "Payment and auto reminders", "Reconciliation status"]),
    ("Communication", ["Chat with teachers", "Receive announcements", "PTA updates",
        "Push notifications", "WhatsApp adaptability (for ease of communication)"]),
    ("Academics", ["View report cards", "Assignment tracking", "Teacher comments",
        "Performance analytics", "Behavioural analytics"]),
    ("Transport", ["Pickup / drop-off information"]),
    ("Lending / Credit Facility", ["Apply for school-fee loans", "View repayment plans",
        "Loan status tracking", "Auto-debit setup", "Repayment history"]),
])

module(3, "Teacher App (Web + Mobile)", "Digitise classroom operations.", [
    ("Dashboard", ["Assigned classes", "Subjects", "Timetable", "Class size"]),
    ("Attendance", ["Mark attendance (works offline)", "Bulk attendance", "Attendance reports"]),
    ("Result Management", ["Enter CA and exam scores", "Auto calculations and grade auto-apply",
        "Comment entry", "Result submission and approval workflow"]),
    ("Assignments & Homework", ["Create assignments", "Upload files", "Set deadlines",
        "Assignment analytics"]),
    ("Communication", ["Message parents", "Broadcast announcements", "Chat with admin"]),
    ("Lesson Planning", ["Lesson notes", "Weekly plans", "Scheme of work", "Topic tracking"]),
])

module(4, "Financial Management", "Become the financial operating system for schools.", [
    ("Fees & Billing", ["Fee-structure setup", "Class-based fees", "Discounts and scholarships",
        "Invoice generation", "Instalment plans", "Multi-child billing"]),
    ("Revenue Management", ["Revenue dashboard", "Outstanding fees", "Debtors list", "Cashflow analytics"]),
    ("Expense Management", ["Expense recording", "Budgeting", "P&L reporting", "Cashbook"]),
    ("Payment Reconciliation", ["Auto reconciliation", "Automated payment matching", "Failed-payment tracking"]),
    ("Financial Reporting", ["Income statements", "Revenue analytics", "Audit-ready reports"]),
])

module(5, "Payments Integration (Paystack)", "Provide reliable, multi-channel fee collection.", [
    ("", ["Card payments", "Bank transfer", "USSD", "Wallet funding",
        "Transaction verification", "Webhook handling", "Receipt generation"]),
], status="Prototype flow complete — requires live Paystack keys for production")

module(6, "Lending Engine", "Provide education financing.", [
    ("Loan Types", ["School-fee loan (parents)", "Salary advance (teachers)", "Working capital (schools)"]),
    ("Features", ["Loan application", "Eligibility scoring", "Repayment schedules",
        "Interest calculation", "Loan monitoring", "Default tracking"]),
])

module(7, "Risk Engine", "Manage lending risk.", [
    ("Credit Scoring — inputs", ["Payment history", "Salary pattern", "School relationship",
        "Attendance / payment consistency"]),
    ("Fraud Detection", ["Suspicious transactions", "Duplicate identities", "Payment anomalies"]),
    ("Default Prediction", ["AI risk scoring", "Early-warning alerts"]),
])

module(8, "AI Assistant", "Automate school workflows.", [
    ("Features", ["AI-generated report-card comments", "FAQ automation"]),
    ("Predictive Analytics", ["Dropout risk", "Fee-default risk", "Performance decline"]),
    ("Personalised Learning", ["Recommendations based on student performance"]),
])

module(9, "Communication System", "Deliver WhatsApp-like school communication.", [
    ("", ["Teacher ↔ Parent chat", "Broadcast messaging", "Push notifications",
        "SMS integration", "Email notifications"]),
])

module(10, "Admin Portal / Dashboard", "Platform-wide management.", [
    ("School Management", ["Create schools", "Suspend schools", "Manage subscriptions"]),
    ("Revenue Monitoring", ["SaaS revenue", "Transaction revenue", "Lending analytics"]),
    ("System Control", ["Feature toggles", "Permission management", "Audit logs"]),
    ("Business Intelligence", ["School-growth analytics", "Usage analytics", "Financial insights"]),
])

doc.add_heading("Offline-First Capability", level=2)
para("Objective: support unreliable-internet environments.", bold=True)
bullets(["Local device caching", "Offline attendance", "Offline assignment creation",
         "Background sync", "Retry queues", "Teachers can mark attendance without internet"])

# ============================ USER STORIES ============================
doc.add_page_break()
doc.add_heading("6. User Stories & Acceptance Criteria", level=1)
para("User stories are organised by module so that every MVP module has explicit, testable coverage. Each story "
     "follows the 'As a [role], I want [capability] so that [outcome]' format, followed by the step-by-step user "
     "Flow through the actual views and an Acceptance Criteria table (criterion → expected outcome). Flows and "
     "criteria reflect the behaviour implemented in the validated prototype — including the specific screens, "
     "modals, and computed values a user encounters.", italic=True, color=GREY)


def mod_group(title):
    doc.add_heading(title, level=2)


def story(code, title, role, want, so_that, rows, flow=None):
    doc.add_heading(f"{code}  {title}", level=3)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.add_run("User Story: ").bold = True
    p.add_run(f"As a {role}, I want to {want} so that {so_that}")
    if flow:
        fp = doc.add_paragraph()
        fp.paragraph_format.space_after = Pt(2)
        fr = fp.add_run("Flow")
        fr.bold = True
        fr.font.color.rgb = ACCENT
        for i, step in enumerate(flow, 1):
            s = doc.add_paragraph(style="List Number")
            s.paragraph_format.space_after = Pt(1)
            s.paragraph_format.line_spacing = 1.05
            s.add_run(step)
        cap = doc.add_paragraph()
        cap.paragraph_format.space_after = Pt(2)
        cr = cap.add_run("Acceptance Criteria")
        cr.bold = True
        cr.font.color.rgb = ACCENT
    add_table(["Acceptance Criteria", "Expected Outcome"], rows, widths=[3.1, 3.3])


# ---------------- MODULE 1 — SCHOOL ERP ----------------
mod_group("Module 1 — School ERP (Core)")

story("US-1.1", "Student Information System (SIS)", "School Admin",
      "register and manage student records",
      "I can maintain centralised academic and administrative information.",
      [["Add Student form validates required fields", "Name, admission no., DOB, gender, class, parent required"],
       ["Admission number auto-generated", "Format BL/2024/NNN pre-filled, editable"],
       ["Student photo upload", "JPG/PNG up to 1MB, initials shown as fallback"],
       ["Supporting documents attached", "Birth certificate, parent ID, passport, immunisation card"],
       ["Parent linked from existing records", "Parent dashboard immediately reflects the child"],
       ["People Hub advanced filters", "Gender, payment status, attendance band, status, scholarship"],
       ["Bulk Excel/CSV upload with template", "Template downloadable; created / replaced / skipped counts and per-row errors returned"],
       ["CSV export of full register", "Admission, class, fees, attendance columns exported"],
       ["Lifecycle actions available", "Promote, transfer, suspend, withdraw, graduate-to-alumni"],
       ["'View as Parent' switch", "Admin previews the parent's dashboard for that child"],
       ["Every change audit-logged", "Action recorded with actor and timestamp"]],
      flow=[
          "Sign in as School Proprietor → open People Hub → Students tab.",
          "Filter/search the register (class, gender, payment, attendance, scholarship) or click Add Student.",
          "Complete the Add Student form: photo, bio-data, class & arm, parent link, documents → Save Student.",
          "To import many at once: Bulk Upload → download CSV template → upload filled file → review created/skipped/error report.",
          "Open any student card to view profile (attendance %, avg score, balance, recent results) or run a lifecycle action (Promote / Transfer / Suspend / Withdraw / Graduate).",
          "Use CSV Export to download the filtered register.",
      ])

story("US-1.2", "Staff Management & HR", "School Admin",
      "manage staff records, assignments, leave, and attendance",
      "teaching and administrative resources are organised and accountable.",
      [["Add Staff with staff type", "Academic / Finance / Administration / Operations / ICT / Transport"],
       ["Academic staff get subjects & classes", "Multi-select assignments reflected on teacher dashboard"],
       ["Per-staff permission toggles set", "7 module permissions configurable at creation"],
       ["Auto-generated login credentials", "Username + temp password sent via Email + WhatsApp"],
       ["Directory split academic / non-academic", "Non-academic grouped by department"],
       ["Performance panel for teachers", "Avg score, pass rate, punctuality, assignment count"],
       ["Leave requests with approval workflow", "Pending → approved/rejected; staff notified"],
       ["Substitute coverage suggested on approval", "Candidate teachers ranked by subject overlap"],
       ["Staff attendance / clock-in records", "Self clock-in/out or admin gate entry; late auto-flagged"]],
      flow=[
          "Open Staff & HR Hub → Staff Directory → Add Staff.",
          "Pick staff type; for academic staff assign subjects and classes; set salary, banking, and permission toggles → Save Staff.",
          "System auto-issues username + temporary password over Email + WhatsApp (confirmation modal).",
          "Switch to HR Hub → Leave Requests to approve/reject leave; on approval, optionally assign a ranked substitute.",
          "Use HR Hub → Staff Attendance to view clock-in records or mark attendance from the gate sign-in book.",
      ])

story("US-1.3", "Timetable Management", "School Admin",
      "build class, exam, and teacher timetables with conflict-free periods",
      "lessons are scheduled without clashes and everyone sees their own schedule.",
      [["Period created from an empty grid cell", "Add Period modal: day, period 1–8, time, subject, teacher"],
       ["Default period times pre-filled", "Each period shows its standard time, editable per cell"],
       ["Teacher double-booking blocked", "Saving a clashing teacher/day/period is rejected with a conflict toast"],
       ["Filled cell editable or removable", "Edit Timetable Cell modal: change time/subject/teacher or Remove"],
       ["Build Week quick-fill grid", "Periods 1–4 × Mon–Fri filled in one pass; empty cells cleared"],
       ["Whole-School View", "All classes in one table (first periods per day) for an overview"],
       ["Bulk CSV upload across classes", "Template with Class, Day, Period, Time, Subject, Teacher Email; replaces existing"],
       ["Teacher timetable auto-derived", "Each teacher's My Schedule shows day/week/month views"]],
      flow=[
          "Open Academic Hub → Timetable tab → choose a class from the selector.",
          "Click an empty cell → Add Period modal → pick day, period (1–8), time, subject, and teacher → Add Period.",
          "If the chosen teacher already teaches that day/period, the save is blocked with 'Conflict: teacher is already teaching [Day] P[Period]'.",
          "Click a filled cell to edit its time/subject/teacher or remove the period.",
          "Use Build Week to fill the whole grid at once, or Bulk Upload CSV to load many classes from a spreadsheet.",
          "Switch to Whole-School View to scan every class; teachers see the result under My Schedule (Day/Week/Month).",
      ])

story("US-1.4", "Curriculum & Scheme of Work", "School Admin",
      "maintain schemes of work and track weekly coverage",
      "teaching stays aligned to the curriculum and coverage is visible.",
      [["Scheme created per class & subject", "Source tagged NERDC / WAEC / Custom"],
       ["NERDC template import", "Scheme generated from uploaded template"],
       ["Week-by-week editor", "Topic, sub-topics, objective, duration, methods, resources per week"],
       ["Coverage tracked", "Weeks marked covered update progress bars and coverage rate"],
       ["Overview / By Class / By Subject tabs", "Coverage rate colour-coded (red < 50%)"],
       ["Scheme links to lesson plans", "Marking a tied lesson plan covers that scheme week"],
       ["Scheme exportable", "Export to PDF available"]],
      flow=[
          "Open Academic Hub → Curriculum → New Scheme (or Import NERDC Template).",
          "Select class, subject, term, and source; add weeks with topics and objectives.",
          "Track progress via Overview / By Class / By Subject tabs and the coverage-rate cards.",
          "Open a scheme to tick covered weeks, edit week details, or Export PDF.",
      ])

story("US-1.5", "Result & Academic Management", "Teacher / Admin",
      "enter scores, auto-grade, approve, and publish broadsheets",
      "result processing is faster, consistent, and error-free.",
      [["Score entry capped to component maxima", "CA1 /20, CA2 /20, Exam /60 auto-capped on input"],
       ["Total and grade recalculated live", "Total and colour-coded grade badge update as you type"],
       ["Comments per student", "Manual entry or AI suggestion per row"],
       ["Submitted results require approval", "Saved as unapproved; admin approves before publish"],
       ["Admin broadsheet by class", "Grid of subjects, average, and position/rank"],
       ["Pending approvals surfaced", "'Approve N pending' action on the broadsheet"],
       ["Broadsheet / report cards exportable", "Printable PDF (report card, transcript, broadsheet)"]],
      flow=[
          "Teacher opens Results → selects class and subject.",
          "Enters CA1, CA2, Exam per student; total and grade compute instantly (scores auto-capped).",
          "Adds a comment per student or uses AI: Suggest (single or all) → Submit for approval.",
          "Admin opens Academic Hub → Results (broadsheet) → reviews and Approve pending.",
          "Admin or parent exports the broadsheet / report card / transcript as PDF.",
      ])

story("US-1.6", "Attendance Management", "Teacher / Admin",
      "mark attendance quickly, offline-capable, and notify parents",
      "the school tracks presence and parents are alerted to absences.",
      [["Per-student Present / Late / Absent", "Colour-coded buttons; in-memory buffer for speed"],
       ["Bulk mark all present / absent", "Whole class set in one tap"],
       ["Works offline", "Offline badge shown; entries queued and synced on reconnect"],
       ["Save & Notify triggers alerts", "Absence and late notices sent to parents via WhatsApp + app"],
       ["Post-save next-step prompt", "Mark next class or jump to results entry"],
       ["Admin attendance overview", "By class and date with present/late/absent counts and source"]],
      flow=[
          "Teacher opens Attendance → selects class and date (defaults to today).",
          "Taps Present/Late/Absent per student, or uses All Present / All Absent.",
          "Taps Save & Notify → records persist; absence and late alerts go to parents (WhatsApp + app).",
          "If offline, entries are cached with an offline badge and auto-sync when connectivity returns.",
          "Admin reviews Academic Hub → Attendance by class/date for the school-wide picture.",
      ])

story("US-1.7", "Discipline & Behaviour", "Teacher / Admin",
      "record commendations and misconduct with points",
      "behaviour is tracked and parents stay informed.",
      [["New record with type and points", "Commendation (+) or misconduct (−) with integer points"],
       ["Note required", "Context captured against the student"],
       ["History per student", "Discipline log shows type, points, note, date"],
       ["Parent alerts", "Parent notified of behaviour events"]],
      flow=[
          "Open Academic Hub → Discipline → New Record.",
          "Choose student, type (commendation/misconduct), points, and a note → Save Record.",
          "Record appears in the discipline history and the parent is notified.",
      ])

story("US-1.8", "Inventory Management", "School Admin",
      "track books, stock, and assets with low-stock alerts",
      "inventory losses are minimised.",
      [["Add item with min stock level", "Category, unit cost, quantity, supplier captured"],
       ["Low-stock alert banner", "Items below minimum flagged with a reorder prompt"],
       ["Stock adjustment with reason", "Recount / purchase / damage / loss / adjustment + notes"],
       ["Adjustment history per item", "Old vs new quantity, reason, who and when"],
       ["Total value computed", "Quantity × unit cost summed across inventory"]],
      flow=[
          "Open Inventory → review stats (items, total value, low-stock count).",
          "Add Inventory to register an item with its minimum stock level.",
          "Use Adjust Stock (with reason + notes) when quantities change; view History per item.",
          "Act on the amber low-stock banner to reorder before shortage.",
      ])

story("US-1.9", "Admissions Intake", "School Admin",
      "receive and process applications via a public link",
      "enrolment is streamlined from application to invoice.",
      [["Public admission link", "Shareable caspaa.com/apply/<schoolId>; no account needed to apply"],
       ["Application pipeline", "Pending → Reviewing → Accepted / Rejected with document checklist"],
       ["Document checklist verified", "Birth certificate, parent ID, immunisation, passport tracked X/4"],
       ["Accept & Enrol automation", "Creates/links parent, creates student (auto admission no.), generates invoice"],
       ["Welcome notification sent", "Parent receives enrolment confirmation + login credentials option"]],
      flow=[
          "Open People Hub → Admissions → copy the public admission link and share it.",
          "Applicants submit details + documents from the public form (no login).",
          "Filter the pipeline (Pending/Reviewing/Accepted/Rejected); open an application to verify documents.",
          "Click Accept & Enrol → parent + student records and an invoice are auto-created and a welcome message is sent.",
      ])

story("US-1.10", "Operations Registers (Sick Bay · Visitors · Library)", "School Admin",
      "run day-to-day operational registers",
      "health, security, and library activity are properly logged.",
      [["Sick bay visit logged", "Complaint, temperature, treatment, referral, parent notification"],
       ["Visitor gate log", "Check-in with purpose/vehicle and check-out timestamp; on-premises count"],
       ["Library catalogue & issue", "Books with copies/location; issue to student with due date"],
       ["Active loans & overdue tracking", "Return updates copies; overdue days computed"],
       ["Parent notified where relevant", "Sick bay can WhatsApp the parent on record"]],
      flow=[
          "Sick Bay → New Record: capture complaint, temperature, treatment, referral and (optionally) notify the parent.",
          "Visitors → Check-in a guest (purpose, vehicle); Check out when they leave — on-premises count updates.",
          "Library → add books to the catalogue, Issue to a student with a due date, and process Returns; overdue tab flags late loans.",
      ])

story("US-1.11", "School Settings & Configuration", "School Admin",
      "configure branding, academic calendar, roles, and integrations",
      "the platform reflects our school's identity and rules.",
      [["Branding", "School name, motto, primary colour, logo — applied to receipts, report cards, emails"],
       ["Academic structure", "Sessions, terms (current highlighted), and class arms"],
       ["Calendar events", "Exams, breaks, meetings with audience targeting"],
       ["Roles & permissions", "System and custom roles across 25+ granular permissions"],
       ["AI assistant toggles", "Report comments, fee reminders, attendance alerts, performance insights"],
       ["Payment gateway & data backup", "Paystack config; nightly multi-region backups with restore request"]],
      flow=[
          "Open Settings and pick a tab: Branding, Academic, Calendar, Roles & Permissions, AI Assistant, Payment Gateway, or Data Backup.",
          "Edit the relevant configuration (e.g., add a term, create a custom role, toggle an AI feature) and Save.",
          "Branding and academic settings immediately propagate to receipts, report cards, and class creation.",
      ])

story("US-1.12", "End-of-Term Close-Out Wizard", "School Admin",
      "close the term through a guided five-step flow",
      "term-end operations are complete and consistent.",
      [["Step 1 — approve pending results", "All unapproved results approved in one action"],
       ["Step 2 — generate broadsheets", "Per-class broadsheets merged to PDF"],
       ["Step 3 — publish report cards", "Report cards sent to parents via WhatsApp + email"],
       ["Step 4 — notify fee debtors", "Reminders sent to families with outstanding balances"],
       ["Step 5 — promote & set next-term fees", "Bulk promotion + fee structure duplication"],
       ["Progress is resumable", "Wizard remembers completed steps across sessions"]],
      flow=[
          "From the dashboard, click Close Term Wizard.",
          "Work through the five steps in any order across several days — approve results, generate broadsheets, publish report cards, remind debtors, promote students.",
          "The wizard marks each step done and remembers progress between sessions.",
      ])

# ---------------- MODULE 2 — PARENT APP ----------------
mod_group("Module 2 — Parent App")

story("US-2.1", "Onboarding & Child Dashboard", "Parent",
      "set up my account and monitor each child at a glance",
      "I can track attendance, performance, fees, and announcements.",
      [["First-login welcome wizard", "Set password → confirm contact/opt-ins → setup complete"],
       ["Per-child cards", "Attendance %, average score %, fees status with progress bar"],
       ["Multi-child overview", "Each child summarised; switch between children"],
       ["Payment summary widget", "Total billed, paid, outstanding with Pay Now"],
       ["Child detail tabs", "Overview, Results, Attendance, Homework"],
       ["Report card & transcript download", "Printable PDF from the Results tab"],
       ["Latest announcements shown", "Most recent school notices on the dashboard"]],
      flow=[
          "Parent signs in for the first time → welcome wizard: set a new password, confirm phone/email and WhatsApp/email opt-ins, review children → finish.",
          "Dashboard shows a payment summary and a card per child (attendance, average, fees status).",
          "Tap a child → tabs for Overview, Results, Attendance, Homework.",
          "From Results, download the Report Card or full Transcript as PDF.",
      ])

story("US-2.2", "Fees, Invoices & Payment", "Parent",
      "view and pay school fees in the app",
      "payment reflects immediately with a downloadable receipt.",
      [["Invoice breakdown per child", "Line items (incl. discounts as negative lines) and balance"],
       ["Paystack-style payment modal", "Amount with ⅓ / ½ / Full chips; method = Card / Transfer / USSD"],
       ["USSD path shows dial code", "*737*50*<amount># with confirm step"],
       ["Failure path is handled", "~10% simulated decline shows error code + Try again"],
       ["Success issues reference + receipt", "CSP-<ref> shown; printable PDF receipt"],
       ["Pay All multi-child queue", "Sequential payment per child with progress and per-child method"],
       ["Instalment plan", "Split balance across 2–6 dated payments with reminders"],
       ["Balance reconciles instantly", "Invoice paid/partial/outstanding status updates"]],
      flow=[
          "Parent opens Fees → reviews each child's invoice and balance.",
          "Taps Pay (or Pay All Fees) → payment modal: choose amount (⅓/½/Full) and method (Card/Transfer/USSD).",
          "Card/Transfer process with a brief animation; USSD shows a dial code to confirm. On decline, an error code + Try again appears.",
          "On success, a CSP reference and printable receipt are issued and the balance updates.",
          "For multiple children, the Pay All queue walks through each invoice, allowing a different method each time.",
      ])

story("US-2.3", "Communication", "Parent",
      "message teachers and receive announcements",
      "I can monitor my child's progress and stay informed.",
      [["Direct chat with teachers", "Real-time thread with delivery"],
       ["Attachments supported", "Images and documents up to 1MB"],
       ["WhatsApp launch button", "Opens the conversation on WhatsApp"],
       ["Announcements & PTA updates", "School broadcasts delivered to the parent"],
       ["Notifications", "Bell alerts for messages, fees, and behaviour events"]],
      flow=[
          "Parent opens Messages → selects or starts a chat with a teacher.",
          "Sends text or an attachment (≤1MB); may tap the WhatsApp button to continue there.",
          "Reads school broadcasts under Announcements; the bell surfaces new alerts.",
      ])

story("US-2.4", "Lending / Credit Facility", "Parent",
      "apply for a school-fee loan and manage repayment",
      "I can spread fees over time with an instant decision.",
      [["Credit score card", "Colour-coded score with eligible limit"],
       ["Application modal", "Select children, amount slider, term buttons, live monthly calc, consent"],
       ["Visible risk assessment", "Stepwise checks (identity, history, score, relationship, decision)"],
       ["Instant decision", "Approved when score ≥ threshold; otherwise declined with guidance"],
       ["Repayment schedule generated", "Monthly breakdown at 5% interest"],
       ["Auto-debit toggle", "Mandate to charge on due date"],
       ["Repayment progress tracked", "Paid/total instalments and next payment shown"]],
      flow=[
          "Parent opens Loans → reviews credit score and eligible limit → Apply for Loan.",
          "Selects which children the loan covers, sets amount (slider) and term, confirms the live monthly figure, and consents to auto-debit → Submit.",
          "A visible risk assessment runs and returns an instant decision.",
          "If approved, funds cover the fees and a repayment schedule appears; the parent can pay instalments or enable auto-debit.",
      ])

# ---------------- MODULE 3 — TEACHER APP ----------------
mod_group("Module 3 — Teacher App")

story("US-3.1", "Teacher Dashboard & Workforce", "Teacher",
      "see my day and manage clock-in, leave, and cover",
      "I can plan my day and handle HR tasks in one place.",
      [["Assigned classes, subjects, today's schedule", "Period-ordered timetable for the day"],
       ["Clock-in / clock-out card", "Self clock-in with automatic late detection; duration on clock-out"],
       ["Leave request widget", "Submit leave; pending/upcoming/total counts; admin notified"],
       ["Substitute coverage requests", "Accept/Decline cover; accepted classes added to schedule"],
       ["Stat cards", "Classes, total students, today marked, pending approvals"]],
      flow=[
          "Teacher signs in → dashboard shows today's schedule and a clock-in card.",
          "Taps Clock In (late auto-flagged) and Clock Out at end of day.",
          "Submits leave via Request Leave (admin is notified) and responds to any substitute-cover requests.",
      ])

story("US-3.2", "Assignments & Homework", "Teacher",
      "create, distribute, and grade assignments",
      "students and parents can track work and results.",
      [["Create assignment", "Title, class, subject, instructions, due date, optional file"],
       ["Parents notified on publish", "Class parents alerted with title and due date"],
       ["Submission tracking", "Submissions count vs class size with progress bar; overdue badge"],
       ["Grade out of 100", "Inline grading; parent notified of the grade"],
       ["Edit / delete", "Assignment editable; delete removes submissions with confirmation"]],
      flow=[
          "Teacher opens Assignments → New Assignment → fills title, class, subject, instructions, due date (and optional file) → Post.",
          "Parents in the class are notified automatically.",
          "Teacher opens an assignment to see submissions and enters a grade /100 per student; the parent is notified.",
      ])

story("US-3.3", "Lesson Planning", "Teacher",
      "create lesson notes tied to the scheme of work",
      "teaching is structured and curriculum coverage updates automatically.",
      [["Create lesson plan", "Class, subject, week, topic, objectives, activities, resources"],
       ["Tie to scheme of work", "Selecting a scheme week auto-fills week, topic, objectives, resources"],
       ["Marking covers the scheme week", "Tied week flagged covered with who/when"],
       ["File attachment", "PDF/Word/image up to 2MB attached and downloadable"]],
      flow=[
          "Teacher opens Lesson Plans → New Lesson Plan → choose class and subject.",
          "Optionally tie to a scheme-of-work week (auto-fills week/topic/objectives) or enter manually.",
          "Add activities, resources, and an optional file → Save Plan; the linked scheme week is marked covered.",
      ])

# ---------------- MODULE 4 — FINANCIAL MANAGEMENT ----------------
mod_group("Module 4 — Financial Management")

story("US-4.1", "Fees & Billing Configuration", "Finance Officer",
      "configure fees by class with a live total",
      "billing is automated and transparent.",
      [["Fee structure per class & term", "Tuition, books, uniform, PTA with live total"],
       ["Due date set", "Defaults provided, editable"],
       ["Edit warning shown", "Existing invoices not changed retroactively"],
       ["Discounts / scholarships", "Negative line item reduces balance immediately"],
       ["Instalment plans", "Schedule of dated payments generated"],
       ["Export", "Fee structure to CSV and branded PDF"]],
      flow=[
          "Finance Officer opens Fees → New Structure → selects class and term.",
          "Enters tuition/books/uniform/PTA; the per-student total updates live → Create Structure.",
          "Discounts/scholarships are applied on an invoice as a negative line item; instalment plans generate a dated schedule.",
          "Export the fee structure to CSV or PDF as needed.",
      ])

story("US-4.2", "Revenue, Invoices & Reminders", "Finance Officer",
      "monitor revenue, invoices, and chase debtors",
      "I have real-time visibility of the school's cash position.",
      [["Revenue dashboard", "Collected vs outstanding doughnut, income-vs-expense chart, completion %"],
       ["Invoice register with filters", "All / Paid / Partial / Outstanding with search"],
       ["Top debtors widget", "Highest balances with one-tap Send Reminder"],
       ["Reminder notifies parent", "Balance reminder delivered and logged"]],
      flow=[
          "Open Finance → Dashboard for collection vs outstanding and recent payments.",
          "Open Invoices to filter by status/search; open Top Debtors to Send Reminder to a parent.",
      ])

story("US-4.3", "Payments, Cash & Reconciliation", "Finance Officer",
      "record manual payments and reconcile incoming transfers",
      "all payments post correctly with minimal manual effort.",
      [["Manual / cash payment recording", "Cash/cheque/transfer/POS with reference, capped at balance"],
       ["Payment posts & notifies", "Invoice updated, parent notified, marked reconciled"],
       ["Auto-reconciliation heuristic", "Student surname matched from transfer narration, flagged with an AI badge"],
       ["Unmatched flagged for manual match", "Items queued; manual match updates balance"],
       ["Ledger search & CSV export", "By student, reference, or method"]],
      flow=[
          "Record Cash / Manual → select student (balance shown), amount, method, reference → Record; the invoice updates and the parent is notified.",
          "Open Reconciliation → the system suggests a student per incoming transfer by matching the narration (AI badge).",
          "Click Match to post, or handle unmatched items manually; export the ledger to CSV.",
      ])

story("US-4.4", "Expenses & Financial Reporting", "Finance Officer",
      "record expenses and produce audit-ready statements",
      "spending is controlled and profitability is clear.",
      [["Expense recording by category", "Salaries, utilities, maintenance, supplies, etc."],
       ["Category summary", "Totals per category aggregated"],
       ["P&L statement", "Revenue, expenses, net profit/loss for the term"],
       ["Debtors list & cashbook", "Top debtors and last-30-days cashbook"],
       ["Export P&L to PDF", "Branded, audit-ready statement"]],
      flow=[
          "Open Expenses → Add Expense (date, category, amount, description).",
          "Open Reports → review the P&L, debtors list, and cashbook → Export P&L (PDF).",
      ])

story("US-4.5", "Payroll (4-Stage Approval)", "Finance Officer",
      "run staff payroll through a controlled approval pipeline",
      "salaries are paid accurately with tax and pension handled.",
      [["Stage 1 — HR compute (draft)", "Roster built from salaries with adjustments (bonus/overtime/deduction)"],
       ["Stage 2 — Accounting validate", "Fund-availability check before approval"],
       ["Stage 3 — disburse", "Net paid via NIBSS; payslips issued; ledger updated"],
       ["Stage 4 — post-payroll", "PAYE and pension remittance + compliance checklist"],
       ["Notifications & expense entry", "Staff, bursar, and proprietor notified; salaries posted as expense"],
       ["Run summary PDF", "Per-staff gross/PAYE/pension/net exportable"]],
      flow=[
          "Open Payroll → New Run (draft); add any adjustments → Submit to Accounting.",
          "Accounting validates fund availability → Approve for Payment.",
          "Disburse → net pay is sent via NIBSS, payslips issue, and an expense is posted.",
          "Complete post-payroll tasks: mark PAYE and pension remitted and tick the compliance checklist; export the run summary PDF.",
      ])

story("US-4.6", "Loan Decisioning (Finance)", "Finance Officer",
      "review and decide parent loan applications",
      "lending decisions are consistent and well-documented.",
      [["Pending applications queue", "Live credit score per applicant, colour-coded"],
       ["Application review modal", "Score band, income tier, tenure, children covered, stated reason"],
       ["Risk flags surfaced", "e.g., loan > 150% of monthly income; score < 600"],
       ["Counter-offer", "Adjust amount/term before approval"],
       ["Approve disburses + schedules", "5% interest, monthly schedule, parent notified"],
       ["Reject with reason", "Mapped reason + note sent to applicant; audit-logged"]],
      flow=[
          "Open Lending → review the pending queue (each shows a live credit score).",
          "Open Review Application → inspect the score band, risk flags, and children covered.",
          "Optionally enter a counter-offer (amount/term), then Approve (disburses + builds schedule) or Reject (with a reason that is sent to the parent).",
      ])

# ---------------- MODULE 5 — PAYMENTS ----------------
mod_group("Module 5 — Payments Integration (Paystack)")

story("US-5.1", "Online Payment Processing", "Parent",
      "pay school fees online through Paystack",
      "payments are fast, secure, and confirmed instantly.",
      [["Card, transfer, and USSD", "Method chosen in the payment modal"],
       ["USSD dial-code path", "Code shown with manual confirm"],
       ["Transaction verification", "Status confirmed (webhook-style) before success"],
       ["Reference + receipt issued", "CSP reference and printable PDF receipt"],
       ["Failure handled gracefully", "Error code shown with Try again; no balance change"]],
      flow=[
          "From an invoice, the parent chooses an amount and a method (Card / Transfer / USSD).",
          "The charge is processed (or a USSD code is shown to confirm) and verified.",
          "On success a reference + receipt are issued; on failure an error code and retry are shown.",
      ])

# ---------------- MODULE 6 — LENDING ENGINE ----------------
mod_group("Module 6 — Lending Engine")

story("US-6.1", "Loan Application & Servicing", "Parent",
      "apply for a school-fee loan and track repayment",
      "I can access financing and manage it transparently.",
      [["Application captured", "Amount, term, children covered, consent"],
       ["Eligibility scored", "Decision from the risk engine"],
       ["Schedule with interest", "Monthly breakdown at 5% generated on approval"],
       ["Auto-debit mandate", "Optional charge on due date"],
       ["Monitoring & default tracking", "Status, repayment progress, and overdue flags"]],
      flow=[
          "Parent applies from the Loans view; the risk engine scores eligibility and returns a decision.",
          "On approval, a monthly repayment schedule is created and the parent can enable auto-debit.",
          "The parent (and platform) track repayment progress and any overdue instalments.",
      ])

# ---------------- MODULE 7 — RISK ENGINE ----------------
mod_group("Module 7 — Risk Engine")

story("US-7.1", "Credit Scoring & Risk Assessment", "Finance Officer / Platform",
      "automatically assess applicant risk before approving a loan",
      "lending decisions are consistent and defaults are reduced.",
      [["Credit score computed", "From payment history, income tier, and school tenure"],
       ["Stepwise assessment visible", "Identity, history, score, relationship, decision steps"],
       ["Risk flags generated", "Loan-to-income and low-score warnings; strong-history auto-approve hint"],
       ["Portfolio-at-risk monitoring", "PAR and risk buckets (A/B/C/D) at platform level"],
       ["Decision recorded", "Outcome stored with rationale for audit"]],
      flow=[
          "On application, the engine computes a credit score from payment history, income, and tenure.",
          "A stepwise assessment runs and surfaces risk flags to the decision-maker.",
          "At platform level, PAR and risk buckets monitor the whole loan book; every decision is recorded.",
      ])

# ---------------- MODULE 8 — AI ASSISTANT ----------------
mod_group("Module 8 — AI Assistant")

story("US-8.1", "AI Report-Card Comments", "Teacher",
      "generate report comments with AI",
      "I can reduce repetitive manual writing.",
      [["Comment generated from score band", "Excellent / good / average / poor templates personalised with name & subject"],
       ["Typed-out generation", "Comment types in and is fully editable"],
       ["Single or bulk", "Suggest one comment or all comments for a class"],
       ["Fast response", "Generated within seconds"]],
      flow=[
          "In Results entry, the teacher clicks the AI button on a row (or 'Suggest all comments').",
          "A score-appropriate comment is generated and typed into the field.",
          "The teacher edits if needed and submits with the results.",
      ])

story("US-8.2", "AI Insights & Automation", "Proprietor / Admin",
      "see AI risk insights and enable smart automation",
      "I can intervene early and reduce manual follow-up.",
      [["At-risk student insights", "Flags students with attendance < 75% or average < 50%"],
       ["AI reconciliation matching", "Suggests student matches from transfer narrations"],
       ["Configurable AI features", "Report comments, fee reminders, attendance alerts, performance insights toggles"],
       ["Insight drill-through", "Each flagged student links to their profile"]],
      flow=[
          "Open Settings → AI Assistant to toggle AI features on/off.",
          "Review AI Performance Insights for flagged at-risk students and drill into each profile.",
          "AI assists reconciliation by suggesting student matches for incoming transfers.",
      ])

# ---------------- MODULE 9 — COMMUNICATION SYSTEM ----------------
mod_group("Module 9 — Communication System")

story("US-9.1", "Messaging & Broadcast", "School Admin",
      "chat one-to-one and broadcast announcements across channels",
      "the whole school community stays informed.",
      [["One-to-one chat with attachments", "Images/documents up to 1MB; WhatsApp launch button"],
       ["Announcement composer", "Title, message, audience (everyone/parents/teachers)"],
       ["Multi-channel send", "In-app plus optional WhatsApp and Email toggles"],
       ["Delivery report", "Per-channel delivered/failed and success % (in-app 100%, WhatsApp ~96%, email ~99%)"]],
      flow=[
          "Open Messages for one-to-one chat (attachments + WhatsApp launch).",
          "Open Announcements → New Announcement → write the notice, pick the audience, toggle WhatsApp/Email → Send Now.",
          "Review the delivery report showing per-channel delivered/failed counts and success rates.",
      ])

# ---------------- MODULE 10 — ADMIN PORTAL ----------------
mod_group("Module 10 — Admin Portal / Dashboard (Super Admin)")

story("US-10.1", "School Onboarding & Subscription Management", "Super Admin",
      "onboard, configure, and manage schools and their subscriptions",
      "the platform is controlled centrally and securely.",
      [["Onboard school with KYC", "Profile + CAC/NIN/accreditation; starts a 14-day trial"],
       ["Plan management", "Essential / Professional / Enterprise; change, renew, auto-renew"],
       ["Suspend / reactivate", "Access disabled while data is retained"],
       ["Per-school feature flags", "WhatsApp, lending, AI, offline, transport, payroll toggles"],
       ["Subscription invoices", "Generate monthly, send reminders, mark paid, export"],
       ["Composite health & onboarding %", "Colour-coded school cards with drill-in"]],
      flow=[
          "Sign in as Super Admin (with 2FA OTP) → Schools → Onboard School (profile + KYC) → 14-day trial begins.",
          "Open a school → Subscription tab to change plan, renew, or toggle auto-renew; Features tab to flip per-school capabilities.",
          "Use Revenue → School Invoices to generate monthly invoices, send reminders, and mark paid.",
          "Suspend or reactivate a school from its detail modal as needed.",
      ])

story("US-10.2", "Lending Oversight & Disbursement", "Super Admin",
      "monitor the loan book and disburse approved loans",
      "platform lending is controlled and auditable.",
      [["Loan book overview", "Active loans, disbursed, interest income, repaid"],
       ["Disbursement console", "Verify recipient/account, confirm NIBSS transfer, ledger entry"],
       ["Loan analytics", "Repayment rate, PAR, risk buckets, delinquency list"],
       ["Delinquency reminders", "Bulk reminders to overdue borrowers"]],
      flow=[
          "Open Lending → Disbursement Console → review approved loans awaiting funds.",
          "Initiate Disbursement → verify the recipient account → confirm; a NIBSS transfer is simulated and the ledger updates.",
          "Use Loan Analytics to monitor repayment rate, PAR, and delinquency, and send reminders to overdue borrowers.",
      ])

story("US-10.3", "Platform Operations (Support · Team · Audit · Analytics)", "Super Admin",
      "run support, manage the CASPAA team, and review platform analytics",
      "operations are accountable and data-driven.",
      [["Support desk with SLA", "Tickets by status/priority/channel; SLA-at-risk surfaced; resolution notes"],
       ["Team RBAC", "Add members by role with granular permission toggles"],
       ["Audit log", "Every privileged action recorded with actor, action, target"],
       ["Analytics suite", "Business, platform usage (DAU/feature adoption), and system performance"],
       ["Platform settings", "Global feature flags and security posture (AES-256, MFA, backups)"]],
      flow=[
          "Open Support Desk to triage tickets, change status/assignment, add notes, and resolve within SLA.",
          "Open CASPAA Team to add members and fine-tune granular permissions; review the Audit Log for privileged actions.",
          "Open Analytics for business, usage, and system-performance views; manage global flags under Platform Settings.",
      ])

# ---------------- OFFLINE-FIRST ----------------
mod_group("Offline-First Capability")

story("US-O.1", "Offline Operations & Sync", "Teacher",
      "keep working when the internet is unavailable",
      "school operations continue in low-connectivity environments.",
      [["Offline mode indicated", "Offline badge shown in the top bar / on the attendance screen"],
       ["Attendance marked offline", "Entries buffered locally without internet"],
       ["Queued actions on reconnect", "'Syncing N items…' indicator processes the queue"],
       ["Retry queue handles failures", "Failed syncs retried without data loss"],
       ["Conflicts resolved gracefully", "Duplicate sync handled without overwrite"]],
      flow=[
          "Teacher toggles offline mode (or loses connectivity) → an offline badge appears.",
          "Attendance and other actions are captured and buffered locally.",
          "On reconnect, a 'Syncing N items…' indicator flushes the queue; failures retry without data loss.",
      ])

# ============================ TECHNICAL REQUIREMENTS ============================
doc.add_page_break()
doc.add_heading("7. Technical Requirements", level=1)

para("Note: the validated prototype is built as a zero-dependency browser app (HTML, Tailwind CSS, vanilla "
     "JavaScript, Chart.js) with a LocalStorage data layer, intentionally mirroring the modular architecture "
     "below so the production build can be a 1:1 port. The target production stack is:", italic=True, color=GREY)

doc.add_heading("Frontend", level=2)
add_table(["Platform", "Technology"], [["Web", "React.js / Next.js"], ["Mobile", "Flutter"]],
          widths=[2.4, 4.0])

doc.add_heading("Backend", level=2)
add_table(["Component", "Technology"],
          [["API", "Node.js (NestJS)"], ["Database", "PostgreSQL"], ["Realtime", "Socket.io"],
           ["Queue", "Redis"], ["AI Services", "OpenAI APIs"], ["Search", "Elasticsearch"]],
          widths=[2.4, 4.0])

doc.add_heading("Cloud Infrastructure", level=2)
add_table(["Service", "Provider"],
          [["Hosting", "AWS"], ["Storage", "S3"], ["CDN", "CloudFront"],
           ["Monitoring", "CloudWatch"], ["CI/CD", "GitHub Actions"]],
          widths=[2.4, 4.0])

doc.add_heading("Database Modules — Core Tables", level=2)
para("School, Students, Parents, Teachers, Classes, Attendance, Results, Payments, Loans, Wallets, "
     "Transactions, Messages, Assignments, Inventory, Audit Logs.")

doc.add_heading("API Structure — REST APIs", level=2)
add_table(["Module", "Endpoint Example"],
          [["Auth", "/api/auth"], ["Students", "/api/students"], ["Attendance", "/api/attendance"],
           ["Payments", "/api/payments"], ["Loans", "/api/loans"], ["Messaging", "/api/messages"]],
          widths=[2.4, 4.0])

doc.add_heading("Authentication", level=2)
bullets(["JWT", "OAuth", "Role-based access control"])

# ============================ NFRs ============================
doc.add_heading("8. Non-Functional Requirements (NFRs)", level=1)

doc.add_heading("Performance", level=2)
add_table(["Requirement", "Target"],
          [["API response time", "< 300 ms"], ["Concurrent users", "50,000+"],
           ["Real-time messaging latency", "< 2 sec"], ["Result generation", "< 10 sec"]],
          widths=[3.0, 3.4])

doc.add_heading("Offline & Progressive Web App", level=2)
add_table(["Requirement", "Description"],
          [["Offline attendance", "Teachers can mark attendance without internet"],
           ["Local caching", "Device stores temporary data"],
           ["Background sync", "Auto-sync when internet returns"],
           ["Progressive Web App", "Installable, lightweight app"],
           ["Conflict resolution", "Duplicate sync handled gracefully"]],
          widths=[2.4, 4.0])

doc.add_heading("Security", level=2)
para("Educational and financial data requires enterprise-grade protection.")
add_table(["Requirement", "Description"],
          [["Encryption at rest", "AES-256"], ["Encryption in transit", "TLS 1.3"],
           ["Role-based access", "Granular permissions"], ["Audit logging", "Every action tracked"],
           ["Session timeout", "Auto logout"], ["MFA (admin roles)", "Required"],
           ["Device verification", "Suspicious-login detection"]],
          widths=[2.4, 4.0])

doc.add_heading("Scalability", level=2)
add_table(["Requirement", "Description"],
          [["Multi-tenant architecture", "Thousands of schools"],
           ["Horizontal scaling", "Auto-scale backend"],
           ["Queue-based processing", "Async jobs"],
           ["CDN optimisation", "Faster media delivery"]],
          widths=[2.4, 4.0])

doc.add_heading("Reliability", level=2)
add_table(["Requirement", "Description"],
          [["Uptime target", "99.9%"], ["Automated backups", "Daily snapshots"],
           ["Disaster recovery", "Multi-region backups"], ["Failover support", "Redundant services"]],
          widths=[2.4, 4.0])

doc.add_heading("Usability", level=2)
add_table(["Requirement", "Description"],
          [["Mobile-first UX", "Android-optimised"], ["Low-bandwidth optimisation", "Compressed payloads"],
           ["Simple navigation", "Friendly to low-tech users"], ["Accessibility", "Large fonts, simple flows"]],
          widths=[2.4, 4.0])

doc.add_heading("Compliance", level=2)
add_table(["Requirement", "Description"],
          [["Data privacy principles", "FERPA-like compliance"],
           ["Payment compliance", "PCI DSS considerations"],
           ["NDPR compliance", "Nigerian Data Protection Regulation"]],
          widths=[2.4, 4.0])

# ============================ ROLES ============================
doc.add_heading("9. System Roles & Permissions", level=1)
add_table(["Role", "Key Access"],
          [["Super Admin", "Platform-wide control"],
           ["Proprietor / Admin", "School management"],
           ["Finance Officer", "Payments and reconciliation"],
           ["Teacher", "Academics and attendance"],
           ["Parent", "Child monitoring and payments"],
           ["Student", "Assignments and results"]],
          widths=[2.4, 4.0])

# ============================ SUCCESS ============================
doc.add_heading("10. MVP Success Definition", level=1)
para("The MVP will be considered successful if:")
bullets(["Schools actively use the ERP daily",
         "Parents regularly pay fees digitally",
         "Loan repayments remain healthy",
         "Payment reconciliation reduces manual work",
         "Engagement levels validate product-market fit",
         "The platform proves scalable in Lagos schools",
         "Onboarding is fast and easy"])

doc.add_heading("11. Success Metrics & KPIs", level=1)

doc.add_heading("Business KPIs", level=2)
add_table(["KPI", "Year 1 Target"],
          [["Schools onboarded", "60"], ["Paid schools", "50"], ["Student records digitised", "6,000"],
           ["Payment volume", "₦1B+"], ["Loan disbursement", "₦300M"], ["Revenue", "₦138M+"]],
          widths=[3.2, 3.2])

doc.add_heading("Product KPIs", level=2)
add_table(["KPI", "Target"],
          [["Monthly Active Users", "> 85%"], ["Parent app engagement", "> 70% weekly"],
           ["Attendance submission rate", "> 85%"], ["Assignment completion rate", "> 85%"],
           ["Payment success rate", "> 95%"]],
          widths=[3.2, 3.2])

doc.add_heading("FinTech KPIs", level=2)
add_table(["KPI", "Target"],
          [["Loan repayment rate", "> 92%"], ["Payment reconciliation accuracy", "> 98%"],
           ["Failed payment rate", "< 3%"], ["Loan approval turnaround", "< 24 hrs"]],
          widths=[3.2, 3.2])

doc.add_heading("Technical KPIs", level=2)
add_table(["KPI", "Target"],
          [["Uptime", "99.9%"], ["API response time", "< 300 ms"],
           ["Crash-free sessions", "> 99%"], ["Offline sync success rate", "> 95%"]],
          widths=[3.2, 3.2])

doc.add_heading("Customer Success KPIs", level=2)
add_table(["KPI", "Target"],
          [["School retention", "> 90%"], ["Parent satisfaction score", "> 4 / 5"],
           ["NPS score", "> 50"], ["Support resolution SLA", "< 12 hrs"]],
          widths=[3.2, 3.2])

# ============================ OUT OF SCOPE ============================
doc.add_page_break()
doc.add_heading("12. Out of Scope (Post-MVP Roadmap)", level=1)
para("The following features are intentionally excluded from the MVP to reduce complexity and accelerate launch.")

doc.add_heading("FinTech Features", level=2)
add_table(["Feature", "Reason"],
          [["Teacher salary advance", "Phase 2 lending expansion"],
           ["Wallet system", "Requires additional licensing"],
           ["Savings accounts", "Banking-layer expansion"],
           ["Investment lock features", "Future wealth products"],
           ["Bulk disbursement", "Requires treasury infrastructure"],
           ["School working-capital loans", "Post-risk-engine maturity"],
           ["Asset-acquisition financing (BNPL)", "Later lending vertical"],
           ["Teacher BNPL", "Future payroll integration"],
           ["Micro / business loans", "Future SME lending"]],
          widths=[3.0, 3.4])

doc.add_heading("Payments & Infrastructure", level=2)
add_table(["Feature", "Reason"],
          [["NIBSS integration", "Requires advanced settlement infrastructure"],
           ["Offline payment authorisation", "Complex reconciliation risk"],
           ["Cross-border payments", "Requires FX compliance"],
           ["Multi-currency support", "Not an initial focus"]],
          widths=[3.0, 3.4])

doc.add_heading("Advanced AI & Risk", level=2)
add_table(["Feature", "Reason"],
          [["Behavioural credit-scoring engine", "Requires historical datasets"],
           ["AI loan-underwriting automation", "Future iteration"]],
          widths=[3.0, 3.4])

doc.add_heading("Advanced School Operations", level=2)
add_table(["Feature", "Reason"],
          [["Transport management", "Operationally heavy"],
           ["Clock-in device integration", "Hardware dependency"],
           ["Internal control & audit workflows", "Enterprise phase"],
           ["e-Learning / online classes", "Separate product vertical"],
           ["Story books / digital library", "Content licensing required"]],
          widths=[3.0, 3.4])

doc.add_heading("Language Support", level=2)
add_table(["Feature", "Reason"],
          [["French language support", "West Africa expansion phase"],
           ["Portuguese support", "Lusophone Africa expansion"]],
          widths=[3.0, 3.4])

# ============================ ADDENDUM (v1.2) ============================
doc.add_page_break()
doc.add_heading("13. Addendum — New & Enhanced Modules (Pilot v1.2)", level=1)
para("The following modules were added to the CASPAA MVP scope following stakeholder engagements with school "
     "administrators and proprietors during pilot validation. They are intended to improve product-market fit, "
     "increase adoption, enhance operational efficiency, expand student engagement, and deepen the EdTech × FinTech "
     "infrastructure. All of these modules have now been built end-to-end in the working prototype, including a "
     "dedicated Student Portal with login, gamification (stars, points, badges), CBT exams with auto-grading, an "
     "LMS integrated with assignments and a full submit-grade loop, digital consent, HR appraisal, accounting "
     "statements, and in-app support with live chat. Implementation priority is retained in Section 16; the pilot "
     "UX refinements applied to the prototype are listed in Section 17.", italic=True, color=GREY)

add_table(["Module", "Status"],
          [["Student Portal", "NEW"], ["CBT Learning", "NEW"], ["LMS (Learning Management System)", "ENHANCED"],
           ["HR & Staff Appraisal", "ENHANCED"], ["Digital Consent", "NEW"],
           ["Accounting & Financial Reporting", "ENHANCED"], ["Advanced Support & Live Chat", "ENHANCED"]],
          widths=[4.0, 2.4])

NEW = "NEW — Implemented (prototype, Addendum v1.2)"
ENH = "ENHANCED — Implemented (prototype, Addendum v1.2)"

module(11, "Student Portal", "Provide a dedicated student-facing portal and mobile experience.", [
    ("Access", ["Secure student login", "Personalised dashboard", "Assignments / homework access",
        "CBT access (digital assessments)", "LMS access (learning materials)", "Timetable access", "Notifications"]),
    ("Academics & Conduct", ["View report cards / academic results", "Behavioural dashboard (conduct tracking)",
        "Commendations & rewards", "Stars & points", "Trophies & badges (gamification)"]),
], status=NEW)

story("US-11.1", "Student Login & Learning", "Student",
      "securely log in to CASPAA",
      "I can access assignments, CBT exams, academic records, and learning materials independently.",
      [["Student can log in securely", "Role-scoped session created"],
       ["Student sees only assigned content", "Access limited to own class/subjects"],
       ["Student can access assignments and results", "Homework and report cards visible"],
       ["Student can participate in CBT exams", "Active assessments launchable"],
       ["Dashboard loads quickly", "Under 3-second load"],
       ["Parents can monitor activity", "Parent view reflects student engagement"]],
      flow=[
          "Student signs in with their credentials (access level scales by section: nursery limited, primary guided, secondary independent).",
          "Dashboard shows current assignments, upcoming CBT exams, attendance, performance, stars/points, and commendations.",
          "Student opens assignments or learning materials, takes CBT assessments, and views results — all scoped to their own record.",
      ])

module(12, "CBT Learning", "Conduct digital tests, quizzes, practice assessments, and examinations.", [
    ("Authoring", ["CBT exam/test creation", "Reusable question bank", "Objective (MCQ) questions",
        "Theory (essay) questions", "Randomised questions", "Practice tests"]),
    ("Delivery & Grading", ["Timer management (timed assessments)", "Auto-grading of objective questions",
        "Result synchronisation with report cards", "Basic anti-cheating / integrity controls", "Mobile-device support"]),
], status=NEW)

story("US-12.1", "CBT Authoring", "Teacher",
      "create and publish computer-based tests",
      "assessments are digital, paper-free, and faster to grade.",
      [["Exam created for a subject & class", "Test saved against the class"],
       ["Questions added from bank or new", "Objective and theory items supported"],
       ["Duration and rules set", "Timer and integrity controls configured"],
       ["Exam published to students", "Visible to the assigned class"]],
      flow=[
          "Teacher creates a CBT exam/test → selects subject & class.",
          "Adds questions (objective and/or theory), optionally pulling from the question bank, and enables randomisation.",
          "Sets duration and rules → publishes the CBT to the class.",
      ])

story("US-12.2", "CBT Participation", "Student",
      "take tests online",
      "I can complete assessments digitally and receive instant feedback.",
      [["Student accesses CBT securely", "Only assigned tests are launchable"],
       ["Timer functions correctly", "Countdown enforced; auto-submit at expiry"],
       ["Objective questions auto-grade instantly", "Score computed on submission"],
       ["Results sync to report card", "Scores flow into Result Management"],
       ["Teacher can review submissions", "Theory answers available for marking"],
       ["Works on mobile devices", "Responsive on phones/tablets"]],
      flow=[
          "Student opens an active CBT and starts the timed test.",
          "On submit (or timer expiry) the system auto-grades objective questions instantly.",
          "Scores synchronise to the Result module; the teacher reviews theory submissions.",
      ])

module(13, "Learning Management System (LMS)", "Expand CASPAA into a hybrid learning platform.", [
    ("Content & Resources", ["Upload notes", "Upload videos", "Assignment upload", "Resource library (central repository)"]),
    ("Tracking", ["Grade assignments", "Topic tracking (curriculum progress)", "Student progress monitoring (learning analytics)"]),
    ("Synchronisation (key requirement)", ["Syncs with Result Management", "Syncs with Student Dashboard",
        "Syncs with Teacher Analytics", "Syncs with Parent Academic View"]),
], status=ENH)

story("US-13.1", "Digital Learning Materials", "Teacher",
      "upload notes, videos, and assignments",
      "students can access learning materials digitally and progress is trackable.",
      [["Teachers can upload resources", "Notes, videos, and files stored in the library"],
       ["Students can download materials", "Resources accessible from the student portal"],
       ["Assignments sync to grading system", "Submissions feed the grading workflow"],
       ["Learning progress is trackable", "Topic coverage and analytics updated"]],
      flow=[
          "Teacher uploads notes/videos and posts assignments to the resource library.",
          "Students download materials and submit work from the student portal.",
          "Submissions sync to grading; progress and topic coverage update across Result Management, dashboards, and the parent academic view.",
      ])

module(14, "HR & Staff Appraisal", "Digitise school HR operations and staff performance management.", [
    ("Appraisal & KPIs", ["Staff appraisal (performance reviews)", "KPI tracking (productivity)",
        "Staff performance reports / analytics"]),
    ("Workflow", ["Salary advance requests (HR financing workflow)", "Approval hierarchy (multi-level approvals)"]),
    ("Appraisal metrics", ["Attendance / punctuality", "Result submission (academic performance)",
        "Parent feedback (engagement)", "Classroom performance (quality metrics)"]),
], status=ENH)

story("US-14.1", "Digital Staff Appraisal", "School Admin",
      "appraise teachers digitally",
      "I can monitor staff performance efficiently.",
      [["Appraisal forms configurable", "Metrics and weightings adjustable"],
       ["Workflow approvals functional", "Multi-level approval hierarchy enforced"],
       ["Reports downloadable", "Performance reports exportable"],
       ["Role-based access enforced", "Only authorised roles can appraise"]],
      flow=[
          "Admin opens HR & Staff Appraisal → selects a staff member and an appraisal form.",
          "Scores the metrics (attendance, result submission, parent feedback, classroom performance); KPIs compute.",
          "The appraisal routes through the approval hierarchy; the final report is downloadable.",
      ])

module(15, "Digital Consent", "Digitise parental authorisation and approvals.", [
    ("Use cases", ["Excursion consent (field trips)", "PTA consent (parent approvals)",
        "Media consent (photo permissions)", "Policy acceptance (terms agreement)"]),
    ("Features", ["Digital consent forms", "E-signature support", "Timestamp logging (audit trail)",
        "Consent repository (secure storage)", "Parent notifications (approval requests)", "Optional PDF upload"]),
], status=NEW)

story("US-15.1", "Digital Parental Consent", "Parent",
      "approve school activities digitally",
      "I do not need physical paperwork.",
      [["Parent can approve digitally", "E-signature / Agree captured"],
       ["Timestamp recorded", "Approval time logged for audit"],
       ["Consent retrievable later", "Stored securely in the repository"],
       ["Notifications sent successfully", "Parent notified of each request"]],
      flow=[
          "School creates a consent form → parent receives a notification.",
          "Parent reviews the terms and clicks Agree (e-signature).",
          "The timestamp is recorded and the consent is stored securely and retrievable later.",
      ])

module(16, "Accounting & Financial Reporting", "Expand financial management into a mini-school accounting system.", [
    ("Statements", ["Trial balance", "Cash flow statement", "Income statement", "Profit & loss", "Balance sheet"]),
    ("Operations", ["Expense management", "Budget management", "Bank reconciliation (payment matching)"]),
], status=ENH)

story("US-16.1", "Automated Accounting Reports", "Bursar",
      "generate accounting reports automatically",
      "financial reporting becomes easier and more accurate.",
      [["Reports generated automatically", "Statements computed from captured transactions"],
       ["Data export supported", "Reports exportable (PDF/Excel)"],
       ["Reconciliation accurate", "Bank matching balances correctly"],
       ["Financial dashboard updates in real time", "Admin dashboard reflects latest figures"]],
      flow=[
          "Transactions are captured → the ledger updates → accounting entries are generated.",
          "Financial reports (trial balance, cash flow, income statement, P&L, balance sheet) are computed.",
          "The admin dashboard updates in real time; the bursar exports audit-ready statements.",
      ])

module(17, "Advanced Support & Live Chat", "Provide enterprise-grade customer-support infrastructure.", [
    ("Features", ["Live chat support (real-time assistance)", "Ticketing system (issue management)",
        "In-app help centre (knowledge base)", "SLA tracking (support monitoring)",
        "Multi-agent support (team collaboration)"]),
], status=ENH)

story("US-17.1", "In-App Support", "School Administrator",
      "get immediate support",
      "operational issues are resolved quickly.",
      [["Live chat accessible in-app", "Chat reachable from any screen"],
       ["Tickets trackable", "Status visible end-to-end"],
       ["Response time under SLA", "SLA timers enforced and surfaced"],
       ["Notifications functional", "Updates delivered to the requester"]],
      flow=[
          "Admin opens in-app live chat or the help centre and raises an issue.",
          "A ticket is created and tracked against its SLA, handled by one or more support agents.",
          "The admin receives notifications until the ticket is resolved.",
      ])

# ---------------- 14. PLATFORM CONFIGURABILITY ----------------
doc.add_heading("14. Platform Configurability", level=1)
para("Strategic requirement: schools must be able to configure major platform settings independently.")
add_table(["Setting", "Configurable"],
          [["Grading system", "Yes"], ["Timetable structure", "Yes"], ["Fee structure", "Yes"],
           ["Behavioural points", "Yes"], ["Academic sessions", "Yes"], ["Extracurricular", "Yes"],
           ["Notifications", "Yes"]],
          widths=[3.6, 2.8])

# ---------------- 15. PERFORMANCE & UX (ADDENDUM) ----------------
doc.add_heading("15. Performance & UX Requirements (Addendum)", level=1)
doc.add_heading("Performance", level=2)
add_table(["Requirement", "Target"],
          [["Dashboard load time", "< 3 seconds"], ["Payment processing", "< 10 seconds"],
           ["CBT submission", "< 2 seconds"], ["Live chat response", "Real-time"]],
          widths=[3.4, 3.0])
doc.add_heading("Usability", level=2)
add_table(["Requirement", "Description"],
          [["Mobile-first UX", "Android priority"], ["Simple navigation", "Low-tech users"],
           ["Low-bandwidth optimisation", "Emerging markets"], ["Offline capability", "Attendance & assignments"]],
          widths=[2.6, 3.8])

# ---------------- 16. IMPLEMENTATION PRIORITY ----------------
doc.add_heading("16. Implementation Priority", level=1)
add_table(["Module", "Priority"],
          [["Student Module", "HIGH"], ["CBT Module", "HIGH"], ["LMS Enhancements", "HIGH"],
           ["Digital Consent", "MEDIUM"], ["HR Appraisal", "MEDIUM"], ["Accounting Reports", "HIGH"],
           ["Live Chat Support", "HIGH"]],
          widths=[4.0, 2.4])

# ---------------- 17. PILOT REVIEW ADJUSTMENTS ----------------
doc.add_heading("17. Pilot Review Adjustments (Applied to Prototype)", level=1)
para("UX/UI refinements gathered from the pilot demo-script review. All of the following have been applied to the "
     "working prototype in v1.2.", italic=True, color=GREY)
add_table(["Area", "Adjustment", "Status"],
          [["Login hero", "Reworded to 'From School Operations, Payment, Financing, Attendance, Learning, CBT and Engagement Infrastructure — a unified solution replacing several tools'", "Applied"],
           ["Login feature tiles", "Renamed 'Paystack Built-in' → 'Payment'; added 'CBT Learnings' and 'Digital Consent' tiles", "Applied"],
           ["Navigation", "Renamed 'People' to 'Pupils'", "Applied"],
           ["Dashboard date", "Added a date picker to review a previous day's analytics", "Applied"],
           ["Terminology", "Outstanding-fees count now reads 'students owing' (was 'families')", "Applied"],
           ["Average-score metric", "Relabelled 'Academic Avg Score' to clarify it is academic, not payment %", "Applied"],
           ["School notifications", "Each notification is now clickable and deep-links to the relevant page", "Applied"],
           ["Gender insight", "Added a student gender pie chart (boys/girls) for supplies planning", "Applied"],
           ["Revenue chart", "Changed from 6-month view to revenue per term (Nigerian termly model)", "Applied"],
           ["Student profile", "'AVG Score' relabelled 'Academic Avg Score'; row action icons given clear tooltips & hover", "Applied"],
           ["Results / broadsheet", "Per-student 'Generate' result action that publishes and shares with the parent", "Applied"],
           ["Invoices (Finance)", "Generate & send invoice to parent via WhatsApp + email", "Applied"],
           ["Receipts (Finance)", "Generate & send receipt to parent (WhatsApp + email) on part/full payment", "Applied"],
           ["Payroll", "Renamed 'Export Roster' → 'Export Payroll Schedule'", "Applied"],
           ["Fees", "Added 'Extracurricular' fee line (swimming, ballet, music, etc.) across structures & invoices", "Applied"],
           ["Staff attendance", "Admins can clock a staff member out from the gate-book records", "Applied"],
           ["Messaging", "Empty-state CTA now reads 'Start a chat with a parent' (role-aware)", "Applied"]],
          widths=[1.7, 3.9, 0.8])

doc.save("PRD _EDTECH_CASPAA _AFRISPRINGS.docx")
print("Saved PRD _EDTECH_CASPAA _AFRISPRINGS.docx")
