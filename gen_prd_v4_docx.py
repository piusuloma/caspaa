# -*- coding: utf-8 -*-
"""Render CASPAA_PRD_v4.md -> CASPAA_PRD_v4.docx, styled on the Academy of
Product Management PRD template (brand headings, shaded table headers, code
blocks, scope callout). Lightweight Markdown subset converter."""
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r"C:\Users\USER\Desktop\CASPAA\CASPAA_PRD_v4.md"
OUT = r"C:\Users\USER\Desktop\CASPAA\CASPAA_PRD_v4.docx"

BRAND   = RGBColor(0x1F, 0x4E, 0x79)
ACCENT  = RGBColor(0x2E, 0x75, 0xB6)
GREY    = RGBColor(0x59, 0x59, 0x59)
HDRFILL = "1F4E79"
CALLOUT = "FFF4CE"   # soft amber for the scope callout / blockquotes
CODEFILL = "F4F5F7"

doc = Document()

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.13

for lvl, sz, col, bold in [("Heading 1", 17, BRAND, True), ("Heading 2", 13.5, BRAND, True),
                           ("Heading 3", 11.5, ACCENT, True), ("Heading 4", 10.5, ACCENT, True)]:
    st = doc.styles[lvl]
    st.font.name = "Calibri"; st.font.size = Pt(sz); st.font.color.rgb = col; st.font.bold = bold
    st.paragraph_format.space_before = Pt(10); st.paragraph_format.space_after = Pt(4)


def shade(el, fill):
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), fill)
    el.append(sh)


def set_cell_border(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "4"); e.set(qn("w:color"), "BFBFBF")
        borders.append(e)
    tcPr.append(borders)


INLINE = re.compile(r"(\*\*.+?\*\*|`.+?`|\*.+?\*)")

def add_runs(p, text):
    text = text.replace("&amp;", "&")
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2]); r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(9.5)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            r = p.add_run(part[1:-1]); r.italic = True
        else:
            p.add_run(part)


def add_table(rows):
    header, body = rows[0], rows[1:]
    t = doc.add_table(rows=1, cols=len(header))
    t.autofit = True
    hcells = t.rows[0].cells
    for i, htext in enumerate(header):
        set_cell_border(hcells[i])
        shade(hcells[i]._tc.get_or_add_tcPr(), HDRFILL)
        para = hcells[i].paragraphs[0]; para.paragraph_format.space_after = Pt(2)
        r = para.add_run(htext.replace("**", "")); r.bold = True; r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); r.font.size = Pt(9.5)
    for row in body:
        cells = t.add_row().cells
        for i, ctext in enumerate(row if len(row) == len(header) else (row + [""] * (len(header) - len(row)))[:len(header)]):
            set_cell_border(cells[i])
            para = cells[i].paragraphs[0]; para.paragraph_format.space_after = Pt(2)
            for r in para.runs:
                r.font.size = Pt(9.5)
            add_runs(para, ctext)
            for r in para.runs:
                r.font.size = Pt(9.5)
    doc.add_paragraph()


def add_code(lines):
    for ln in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0); p.paragraph_format.left_indent = Pt(6)
        shade(p._p.get_or_add_pPr(), CODEFILL)
        r = p.add_run(ln if ln else " "); r.font.name = "Consolas"; r.font.size = Pt(9)
    doc.add_paragraph()


def add_callout(text):
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]; set_cell_border(cell)
    shade(cell._tc.get_or_add_tcPr(), CALLOUT)
    p = cell.paragraphs[0]
    add_runs(p, text)
    doc.add_paragraph()


def parse_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


with open(SRC, encoding="utf-8") as f:
    lines = f.read().split("\n")

i = 0
first_h1 = True
while i < len(lines):
    line = lines[i]

    # code fence
    if line.strip().startswith("```"):
        block = []
        i += 1
        while i < len(lines) and not lines[i].strip().startswith("```"):
            block.append(lines[i]); i += 1
        add_code(block); i += 1; continue

    # table
    if line.strip().startswith("|") and i + 1 < len(lines) and re.match(r"^\s*\|[\s:\-|]+\|\s*$", lines[i + 1]):
        rows = [parse_row(line)]
        i += 2
        while i < len(lines) and lines[i].strip().startswith("|"):
            rows.append(parse_row(lines[i])); i += 1
        add_table(rows); continue

    stripped = line.strip()

    if stripped.startswith("#### "):
        doc.add_paragraph(stripped[5:], style="Heading 4")
    elif stripped.startswith("### "):
        doc.add_paragraph(stripped[4:], style="Heading 3")
    elif stripped.startswith("## "):
        doc.add_paragraph(stripped[3:], style="Heading 2")
    elif stripped.startswith("# "):
        title = stripped[2:]
        if first_h1:
            p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
            r = p.add_run(title); r.bold = True; r.font.size = Pt(24); r.font.color.rgb = BRAND
            first_h1 = False
        else:
            doc.add_paragraph(title, style="Heading 1")
    elif stripped.startswith("> "):
        add_callout(stripped[2:])
    elif stripped in ("---", "***", "___"):
        pass  # section rule; headings already provide separation
    elif stripped.startswith(("- ", "* ")):
        p = doc.add_paragraph(style="List Bullet"); add_runs(p, stripped[2:])
    elif re.match(r"^\d+\.\s", stripped):
        p = doc.add_paragraph(style="List Number"); add_runs(p, re.sub(r"^\d+\.\s", "", stripped))
    elif stripped == "":
        pass
    else:
        p = doc.add_paragraph(); add_runs(p, stripped)
    i += 1

doc.save(OUT)
print("Wrote", OUT)
