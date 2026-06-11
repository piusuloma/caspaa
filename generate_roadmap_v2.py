from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colours ──────────────────────────────────────────────────────────────────
DARK_NAVY   = RGBColor(0x0a, 0x25, 0x40)
EMERALD     = RGBColor(0x0d, 0x6e, 0x3f)
LIGHT_MINT  = RGBColor(0xec, 0xfb, 0xf4)
RED         = RGBColor(0xdc, 0x26, 0x26)
LIGHT_RED   = RGBColor(0xfe, 0xe2, 0xe2)
AMBER       = RGBColor(0xd9, 0x77, 0x06)
LIGHT_AMBER = RGBColor(0xff, 0xf8, 0xe8)
BLUE        = RGBColor(0x1d, 0x4e, 0xd8)
LIGHT_BLUE  = RGBColor(0xeb, 0xf4, 0xff)
PURPLE      = RGBColor(0x6d, 0x28, 0xd9)
LIGHT_PURPLE= RGBColor(0xf5, 0xf3, 0xff)
GRAY_BG     = RGBColor(0xf8, 0xf9, 0xfa)
GRAY_LINE   = RGBColor(0xe9, 0xec, 0xef)
GRAY_TEXT   = RGBColor(0x6b, 0x72, 0x80)
MID_GRAY    = RGBColor(0x37, 0x41, 0x51)
WHITE       = RGBColor(0xff, 0xff, 0xff)

ACTOR_COLORS = {
    'Admin':    (DARK_NAVY,  RGBColor(0xe8, 0xf0, 0xfe)),
    'Teacher':  (EMERALD,    LIGHT_MINT),
    'Parent':   (AMBER,      LIGHT_AMBER),
    'Student':  (PURPLE,     LIGHT_PURPLE),
    'Finance':  (RED,        LIGHT_RED),
    'System':   (RGBColor(0x4b,0x55,0x63), GRAY_BG),
    'Principal':(BLUE,       LIGHT_BLUE),
}


# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  str(rgb))
    for ex in tcPr.findall(qn('w:shd')):
        tcPr.remove(ex)
    tcPr.append(shd)

def set_col_widths(table, widths):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths):
                cell.width = widths[i]

def set_row_height(row, twips):
    trPr = row._tr.get_or_add_trPr()
    h = OxmlElement('w:trHeight')
    h.set(qn('w:val'), str(twips))
    trPr.append(h)

def set_doc_margins(doc, top=1.0, bottom=1.0, left=1.1, right=1.1):
    for s in doc.sections:
        s.top_margin    = Inches(top)
        s.bottom_margin = Inches(bottom)
        s.left_margin   = Inches(left)
        s.right_margin  = Inches(right)

def add_run(para, text, bold=False, italic=False, size=10.5, color=MID_GRAY):
    r = para.add_run(text)
    r.bold   = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return r

def para(doc, text='', bold=False, size=10.5, color=MID_GRAY,
         align=WD_ALIGN_PARAGRAPH.LEFT, sb=0, sa=4, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if text:
        add_run(p, text, bold=bold, size=size, color=color, italic=italic)
    return p

def page_break(doc):
    doc.add_page_break()


# ── Cover ─────────────────────────────────────────────────────────────────────

def build_cover(doc):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.rows[0].cells[0]
    set_cell_bg(c, DARK_NAVY)
    set_row_height(t.rows[0], 9200)

    p = c.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(70)
    add_run(p, 'CASPAA', bold=True, size=40, color=WHITE)

    p2 = c.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p2, 'SCHOOL OPERATING SYSTEM', size=9,
            color=RGBColor(0x34, 0xd3, 0x99))

    p3 = c.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(30)
    p3.paragraph_format.space_after  = Pt(8)
    add_run(p3, 'New Modules Implementation Roadmap', bold=True, size=26, color=WHITE)

    p4 = c.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p4, 'Real-Life Step-by-Step User Flows', italic=True, size=14,
            color=RGBColor(0x93, 0xc5, 0xfd))

    p5 = c.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p5.paragraph_format.space_before = Pt(18)
    add_run(p5,
        'Every flow below describes exactly what a real user clicks, sees, types,\n'
        'and receives — and what the system does behind the scenes at each moment.',
        size=11, color=RGBColor(0xb0, 0xc4, 0xde))

    p6 = c.add_paragraph()
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p6.paragraph_format.space_before = Pt(48)
    add_run(p6,
        'Phase 1: Weeks 1–4   |   Phase 2: Weeks 5–10   |   Phase 3: Weeks 11–16+   |   14 Modules',
        size=9, color=RGBColor(0x64, 0x80, 0x9a))

    page_break(doc)


# ── Section header ────────────────────────────────────────────────────────────

def section_header(doc, label, title, intro=''):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    add_run(p, label.upper(), bold=True, size=8.5, color=EMERALD)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(4)
    add_run(p2, title, bold=True, size=20, color=DARK_NAVY)

    if intro:
        p3 = doc.add_paragraph()
        p3.paragraph_format.space_after = Pt(16)
        add_run(p3, intro, size=10.5, color=GRAY_TEXT)


# ── Module banner ─────────────────────────────────────────────────────────────

def module_banner(doc, num, title, subtitle, phase_label, phase_color, effort, effort_color):
    t = doc.add_table(rows=1, cols=1)
    c = t.rows[0].cells[0]
    set_cell_bg(c, DARK_NAVY)

    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(8)
    add_run(p, f'MODULE {num}  ', bold=True, size=9, color=RGBColor(0x34,0xd3,0x99))
    add_run(p, f'{title}', bold=True, size=14, color=WHITE)

    p2 = c.add_paragraph()
    p2.paragraph_format.space_after = Pt(4)
    add_run(p2, subtitle, size=10, color=RGBColor(0xa0,0xb8,0xd0))

    p3 = c.add_paragraph()
    p3.paragraph_format.space_after = Pt(8)
    add_run(p3, f'Phase: ', bold=True, size=9, color=RGBColor(0x80,0xa0,0xc0))
    add_run(p3, phase_label, bold=True, size=9, color=phase_color)
    add_run(p3, '    Effort: ', bold=True, size=9, color=RGBColor(0x80,0xa0,0xc0))
    add_run(p3, effort, bold=True, size=9, color=effort_color)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ── Flow header (e.g. "Flow 1 of 2: Admin creates an event") ─────────────────

def flow_header(doc, flow_num, total, title):
    t = doc.add_table(rows=1, cols=1)
    c = t.rows[0].cells[0]
    set_cell_bg(c, GRAY_BG)

    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(5)
    add_run(p, f'FLOW {flow_num} of {total}  ', bold=True, size=8.5, color=GRAY_TEXT)
    add_run(p, title, bold=True, size=11.5, color=DARK_NAVY)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ── Pre-conditions box ────────────────────────────────────────────────────────

def preconditions(doc, items):
    t = doc.add_table(rows=1, cols=1)
    c = t.rows[0].cells[0]
    set_cell_bg(c, LIGHT_BLUE)

    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(5)
    add_run(p, 'PRE-CONDITIONS: ', bold=True, size=8.5, color=BLUE)
    add_run(p, '  '.join(f'[{i}]' for i in items), size=9, color=BLUE)

    for item in items:
        pi = c.add_paragraph()
        pi.paragraph_format.space_after = Pt(2)
        pi.paragraph_format.left_indent = Inches(0.2)
        add_run(pi, f'• {item}', size=9, color=MID_GRAY)

    c.paragraphs[-1].paragraph_format.space_after = Pt(6)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ── Step row ─────────────────────────────────────────────────────────────────
# actor: 'Admin' | 'Teacher' | 'Parent' | 'Student' | 'Finance' | 'System' | 'Principal'

def step(doc, num, actor, action, detail='', is_system=False):
    txt_color, bg_color = ACTOR_COLORS.get(actor, (DARK_NAVY, GRAY_BG))

    t = doc.add_table(rows=1, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Col 0: step number
    c0 = t.rows[0].cells[0]
    set_cell_bg(c0, bg_color)
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_before = Pt(5)
    p0.paragraph_format.space_after  = Pt(5)
    add_run(p0, str(num), bold=True, size=11, color=txt_color)

    # Col 1: actor badge
    c1 = t.rows[0].cells[1]
    set_cell_bg(c1, bg_color)
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(5)
    p1.paragraph_format.space_after  = Pt(5)
    add_run(p1, actor.upper(), bold=True, size=7.5, color=txt_color)

    # Col 2: action + detail
    c2 = t.rows[0].cells[2]
    set_cell_bg(c2, WHITE)
    p2 = c2.paragraphs[0]
    p2.paragraph_format.space_before = Pt(5)
    p2.paragraph_format.space_after  = Pt(2)
    add_run(p2, action, bold=True, size=10.5, color=DARK_NAVY)

    if detail:
        pd = c2.add_paragraph()
        pd.paragraph_format.space_after  = Pt(5)
        pd.paragraph_format.left_indent  = Inches(0.05)
        add_run(pd, detail, size=9.5, color=GRAY_TEXT, italic=is_system)

    set_col_widths(t, [Inches(0.38), Inches(0.75), Inches(5.37)])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


# ── Outcome box ───────────────────────────────────────────────────────────────

def outcome(doc, items):
    t = doc.add_table(rows=1, cols=1)
    c = t.rows[0].cells[0]
    set_cell_bg(c, LIGHT_MINT)

    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    add_run(p, 'OUTCOME', bold=True, size=8.5, color=EMERALD)

    for item in items:
        pi = c.add_paragraph()
        pi.paragraph_format.left_indent = Inches(0.15)
        pi.paragraph_format.space_after = Pt(2)
        add_run(pi, f'+ {item}', size=9.5, color=MID_GRAY)

    c.paragraphs[-1].paragraph_format.space_after = Pt(8)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ── Edge cases box ────────────────────────────────────────────────────────────

def edge_cases(doc, items):
    t = doc.add_table(rows=1, cols=1)
    c = t.rows[0].cells[0]
    set_cell_bg(c, LIGHT_RED)

    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(5)
    add_run(p, 'EDGE CASES & ERROR HANDLING', bold=True, size=8.5, color=RED)

    for item in items:
        pi = c.add_paragraph()
        pi.paragraph_format.left_indent = Inches(0.15)
        pi.paragraph_format.space_after = Pt(2)
        add_run(pi, f'! {item}', size=9.5, color=RED)

    c.paragraphs[-1].paragraph_format.space_after = Pt(8)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)


# ════════════════════════════════════════════════════════════════════════════
# BUILD THE DOCUMENT
# ════════════════════════════════════════════════════════════════════════════

def build():
    doc = Document()
    set_doc_margins(doc)
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10.5)

    build_cover(doc)

    # ── HOW TO READ THIS DOCUMENT ─────────────────────────────────────────────
    section_header(doc, 'How to read this document', 'Reading Guide')

    guide = doc.add_table(rows=1, cols=1)
    gc = guide.rows[0].cells[0]
    set_cell_bg(gc, GRAY_BG)
    gp = gc.paragraphs[0]
    gp.paragraph_format.space_before = Pt(10)
    add_run(gp, 'Each module is broken into numbered Flows. ', bold=True, size=10, color=DARK_NAVY)
    add_run(gp, 'A Flow is a complete real-life scenario from the first click to the final result.', size=10, color=GRAY_TEXT)

    actors = [('Admin', DARK_NAVY, RGBColor(0xe8,0xf0,0xfe)),
              ('Teacher', EMERALD, LIGHT_MINT),
              ('Parent', AMBER, LIGHT_AMBER),
              ('Student', PURPLE, LIGHT_PURPLE),
              ('Finance', RED, LIGHT_RED),
              ('System', RGBColor(0x4b,0x55,0x63), GRAY_BG),
              ('Principal', BLUE, LIGHT_BLUE)]

    gp2 = gc.add_paragraph()
    gp2.paragraph_format.space_before = Pt(6)
    add_run(gp2, 'Actor colour codes:  ', bold=True, size=9.5, color=DARK_NAVY)
    for name, tc, _ in actors:
        add_run(gp2, f'[{name}]  ', bold=True, size=9, color=tc)

    gp3 = gc.add_paragraph()
    gp3.paragraph_format.space_after = Pt(10)
    add_run(gp3,
        'Steps labelled [System] happen automatically — no user action required. '
        'Steps labelled with a role show what that person actually does in the app.',
        size=9.5, color=GRAY_TEXT)

    doc.add_paragraph()
    page_break(doc)


    # ════════════════════════════════════════════════════════════════════════
    # PHASE 1
    # ════════════════════════════════════════════════════════════════════════

    section_header(doc, 'Phase 1 — Weeks 1–4',
        'Core Content & Communication Layer',
        'Six modules that unlock daily student, parent, and teacher usage. '
        'These close the biggest visible gap between CASPAA and EDVES.')

    # ─────────────────────────────────────────────────────────────────────────
    # MODULE 1: EVENT CALENDAR
    # ─────────────────────────────────────────────────────────────────────────
    module_banner(doc, 1, 'Event Calendar',
        'School-wide events visible to every role — Admin creates, everyone sees.',
        'Phase 1 — Critical', RGBColor(0xdc,0x26,0x26), 'Medium', AMBER)

    flow_header(doc, 1, 2, 'Admin creates a new school event')
    preconditions(doc, [
        'Admin is logged in',
        'At least one class exists in the school',
    ])
    step(doc, 1, 'Admin', 'Opens Communications hub',
         'Admin clicks "Communications" in the left sidebar. Hub shows three tabs: Messages | Announcements | Calendar.')
    step(doc, 2, 'Admin', 'Clicks the "Calendar" tab',
         'A monthly grid renders. Today\'s date is highlighted. Any existing events appear as colour-coded dots on their dates. '
         'If no events exist yet, an empty-state message reads: "No events this month — add your first one."')
    step(doc, 3, 'Admin', 'Clicks "+ Add Event" button (top right)',
         'A modal opens with the following fields:\n'
         '  • Event Title (required)\n'
         '  • Start Date  (required, date picker)\n'
         '  • End Date    (optional — for multi-day events)\n'
         '  • Event Type  (dropdown): Holiday / Exam / Sports Day / PTA Meeting / School Trip / Other\n'
         '  • Audience    (dropdown): Everyone / Students Only / Parents Only / Teachers Only\n'
         '  • Description (textarea, optional)')
    step(doc, 4, 'Admin', 'Fills in the form',
         'Example: Title = "Mid-Term Break", Start = 16 Nov 2026, End = 20 Nov 2026, '
         'Type = Holiday, Audience = Everyone, Description = "School resumes Monday 23rd November."')
    step(doc, 5, 'Admin', 'Clicks "Save Event"',
         'Form validates: Title and Start Date must not be empty. If valid, the modal closes.')
    step(doc, 6, 'System', 'Saves event to the Events table',
         'Record: { id, schoolId, title, startDate, endDate, type, audience, description, createdBy, createdAt }', is_system=True)
    step(doc, 7, 'System', 'Pushes event to all relevant dashboards',
         '"Upcoming Events" widget on Admin, Teacher, Parent, and Student dashboards now shows '
         '"Mid-Term Break — 16 Nov" within the next page load. Audience filter is applied: '
         'if Audience = "Parents Only", only parent dashboards show it.', is_system=True)
    step(doc, 8, 'Admin', 'Sees the event on the calendar grid',
         'A red dot (Holiday colour) appears on 16–20 Nov. Hovering or tapping the dot shows a tooltip: '
         '"Mid-Term Break — 16–20 Nov 2026 — Holiday".')
    outcome(doc, [
        'Event is saved and visible on the school calendar for all targeted roles.',
        '"Upcoming Events" widget on all relevant dashboards shows the new event.',
        'Admin can click any event to Edit or Delete it.',
    ])
    edge_cases(doc, [
        'Title left blank → "Event title is required" validation error shown inline. Form does not submit.',
        'End Date set before Start Date → "End date cannot be before start date" error.',
        'Admin tries to create an event on a past date → allowed (for recording purposes) but a warning shows: "This date is in the past."',
    ])

    flow_header(doc, 2, 2, 'Teacher / Parent views the calendar')
    step(doc, 1, 'Teacher', 'Logs in and sees dashboard',
         '"Upcoming Events" widget shows the next 3 events. Example: "Mid-Term Break — 16 Nov  |  PTA Meeting — 28 Oct  |  Sports Day — 5 Dec".')
    step(doc, 2, 'Teacher', 'Clicks "View Full Calendar" link in the widget',
         'Navigates to the teacher calendar view. Month grid shows all events colour-coded. Teachers see events where Audience = "Everyone" or "Teachers Only".')
    step(doc, 3, 'Teacher', 'Clicks an event on the grid',
         'A small detail card pops up showing: Title, Date range, Type badge, Description. No edit controls (read-only for teachers).')
    step(doc, 4, 'Parent', 'Same experience as Teacher',
         'Parents see events where Audience = "Everyone" or "Parents Only". They access the calendar from the parent portal menu.')
    outcome(doc, [
        'All roles have a read-only calendar view filtered to their audience.',
        'No teacher or parent can create or edit events — Admin-only.',
    ])

    page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # MODULE 2: BULK SMS & EMAIL
    # ─────────────────────────────────────────────────────────────────────────
    module_banner(doc, 2, 'Bulk SMS & Email Notifications',
        'Admin sends mass communications; system auto-sends on fee payments, results, and events.',
        'Phase 1 — Critical', RGBColor(0xdc,0x26,0x26), 'Medium', AMBER)

    flow_header(doc, 1, 3, 'Admin sends a manual bulk SMS')
    preconditions(doc, [
        'SMS provider (Africa\'s Talking or Termii) API key configured in School Settings',
        'School has SMS units credited (e.g. 500 units)',
        'Parents/teachers have phone numbers on their profiles',
    ])
    step(doc, 1, 'Admin', 'Goes to Communications > SMS',
         'Sees the SMS dashboard. Top of page shows: "SMS Balance: 500 units remaining." '
         'Below is a log of past SMS sends (empty for first use).')
    step(doc, 2, 'Admin', 'Clicks "Compose SMS"',
         'A two-step compose form opens.')
    step(doc, 3, 'Admin', 'Step 1 — Selects Audience',
         'Dropdown options:\n'
         '  • All Parents (shows count: "187 parents")\n'
         '  • All Teachers (shows count: "24 teachers")\n'
         '  • Specific Class (sub-dropdown appears — select JSS 2A → "32 parents of JSS 2A students")\n'
         '  • Custom (paste phone numbers manually)\n\n'
         'Admin selects "All Parents". Recipient count updates to "187 recipients".')
    step(doc, 4, 'Admin', 'Step 2 — Types the message',
         'Textarea with live character counter below: "0 / 160 characters". At 160 chars = 1 SMS unit per recipient.\n\n'
         'Admin types: "Dear Parent, school resumes Monday 23rd November. Please ensure fees are paid before resumption. Thank you — Bright Lights Academy"\n\n'
         'Counter shows: "142 / 160 characters — 1 unit per recipient — Total: 187 units needed"\n\n'
         'Yellow warning banner: "Note: Numbers on Do-Not-Disturb (DND) consume extra units. Actual delivery may vary."')
    step(doc, 5, 'Admin', 'Clicks "Preview & Send"',
         'Confirmation modal:\n'
         '"You are about to send to 187 parents.\n'
         'Units required: 187\n'
         'Units remaining after send: 313\n'
         'Proceed?"')
    step(doc, 6, 'Admin', 'Clicks "Confirm Send"',
         'Button changes to "Sending..." with a spinner. Admin cannot close the modal.')
    step(doc, 7, 'System', 'Makes API call to SMS provider',
         'Sends the message to all 187 phone numbers. Receives delivery report from provider.', is_system=True)
    step(doc, 8, 'System', 'Updates SMS balance and logs the send',
         'Balance: 500 → 313 units. SMS Log entry created:\n'
         '{ audience: "All Parents", recipientCount: 187, unitsUsed: 187, message: "...", status: "Sent", sentAt: [timestamp], sentBy: [adminId] }', is_system=True)
    step(doc, 9, 'Admin', 'Sees success confirmation',
         'Modal closes. Toast: "SMS sent to 187 recipients." '
         'New row appears in the SMS log table: "All Parents — 187 sent — 187 units — Just now."')
    outcome(doc, [
        'Message delivered to all 187 parent phone numbers.',
        'SMS balance deducted from school account.',
        'Full audit log of send with timestamp and who sent it.',
    ])
    edge_cases(doc, [
        'Message exceeds 160 characters → counter turns red. Warning: "This will use 2 units per recipient (374 units total)." Admin can still send.',
        'SMS balance too low → "Insufficient units. You need 187 units but only have 50. Please top up." Send button is disabled.',
        'API call fails → toast error: "SMS send failed. Please try again or contact support." Balance is NOT deducted. Log entry marked "Failed".',
        'No phone numbers in audience → "No valid phone numbers found for selected audience." Send is blocked.',
    ])

    flow_header(doc, 2, 3, 'System auto-sends SMS when a payment is received')
    step(doc, 1, 'System', 'Detects successful fee payment',
         'Parent pays ₦85,000 online. Payment verified via Paystack callback.', is_system=True)
    step(doc, 2, 'System', 'Fires auto-SMS to parent',
         'Message: "Dear Mr. Tunde, payment of ₦85,000 for Tobi Okafor has been received. '
         'Receipt: RCP-2026-4821. Balance: ₦0. Thank you — Bright Lights Academy."\n\n'
         'SMS Log records this as Type = "Auto / Fee Payment".', is_system=True)
    step(doc, 3, 'System', 'Fires auto-SMS to Finance Officer',
         'Message: "New payment: ₦85,000 from Tunde Okafor for Tobi Okafor (JSS 2A). Ref: PAY-2026-0342."', is_system=True)
    outcome(doc, [
        'Parent receives instant confirmation without having to refresh the app.',
        'Finance Officer is alerted even if not logged into CASPAA at that moment.',
    ])

    flow_header(doc, 3, 3, 'Admin sends a bulk email announcement')
    preconditions(doc, [
        'Email provider (Resend/SendGrid) configured in School Settings',
        'Parents/teachers have email addresses on their profiles',
    ])
    step(doc, 1, 'Admin', 'Goes to Communications > Email',
         'Same layout as SMS. Shows email log and "Compose Email" button.')
    step(doc, 2, 'Admin', 'Composes email',
         'Fields: Audience (same options as SMS), Subject line, Body (rich text — bold, lists, links), Optional attachment (max 5MB).')
    step(doc, 3, 'Admin', 'Previews email',
         'A preview panel renders how the email will look in a recipient\'s inbox — school logo at top, message body, school contact footer.')
    step(doc, 4, 'Admin', 'Sends',
         'System sends via email provider. Each email is personalised: "Dear Mr. Tunde Okafor" (not "Dear Parent").')
    step(doc, 5, 'System', 'Logs delivery status',
         'Email log shows: Delivered / Bounced / Opened per recipient (if provider supports open tracking).', is_system=True)
    outcome(doc, [
        'All targeted recipients receive a personalised, professionally formatted email.',
        'Bounced emails are flagged — admin can correct the email address on the parent\'s profile.',
    ])
    edge_cases(doc, [
        'Parent has no email address → excluded from send. Admin sees "12 recipients skipped (no email on file)".',
        'Attachment exceeds 5MB → error shown before sending. Admin must reduce file size.',
    ])

    page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # MODULE 3: CLASS NOTES
    # ─────────────────────────────────────────────────────────────────────────
    module_banner(doc, 3, 'Class Notes',
        'Teachers post digital lesson notes. Students read them in the learning portal.',
        'Phase 1 — Critical', RGBColor(0xdc,0x26,0x26), 'Low', EMERALD)

    flow_header(doc, 1, 2, 'Teacher posts a class note')
    preconditions(doc, [
        'Teacher is assigned to at least one class and subject',
        'At least one student is enrolled in that class',
    ])
    step(doc, 1, 'Teacher', 'Goes to Lessons > Notes tab',
         'Notes tab shows a list of all previously posted notes. First use → empty state: "No notes posted yet. Add your first note."')
    step(doc, 2, 'Teacher', 'Clicks "+ Add Note"',
         'Form opens with fields:\n'
         '  • Class         (dropdown — shows only teacher\'s assigned classes)\n'
         '  • Subject       (auto-filters based on class selection)\n'
         '  • Week          (Week 1 through Week 13 dropdown)\n'
         '  • Title         (text input, required)\n'
         '  • Content       (large textarea — supports basic formatting)\n'
         '  • Attachment    (optional: PDF, Word, image — max 10MB)')
    step(doc, 3, 'Teacher', 'Fills in the note',
         'Class = JSS 2A, Subject = Mathematics, Week = 8, '
         'Title = "Introduction to Quadratic Equations", '
         'Content = "A quadratic equation is any equation of the form ax² + bx + c = 0 where a ≠ 0..." '
         '(teacher types full note content). Attaches a PDF worksheet.')
    step(doc, 4, 'Teacher', 'Clicks "Post Note"',
         'Validation: Title and Content must not be empty. Class and Subject must be selected.')
    step(doc, 5, 'System', 'Saves note and notifies students',
         'ClassNotes record saved. In-app notifications created for all 32 students in JSS 2A:\n'
         '"New Mathematics note from Mr. Ibrahim: Introduction to Quadratic Equations."', is_system=True)
    step(doc, 6, 'Teacher', 'Sees the note in the list',
         'Note card shows: Title, Subject, Week, Date posted, and a read counter: "0 / 32 students read".')
    outcome(doc, [
        'Note is available instantly in the student Learning portal.',
        'All 32 students in JSS 2A receive an in-app notification.',
        'Teacher can track exactly who has and hasn\'t read the note.',
    ])

    flow_header(doc, 2, 2, 'Student reads the class note')
    step(doc, 1, 'Student', 'Logs in — sees notification badge on Learning hub',
         'Bell icon or "Learning" menu item shows a badge: "1". Student clicks Learning.')
    step(doc, 2, 'Student', 'Goes to Class Notes tab',
         'A card appears: "Introduction to Quadratic Equations — Mathematics — Mr. Ibrahim — Week 8 — [NEW]"\n'
         'The [NEW] badge is shown because the student has not yet opened this note.')
    step(doc, 3, 'Student', 'Clicks the note card to open it',
         'Full note renders: teacher name, date, subject, week, full content text, and a "Download Attachment" button for the PDF.')
    step(doc, 4, 'System', 'Marks note as read for this student',
         'NoteViews record created: { noteId, studentId, viewedAt }. The [NEW] badge disappears.', is_system=True)
    step(doc, 5, 'Teacher', 'Later checks the read count',
         'Teacher\'s notes list now shows: "1 / 32 students read" for that note. '
         'Teacher can click to see which specific students have read it and which haven\'t.')
    outcome(doc, [
        'Student can access the full note and download any attached files.',
        'Teacher has full visibility into student engagement with posted notes.',
    ])
    edge_cases(doc, [
        'Student is not in the class the note was posted to → note does not appear for that student.',
        'Teacher tries to post a note to a class they are not assigned to → class does not appear in the dropdown.',
        'Attachment fails to upload (network error) → note is saved without attachment. Toast: "Note saved. Attachment failed — please try re-attaching."',
    ])

    page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # MODULE 4: VIDEO LESSONS
    # ─────────────────────────────────────────────────────────────────────────
    module_banner(doc, 4, 'Video Lessons',
        'Teachers link or upload video content per subject. Students watch from the learning portal.',
        'Phase 1 — Critical', RGBColor(0xdc,0x26,0x26), 'Low', EMERALD)

    flow_header(doc, 1, 2, 'Teacher adds a video lesson')
    step(doc, 1, 'Teacher', 'Goes to Lessons > Videos tab',
         'Empty state on first use: "No video lessons yet. Add your first video."')
    step(doc, 2, 'Teacher', 'Clicks "+ Add Video"',
         'Form opens:\n'
         '  • Class / Subject (same dropdowns as Class Notes)\n'
         '  • Week\n'
         '  • Title\n'
         '  • Description (optional)\n'
         '  • Video Source: [Paste YouTube/Vimeo link]  OR  [Upload video file]\n\n'
         'Teacher selects "Paste YouTube link" and enters: https://www.youtube.com/watch?v=abc123')
    step(doc, 3, 'System', 'Auto-detects the video and generates a preview',
         'The system extracts the YouTube video ID, generates:\n'
         '  • Thumbnail from YouTube\'s thumbnail API\n'
         '  • Embed URL: https://www.youtube.com/embed/abc123\n'
         'A preview thumbnail appears instantly below the link field.', is_system=True)
    step(doc, 4, 'Teacher', 'Sees the thumbnail preview and confirms the details',
         'Teacher sees the video thumbnail, confirms title and class, then clicks "Save Video".')
    step(doc, 5, 'System', 'Saves video record and notifies students',
         'Videos record saved. Notifications sent to students in that class: "New Mathematics video: Solving Quadratic Equations."', is_system=True)
    outcome(doc, [
        'Video appears in the student Learning portal under "Video Lessons" tab.',
        'Students receive an in-app notification.',
    ])

    flow_header(doc, 2, 2, 'Student watches a video lesson')
    step(doc, 1, 'Student', 'Opens Learning > Video Lessons tab',
         'A grid of video cards is shown. Each card displays: thumbnail, title, subject, teacher\'s name, and a "Watch" button.')
    step(doc, 2, 'Student', 'Clicks a video card or "Watch" button',
         'A full-screen modal opens. The video plays inline (YouTube embed / Vimeo embed). '
         'Student can play, pause, seek, and use subtitles if available from the source.')
    step(doc, 3, 'System', 'Records that the student watched this video',
         'VideoView record created: { videoId, studentId, watchedAt }.', is_system=True)
    step(doc, 4, 'Student', 'Closes the video',
         'Returns to the video grid. A "Watched" tick appears on the card they just viewed.')
    step(doc, 5, 'Student', 'Student dashboard shows "Continue Learning"',
         'On next login, the dashboard shows a "Recently Watched" or "Continue Learning" row with unwatched videos from their classes.')
    outcome(doc, [
        'Student can watch any lesson video without leaving CASPAA.',
        'Teacher and admin can see how many students have watched each video.',
    ])
    edge_cases(doc, [
        'Invalid YouTube URL pasted → "Unable to load video preview. Please check the URL." Save is blocked until a valid URL is entered.',
        'Uploaded video file exceeds 200MB → "File too large. Please upload a shorter clip or use a YouTube link instead."',
        'Student has no internet when clicking the video → video player shows standard browser error. A message below reads: "Having trouble loading? Check your connection."',
    ])

    page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # MODULE 5: PhET SIMULATIONS
    # ─────────────────────────────────────────────────────────────────────────
    module_banner(doc, 5, 'Maths & Science Simulations (PhET)',
        'Students launch free interactive science and maths simulations from within CASPAA.',
        'Phase 1 — Critical', RGBColor(0xdc,0x26,0x26), 'Low', EMERALD)

    flow_header(doc, 1, 2, 'Student explores simulations')
    step(doc, 1, 'Student', 'Opens Learning > Simulations tab',
         'A catalogue of simulation cards is displayed. Subject filter buttons at the top: All | Mathematics | Physics | Chemistry | Biology')
    step(doc, 2, 'Student', 'Filters by subject',
         'Student clicks "Physics". Cards filter to show only Physics simulations:\n'
         '"Wave on a String", "Projectile Motion", "Ohm\'s Law", "Circuit Construction Kit", etc.\n'
         'Each card shows: simulation name, a screenshot thumbnail, a short description, and a "Launch" button.')
    step(doc, 3, 'Student', 'Clicks "Launch" on "Projectile Motion"',
         'A full-screen modal opens. The PhET simulation loads in an iframe. '
         'The student can interact with the simulation fully — adjust angles, velocity, mass — as if they were on the PhET website.')
    step(doc, 4, 'Student', 'Uses the simulation',
         'Student fires the cannon at different angles, observes trajectories. '
         'All simulation controls work natively inside CASPAA — no redirect to an external site.')
    step(doc, 5, 'Student', 'Closes the simulation',
         'Clicks the X button on the modal. Returns to the Simulations catalogue.')
    outcome(doc, [
        'Student gets a full interactive lab experience without any login or external account.',
        'PhET simulations are free and require no API key or subscription.',
    ])

    flow_header(doc, 2, 2, 'Teacher assigns a simulation to a class')
    step(doc, 1, 'Teacher', 'Browses the Simulations catalogue',
         'Teacher goes to Lessons > Simulations. Same catalogue view as students, but with an "Assign to Class" button on each card.')
    step(doc, 2, 'Teacher', 'Clicks "Assign to Class" on "Ohm\'s Law"',
         'A small form appears: Select Class, Select Week. Teacher picks JSS 3A, Week 5.')
    step(doc, 3, 'Teacher', 'Clicks "Assign"',
         'Assignment saved.')
    step(doc, 4, 'System', 'Pushes simulation to student dashboards',
         'All students in JSS 3A see a widget on their dashboard: "Simulation this week: Ohm\'s Law — Physics — Mr. Ibrahim. Launch now."', is_system=True)
    outcome(doc, [
        'Students are guided to the right simulation each week by their teacher.',
        'Simulation appears prominently on the student dashboard — no hunting through the catalogue.',
    ])
    edge_cases(doc, [
        'Student\'s device does not support the simulation (very old browser) → PhET shows its own "browser not supported" message inside the iframe. A note below reads: "Try opening this on a different device."',
        'Teacher assigns same simulation twice to same class → system deduplicates, shows only once on student dashboards.',
    ])

    page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # MODULE 6: PAYMENT GATEWAY
    # ─────────────────────────────────────────────────────────────────────────
    module_banner(doc, 6, 'Online Payment Gateway',
        'Parents pay school fees directly from CASPAA using card or bank transfer via Paystack.',
        'Phase 1 — Critical', RGBColor(0xdc,0x26,0x26), 'High', RED)

    flow_header(doc, 1, 2, 'Parent pays outstanding fees online')
    preconditions(doc, [
        'Admin has configured Paystack API keys in School Settings',
        'Student has an outstanding fee invoice',
        'Parent is logged in and has a valid card or bank account',
    ])
    step(doc, 1, 'Parent', 'Opens Fees from the parent portal menu',
         'Sees a list of children\'s invoices. Example:\n'
         '"Tobi Okafor — JSS 2A — Term 1 2026/27\n'
         'Total: ₦185,000 | Paid: ₦100,000 | Outstanding: ₦85,000 | Status: Partial"\n\n'
         'A red "Pay Now" button appears next to the outstanding balance.')
    step(doc, 2, 'Parent', 'Clicks "Pay Now"',
         'A payment summary modal opens:\n\n'
         '"Payment Summary\n'
         'Student: Tobi Okafor\n'
         'Amount Due: ₦85,000\n'
         'Transaction Fee (1.5%): ₦1,275\n'
         'Total Charged to Card: ₦86,275\n\n'
         'Note: Transaction fees are charged by the payment provider and are non-refundable."\n\n'
         'Two buttons at the bottom: "Pay with Card" | "Cancel"')
    step(doc, 3, 'Parent', 'Clicks "Pay with Card"',
         'The Paystack inline popup appears on top of the CASPAA page. '
         'It is pre-filled with: Amount = ₦86,275, Email = parent\'s email from their profile.')
    step(doc, 4, 'Parent', 'Enters card details in the Paystack popup',
         'Card number: 5399 XXXX XXXX XXXX\n'
         'Expiry: 12/28\n'
         'CVV: 123\n\n'
         'Parent clicks "Pay ₦86,275".')
    step(doc, 5, 'System', 'Paystack processes the payment and sends OTP',
         'Paystack sends a one-time PIN to the parent\'s registered mobile number. '
         'The popup prompts: "Enter the OTP sent to 0803***4567."', is_system=True)
    step(doc, 6, 'Parent', 'Enters the OTP',
         'Types the 6-digit OTP received via SMS and clicks "Confirm".')
    step(doc, 7, 'System', 'Payment authorised — Paystack callback fires',
         'Paystack sends a success callback to CASPAA with the transaction reference.', is_system=True)
    step(doc, 8, 'System', 'CASPAA verifies the transaction with Paystack API',
         'CASPAA makes a server-side API call to Paystack to verify the reference. '
         'Only proceeds if Paystack confirms status = "success" and amount matches. '
         'This prevents manipulation of the callback.', is_system=True)
    step(doc, 9, 'System', 'Creates transaction record and updates invoice',
         'Transactions record: { method: "Card", amount: 85000, fee: 1275, invoiceId, reference: "PAY-REF-7821", status: "Completed", timestamp }\n'
         'Invoice updated: Outstanding = ₦0, Status = "Paid".', is_system=True)
    step(doc, 10, 'System', 'Sends confirmation SMS and email to parent',
         'SMS: "Payment confirmed! ₦85,000 received for Tobi Okafor. Ref: PAY-REF-7821. Balance: ₦0. — Bright Lights Academy"\n'
         'Email: Full receipt with school logo, line items, and receipt number.', is_system=True)
    step(doc, 11, 'System', 'Notifies Finance Officer',
         'In-app notification to Finance Officer: "Online payment of ₦85,000 received from Tunde Okafor for Tobi Okafor (JSS 2A). Ref: PAY-REF-7821."\n'
         'Finance dashboard: new transaction row appears with a "Card" badge. Marked as reconciled automatically.', is_system=True)
    step(doc, 12, 'Parent', 'Sees success screen',
         'Paystack popup closes. CASPAA shows a success card:\n'
         '"Payment Successful!\n'
         'Amount: ₦85,000\n'
         'Receipt: PAY-REF-7821\n'
         '[Download Receipt]  [Back to Fees]"\n\n'
         'The invoice card on the Fees page now shows Status = "Paid" in green.')
    outcome(doc, [
        'Parent\'s invoice is fully cleared and updated in real time.',
        'Finance Officer receives an alert and sees the payment without manual entry.',
        'Parent holds a downloadable receipt with full details.',
        'Transaction is auto-reconciled — no manual Finance Officer action needed for online payments.',
    ])
    edge_cases(doc, [
        'Card declined by bank → Paystack shows "Your card was declined. Try a different card or contact your bank." CASPAA invoice is not updated.',
        'OTP entered incorrectly 3 times → Paystack cancels the transaction. Parent must restart the payment flow.',
        'Parent closes the Paystack popup before completing → invoice remains unpaid. No charge is made. CASPAA shows: "Payment was not completed."',
        'Paystack callback received but server-side verification fails (amount mismatch) → transaction is NOT recorded. Logged as a failed verification attempt. Finance Officer is alerted.',
        'Parent tries to overpay (manually editing the amount) → Paystack amount is set server-side from the invoice. Cannot be altered client-side.',
    ])

    flow_header(doc, 2, 2, 'Finance Officer reviews online payments')
    step(doc, 1, 'Finance', 'Opens Payments in the Finance module',
         'Sees the full transaction list. Online payments have a green "Card" badge. '
         'Manual payments (cash, bank transfer) have their respective badges.')
    step(doc, 2, 'Finance', 'Clicks on an online payment row',
         'Detail view shows: parent name, student name, amount, fee, net received, Paystack reference, timestamp, and auto-reconciled status.')
    step(doc, 3, 'Finance', 'Exports payment report',
         'Clicks "Export" → downloads an Excel/CSV file of all transactions for the selected term.')
    outcome(doc, [
        'Finance Officer has a complete, auditable record of all online payments.',
        'Online payments require zero manual reconciliation — they are auto-verified.',
    ])

    page_break(doc)

    # ════════════════════════════════════════════════════════════════════════
    # PHASE 2
    # ════════════════════════════════════════════════════════════════════════

    section_header(doc, 'Phase 2 — Weeks 5–10',
        'Operational & Engagement Modules',
        'Eight modules that complete daily operations and deepen parent and teacher engagement.')

    # ─────────────────────────────────────────────────────────────────────────
    # MODULE 7: HOUSE POINT SYSTEM
    # ─────────────────────────────────────────────────────────────────────────
    module_banner(doc, 7, 'House Point System',
        'A pastoral care system where teachers award points to students; a live leaderboard shows house standings.',
        'Phase 2 — Important', AMBER, 'Low', EMERALD)

    flow_header(doc, 1, 3, 'Admin sets up houses and assigns students')
    preconditions(doc, ['Admin is logged in', 'Students are enrolled in the school'])
    step(doc, 1, 'Admin', 'Goes to Students > Houses (new menu item)',
         'Empty state: "No houses set up yet. Create your school\'s houses to get started."')
    step(doc, 2, 'Admin', 'Clicks "Create House"',
         'Form: House Name, Colour (colour picker), Emblem/Icon (upload or choose from preset icons). '
         'Admin creates four houses: Red House, Blue House, Green House, Yellow House.')
    step(doc, 3, 'Admin', 'Assigns students to houses',
         'Either manually (select student → assign house) or in bulk (select all students in a class → assign house). '
         'Each student belongs to exactly one house. House assignment is saved on the student\'s profile.')
    outcome(doc, [
        'Houses are created and students are distributed across them.',
        'House leaderboard initialises with all houses at 0 points.',
    ])

    flow_header(doc, 2, 3, 'Teacher awards a house point to a student')
    step(doc, 1, 'Teacher', 'Opens My Classes and selects JSS 2A',
         'The class roster shows all students with their house badge (a coloured dot) next to their name.')
    step(doc, 2, 'Teacher', 'Clicks the house point icon next to a student\'s name',
         'A small form appears: "Award house point to Tobi Okafor (Red House)."\n'
         'Reason dropdown: Excellent Classwork / Outstanding Conduct / Sports Achievement / Community Service / Academic Excellence / Other\n'
         'Points: 1 (default) — teacher can enter 1–5 for special achievements.')
    step(doc, 3, 'Teacher', 'Selects "Excellent Classwork" and clicks "Award"',
         'Confirmation: "1 point awarded to Tobi Okafor — Red House."')
    step(doc, 4, 'System', 'Updates house standing and notifies student',
         'HousePoints record saved. Red House total increases by 1.\n'
         'In-app notification to Tobi: "You earned 1 house point from Mr. Ibrahim for Excellent Classwork! Red House total: 47 points."', is_system=True)
    outcome(doc, [
        'House point is recorded and leaderboard updates in real time.',
        'Student is notified and feels immediate recognition.',
    ])

    flow_header(doc, 3, 3, 'Student and Parent view the leaderboard')
    step(doc, 1, 'Student', 'Logs in and sees house standings on dashboard',
         '"House Leaderboard" widget:\n'
         '1st: Blue House — 61 pts\n'
         '2nd: Red House — 47 pts\n'
         '3rd: Green House — 39 pts\n'
         '4th: Yellow House — 31 pts\n\n'
         '"Your house (Red House): 47 pts — 2nd place. Keep going!"')
    step(doc, 2, 'Parent', 'Logs in and checks child\'s house page',
         'In the Children section, a "House Points" card shows:\n'
         '"Tobi Okafor — Red House — 12 personal points this term — House rank: 2nd."')
    outcome(doc, [
        'Leaderboard creates friendly competition and motivation across all students.',
        'Parents can track their child\'s pastoral progress alongside academic results.',
    ])
    edge_cases(doc, [
        'Admin deactivates a house (e.g. restructuring) → points are preserved in history but the house is removed from the leaderboard. Students are prompted to be reassigned.',
        'Teacher accidentally awards too many points → Admin can view and reverse any HousePoints entry within 24 hours.',
    ])

    page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # MODULE 8: TEACHER PAYSLIP
    # ─────────────────────────────────────────────────────────────────────────
    module_banner(doc, 8, 'Teacher Payslip (Self-View)',
        'Teachers view and download their own monthly payslip from within CASPAA.',
        'Phase 2 — Important', AMBER, 'Low — Quick Win', EMERALD)

    flow_header(doc, 1, 1, 'Teacher views and downloads payslip')
    preconditions(doc, [
        'Finance Officer has processed payroll for the current month',
        'Teacher\'s salary and bank account are configured in their staff profile',
    ])
    step(doc, 1, 'Teacher', 'Clicks "Payslip" in the teacher menu (new item)',
         'The payslip view loads with the most recent month selected by default. '
         'A month/term selector is at the top right: "June 2026 ▼".')
    step(doc, 2, 'Teacher', 'Reads the payslip',
         'The payslip shows:\n'
         '  BRIGHT LIGHTS ACADEMY\n'
         '  ─────────────────────────────\n'
         '  Employee: Mr. Adamu Ibrahim\n'
         '  Staff ID: STF-0012\n'
         '  Month: June 2026\n'
         '  ─────────────────────────────\n'
         '  Gross Salary:         ₦180,000\n'
         '  Deductions:\n'
         '    PAYE Tax:          -₦12,600\n'
         '    Pension (8%):      -₦14,400\n'
         '    Salary Advance:    -₦20,000 [Deducted — Ref: ADV-001]\n'
         '  ─────────────────────────────\n'
         '  Net Pay:              ₦133,000\n'
         '  ─────────────────────────────\n'
         '  Bank: GTBank | Account: 023XXXXXXX\n'
         '  Payment Status: Paid on 28 June 2026')
    step(doc, 3, 'Teacher', 'Clicks "Download PDF"',
         'The browser\'s print dialog opens with the payslip formatted as a clean, printable document '
         '(school logo, payslip details, no navigation bars). Teacher saves as PDF.')
    step(doc, 4, 'Teacher', 'Checks previous months',
         'Teacher clicks the month selector and picks "May 2026". A different payslip loads with that month\'s figures.')
    outcome(doc, [
        'Teacher can access their payslip anytime without going to the Finance Office.',
        'Payslip data is read-only for the teacher — no editing possible from this view.',
    ])
    edge_cases(doc, [
        'Payroll not yet processed for the selected month → "Payslip not available yet for this month. Please check back after payroll is processed."',
        'Teacher tries to access another teacher\'s payslip by changing a URL parameter → system checks session.userId and blocks access. Shows "Unauthorised."',
    ])

    page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # MODULE 9: TRANSPORT MANAGEMENT & SECURE PICKUP
    # ─────────────────────────────────────────────────────────────────────────
    module_banner(doc, 9, 'Transport Management & Secure Pickup',
        'Admin manages bus routes. Parents register authorised pickup persons for their children.',
        'Phase 2 — Important', AMBER, 'High', RED)

    flow_header(doc, 1, 3, 'Admin sets up a bus and route')
    preconditions(doc, ['Admin is logged in', 'At least one student is enrolled'])
    step(doc, 1, 'Admin', 'Goes to Operations > Transport (new module)',
         'Two tabs: "Buses" and "Routes". Empty state on first visit.')
    step(doc, 2, 'Admin', 'Creates a bus in the Buses tab',
         'Form: Plate Number (e.g. LND 452 AA), Capacity (32 seats), Driver Name, Driver Phone, Current Status (Active / Maintenance).')
    step(doc, 3, 'Admin', 'Creates a route',
         'Route Name: "Lekki Route". Stops: Add stops in order with expected arrival time at each:\n'
         '  Stop 1: Admiralty Way Gate — 6:45 AM\n'
         '  Stop 2: Marwa Junction — 7:00 AM\n'
         '  Stop 3: Ajah Roundabout — 7:20 AM\n'
         '  Stop 4: School Gate — 7:45 AM')
    step(doc, 4, 'Admin', 'Assigns bus to route and assigns students to the route',
         'Links LND 452 AA to "Lekki Route". Selects students who use this route (can select by class or individually).')
    outcome(doc, [
        'Route and bus are set up. Students are linked to the route.',
        'Parents of those students will see the route in their portal.',
    ])

    flow_header(doc, 2, 3, 'Parent registers authorised pickup persons')
    step(doc, 1, 'Parent', 'Opens Children > Tobi Okafor > Authorised Pickup tab',
         'Shows current authorised persons. Empty on first use: "No pickup persons added yet."')
    step(doc, 2, 'Parent', 'Clicks "+ Add Pickup Person"',
         'Form: Full Name, Relationship (dropdown: Mother / Father / Sibling / Grandparent / Family Friend / Driver / Other), Phone Number, Upload Photo (required).\n\n'
         'Maximum of 3 pickup persons per child.')
    step(doc, 3, 'Parent', 'Adds "Mrs. Funke Okafor" (Grandmother) with her photo and phone',
         'Parent clicks Save.')
    step(doc, 4, 'System', 'Notifies Admin of new pickup person pending verification',
         'In-app notification to Admin: "New pickup person added for Tobi Okafor — pending verification."', is_system=True)
    step(doc, 5, 'Admin', 'Reviews and verifies the pickup person',
         'Admin opens the Pickup Persons queue, reviews the photo and details, clicks "Verify". '
         'Status changes from "Pending" to "Approved".')
    step(doc, 6, 'System', 'Notifies parent of approval',
         '"Mrs. Funke Okafor has been approved as an authorised pickup person for Tobi Okafor."', is_system=True)
    outcome(doc, [
        'Pickup person is approved and will appear in the school\'s pickup registry.',
        'Only persons on this list are authorised to collect the child from school.',
    ])

    flow_header(doc, 3, 3, 'Parent tracks bus location on pick-up day')
    step(doc, 1, 'Admin', 'Marks bus as "En Route" at departure time',
         'Admin (or school driver coordinator) opens Transport > Lekki Route, clicks "Mark Bus as Departed" at 6:40 AM.')
    step(doc, 2, 'System', 'Sends SMS to all parents on that route',
         '"Bus LND 452 AA has departed for the Lekki Route. Expected arrival at your stop:\n'
         'Admiralty Way Gate: 6:45 AM\n'
         'Marwa Junction: 7:00 AM\n'
         'Ajah Roundabout: 7:20 AM"', is_system=True)
    step(doc, 3, 'Parent', 'Opens the CASPAA app and checks Transport status',
         '"Lekki Route Bus\n'
         'Status: En Route\n'
         'Your stop: Marwa Junction — Est. arrival: 7:00 AM\n'
         'Driver: Mr. Tunde Bello — 0802 XXX XXXX"')
    outcome(doc, [
        'Parents know exactly when to expect the bus without calling the school.',
        'School has a clear record of bus dispatch times and routes.',
    ])
    edge_cases(doc, [
        'Parent adds a 4th pickup person → "Maximum of 3 authorised pickup persons per child. Please remove one before adding another."',
        'Unverified person tries to collect a child → school gate checks the CASPAA pickup list. Person not on the approved list → child is not released. Admin is alerted.',
    ])

    page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # MODULE 10: PARENT FEEDBACK TOOL
    # ─────────────────────────────────────────────────────────────────────────
    module_banner(doc, 10, 'Parent Feedback Tool',
        'Admin creates structured feedback surveys. Parents submit ratings and comments anonymously.',
        'Phase 2 — Important', AMBER, 'Low', EMERALD)

    flow_header(doc, 1, 2, 'Admin creates and publishes a feedback form')
    step(doc, 1, 'Admin', 'Goes to Communications > Feedback (new tab)',
         'Shows existing forms and their response rates. Empty state on first use.')
    step(doc, 2, 'Admin', 'Clicks "+ Create Feedback Form"',
         'Form builder:\n'
         '  • Form Title: "End of Term 1 Parent Survey 2026/27"\n'
         '  • Deadline: 15 December 2026\n'
         '  • Questions (add up to 10):\n'
         '    - Question type: Star Rating (1–5) / Yes/No / Open Text\n'
         '    - Add question text\n\n'
         'Admin adds questions:\n'
         '  1. "How satisfied are you with your child\'s academic progress this term?" — Star Rating\n'
         '  2. "Are teachers responsive when you contact them?" — Yes/No\n'
         '  3. "What can we do better next term?" — Open Text')
    step(doc, 3, 'Admin', 'Clicks "Publish Form"',
         'Form status changes to "Active". All parents with children in the school receive an in-app notification and optional SMS.')
    step(doc, 4, 'System', 'Sends notification to all parents',
         '"Bright Lights Academy has published a feedback form: End of Term 1 Parent Survey. Please complete it by 15 Dec. Your feedback is anonymous."', is_system=True)
    outcome(doc, ['Feedback form is live and accessible to all parents.'])

    flow_header(doc, 2, 2, 'Parent completes the feedback form')
    step(doc, 1, 'Parent', 'Sees alert banner on dashboard',
         '"End of Term 1 Parent Survey is open. Share your feedback — closes 15 Dec." with a "Take Survey" button.')
    step(doc, 2, 'Parent', 'Clicks "Take Survey"',
         'The form renders one question at a time for a clean, mobile-friendly experience.')
    step(doc, 3, 'Parent', 'Answers the questions',
         'Q1: Taps 4 stars.\n'
         'Q2: Taps "Yes".\n'
         'Q3: Types "More frequent progress reports would be helpful."')
    step(doc, 4, 'Parent', 'Clicks "Submit"',
         'Confirmation: "Thank you for your feedback! Your response has been submitted anonymously."')
    step(doc, 5, 'System', 'Saves response',
         'FeedbackResponses record saved. The dashboard alert banner disappears for this parent.', is_system=True)
    step(doc, 6, 'Admin', 'Checks results dashboard',
         'Form results page shows:\n'
         '  • Response rate: 43 / 187 parents (23%)\n'
         '  • Q1 Average: 3.8 / 5 stars (bar chart)\n'
         '  • Q2: 87% Yes / 13% No (pie chart)\n'
         '  • Q3: Open text responses listed (no names attached)')
    outcome(doc, [
        'Admin gets structured, anonymous data on parent satisfaction.',
        'Parents feel heard. Feedback loop improves trust between school and parents.',
    ])
    edge_cases(doc, [
        'Parent tries to submit the same form twice → "You have already submitted this survey. Thank you for your response." Second submission is blocked.',
        'Form deadline passes → form automatically closes. Parents see "This survey is now closed." Admin can re-open manually if needed.',
    ])

    page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # MODULE 11: COMMUNICATION DIARY
    # ─────────────────────────────────────────────────────────────────────────
    module_banner(doc, 11, 'Communication Diary',
        'Structured daily notes from teachers to parents, per student — distinct from the general chat system.',
        'Phase 2 — Important', AMBER, 'Medium', AMBER)

    flow_header(doc, 1, 2, 'Teacher writes a diary entry for a student')
    preconditions(doc, [
        'Teacher is assigned to at least one class',
        'Bulk SMS is configured (optional — for SMS alerts)',
    ])
    step(doc, 1, 'Teacher', 'Opens My Classes > JSS 2A > Diary tab',
         'Shows a list of students with a diary icon next to each name. '
         'An unread reply counter badge appears on students who have replied to a previous entry.')
    step(doc, 2, 'Teacher', 'Clicks the diary icon next to "Tobi Okafor"',
         'A diary panel opens on the right:\n'
         '  • Date: today (auto-filled)\n'
         '  • Category: Homework / Behaviour / Health / Academic / General\n'
         '  • Note: (textarea)\n\n'
         'Previous diary entries for Tobi are shown below in reverse-chronological order, each showing date, category, the note, and parent\'s reply if any.')
    step(doc, 3, 'Teacher', 'Writes the entry',
         'Category = Homework\n'
         'Note = "Tobi did not submit the Mathematics assignment due today. '
         'Please remind him to complete it and submit by Thursday."')
    step(doc, 4, 'Teacher', 'Clicks "Send to Parent"',
         'Entry saved. Status = "Unread by Parent".')
    step(doc, 5, 'System', 'Sends SMS alert to parent (if SMS is configured)',
         '"New diary note from Mr. Ibrahim for Tobi Okafor (Homework). Log in to CASPAA to read and reply."', is_system=True)
    outcome(doc, [
        'Diary entry is created and visible to the parent immediately.',
        'Parent receives an SMS nudge to check the app.',
    ])

    flow_header(doc, 2, 2, 'Parent reads the diary entry and replies')
    step(doc, 1, 'Parent', 'Opens CASPAA — sees notification badge on "Diary"',
         'A new "Diary" item appears in the parent menu with a red badge "1".')
    step(doc, 2, 'Parent', 'Clicks Diary > Tobi Okafor',
         'Timeline of diary entries for Tobi. The newest entry from today is at the top, with a blue "UNREAD" tag.')
    step(doc, 3, 'Parent', 'Reads the entry',
         '"Homework — Today — Mr. Ibrahim\n'
         'Tobi did not submit the Mathematics assignment due today. Please remind him to complete it and submit by Thursday."\n\n'
         'As the parent reads it, the entry status automatically changes to "Read" and a timestamp is recorded.')
    step(doc, 4, 'Parent', 'Types a reply',
         '"Thank you for letting me know. I will speak with him tonight and ensure he submits by Thursday."')
    step(doc, 5, 'Parent', 'Clicks "Reply"',
         'Reply is saved against the diary entry.')
    step(doc, 6, 'System', 'Notifies teacher of the parent\'s reply',
         'In-app notification to the teacher: "Tunde Okafor replied to your diary note for Tobi Okafor."', is_system=True)
    step(doc, 7, 'Teacher', 'Sees the reply on the diary panel',
         'The diary entry now shows the original note and the parent\'s reply below it, with a timestamp. '
         'The unread reply badge on the class roster disappears once the teacher reads it.')
    outcome(doc, [
        'A full written record of teacher–parent communication exists per student, per date, with read receipts.',
        'Both teacher and parent know the other has seen the note.',
        'This diary is separate from the chat system — it is structured record-keeping, not conversation.',
    ])
    edge_cases(doc, [
        'Parent does not read the entry for 3 days → a follow-up SMS auto-fires: "You have an unread diary note from Mr. Ibrahim for Tobi Okafor." (configurable threshold).',
        'Teacher submits an empty note → "Note content cannot be empty." Submission blocked.',
    ])

    page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # MODULE 12: FORMATIVE ASSESSMENT
    # ─────────────────────────────────────────────────────────────────────────
    module_banner(doc, 12, 'Formative Assessment',
        'Teachers record ongoing in-class scores (class tests, oral quizzes, projects) that feed into the term report.',
        'Phase 2 — Important', AMBER, 'Medium', AMBER)

    flow_header(doc, 1, 2, 'Teacher creates a formative assessment and enters scores')
    step(doc, 1, 'Teacher', 'Goes to Results > Formative tab (new tab alongside existing Results)',
         'List of all formative assessments created this term. Empty on first use.')
    step(doc, 2, 'Teacher', 'Clicks "+ New Assessment"',
         'Form:\n'
         '  • Class: JSS 2A\n'
         '  • Subject: Mathematics\n'
         '  • Assessment Type: Class Test / Weekly Quiz / Oral Assessment / Project / Practical\n'
         '  • Title: "Week 8 Class Test — Quadratic Equations"\n'
         '  • Maximum Score: 20\n'
         '  • Date: Today')
    step(doc, 3, 'Teacher', 'Clicks "Create & Enter Scores"',
         'A score-entry table appears:\n'
         '  Student Name          | Score (/20) | Remarks\n'
         '  Tobi Okafor           | ___         | ___\n'
         '  Amina Bello           | ___         | ___\n'
         '  Chukwuemeka Eze       | ___         | ___\n'
         '  (all 32 students listed)\n\n'
         'Teacher types scores directly into the table cells.')
    step(doc, 4, 'Teacher', 'Fills in scores and optional remarks',
         'Tobi = 16/20, Amina = 18/20, Chukwuemeka = 11/20 "Needs to revise factorisation", etc.')
    step(doc, 5, 'Teacher', 'Clicks "Save Scores"',
         'All scores saved. The assessment row in the list shows: "Week 8 Class Test — 32/32 scores entered — Avg: 14.2/20".')
    outcome(doc, [
        'Formative assessment scores are saved and immediately visible to students.',
        'Scores feed into the cumulative CA column on the end-of-term report card.',
    ])

    flow_header(doc, 2, 2, 'Student views their formative scores')
    step(doc, 1, 'Student', 'Goes to Results > Assessments tab',
         'A table shows all formative assessments for the current term:\n'
         '  Week 2 Quiz — Algebra — 8/10 — Class Avg: 7.1\n'
         '  Week 5 Project — Geometry — 18/20 — Class Avg: 15.6\n'
         '  Week 8 Class Test — Quadratic Equations — 16/20 — Class Avg: 14.2\n\n'
         'Student can see their score relative to the class average for context.')
    step(doc, 2, 'Student', 'Clicks an assessment row for details',
         'Shows their score, the teacher\'s remark (if any), and the full score distribution (bar chart showing how many students scored each range).')
    outcome(doc, [
        'Students stay informed of their ongoing academic performance throughout the term — not just at results time.',
        'Teacher remarks give actionable guidance rather than just a number.',
    ])
    edge_cases(doc, [
        'Teacher enters a score higher than the maximum → "Score cannot exceed maximum (20). Please correct."',
        'Teacher saves with some scores blank → warning: "12 students have no score entered. Save anyway?" Teacher can choose to save partial scores.',
    ])

    page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # MODULE 13: HEALTH / SICKBAY
    # ─────────────────────────────────────────────────────────────────────────
    module_banner(doc, 13, 'Health Management / Sickbay',
        'Sickbay staff log student health visits. Parents are notified automatically. Health profiles are maintained.',
        'Phase 2 — Important', AMBER, 'Medium', AMBER)

    flow_header(doc, 1, 2, 'Student visits the sickbay')
    preconditions(doc, [
        'Student has a health profile on CASPAA (blood group, allergies, emergency contact)',
        'Bulk SMS is configured for parent notifications',
    ])
    step(doc, 1, 'Admin', 'Opens Operations > Sickbay',
         'Today\'s sickbay log is shown. Any students currently in the sickbay are listed with their check-in time.')
    step(doc, 2, 'Admin', 'Clicks "+ Check In Student"',
         'Searchable student lookup field. Types "Tobi" — Tobi Okafor appears with class and photo.')
    step(doc, 3, 'Admin', 'Selects Tobi and fills in the visit details',
         'Fields:\n'
         '  • Time In: (auto-filled with current time)\n'
         '  • Complaint: (text field) "Headache and mild fever"\n'
         '  • Temperature: 37.8°C\n'
         '  • Treatment Given: "Paracetamol 500mg administered"\n'
         '  • Referred Out: No\n\n'
         'Tobi\'s health profile panel appears on the right: Blood Group = O+, Allergies = Penicillin.\n'
         '"ALERT: Student is allergic to Penicillin" — shown in red.')
    step(doc, 4, 'Admin', 'Clicks "Check In & Notify Parent"',
         'Visit saved. The "Notify Parent" checkbox is ticked by default.')
    step(doc, 5, 'System', 'Sends SMS to parent immediately',
         '"Hello Mr. Tunde, your child Tobi Okafor visited the sickbay at 10:23 AM today. '
         'Complaint: Headache and mild fever. Treatment: Paracetamol administered. '
         'They are resting and will return to class shortly. For more info call: 0803-XXX-XXXX. — Bright Lights Academy."', is_system=True)
    step(doc, 6, 'Admin', 'Checks Tobi out when recovered',
         'Clicks "Check Out" next to Tobi\'s name. Time Out recorded. Visit is closed with duration: "1 hour 15 mins".')
    outcome(doc, [
        'Full sickbay visit record is saved with complaint, treatment, and duration.',
        'Parent is informed within seconds of their child\'s visit — without needing to call the school.',
        'Allergy alerts prevent dangerous medication errors.',
    ])

    flow_header(doc, 2, 2, 'Parent updates child\'s health profile')
    step(doc, 1, 'Parent', 'Opens Children > Tobi Okafor > Health tab',
         'Shows current health profile: Blood Group, Known Allergies, Chronic Conditions, Emergency Contact, Doctor Name/Phone, Vaccinations.')
    step(doc, 2, 'Parent', 'Clicks "Edit Health Profile"',
         'All fields become editable. Parent adds a new allergy: "Latex" and updates the family doctor\'s phone number.')
    step(doc, 3, 'Parent', 'Clicks "Save"',
         'Health profile updated. Changes are immediately visible to sickbay staff when they next view the student.')
    outcome(doc, [
        'Parent controls and maintains accurate medical information for their child.',
        'Sickbay staff always have up-to-date health information before administering treatment.',
    ])
    edge_cases(doc, [
        'Admin checks in a student who is already checked in → "Tobi Okafor is already logged in the sickbay (since 10:23 AM). Do you want to add a new visit or update the existing one?"',
        'Parent notification SMS fails (network error) → visit is still saved. A "Notification failed" flag appears on the visit row. Admin can manually retry the SMS.',
    ])

    page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # MODULE 14: INVENTORY MANAGEMENT
    # ─────────────────────────────────────────────────────────────────────────
    module_banner(doc, 14, 'Inventory Management',
        'Admin tracks school assets — furniture, electronics, stationery — and manages issues and restocks.',
        'Phase 2 — Important', AMBER, 'Medium', AMBER)

    flow_header(doc, 1, 2, 'Admin adds new stock items')
    step(doc, 1, 'Admin', 'Goes to Operations > Inventory',
         'Item catalogue table. Categories filter at top: All | Furniture | Electronics | Stationery | Equipment | Uniforms.')
    step(doc, 2, 'Admin', 'Clicks "+ Add Item"',
         'Form:\n'
         '  • Item Name: "Student Chair"\n'
         '  • Category: Furniture\n'
         '  • Quantity in Stock: 120\n'
         '  • Unit: "pieces"\n'
         '  • Reorder Level: 20 (alert when stock falls below this)\n'
         '  • Storage Location: "Block A Store Room"')
    step(doc, 3, 'Admin', 'Saves the item',
         'Item appears in the catalogue. A green "Good Stock" badge shows because 120 > 20 reorder level.')
    outcome(doc, ['Item is in the catalogue and stock is tracked going forward.'])

    flow_header(doc, 2, 2, 'Admin issues items and receives a low-stock alert')
    step(doc, 1, 'Admin', 'Clicks "Issue Items" on the Student Chair row',
         'Form: Issued To (search for staff or class), Quantity, Purpose, Date.')
    step(doc, 2, 'Admin', 'Issues 8 chairs to JSS 2A classroom',
         '"Issued To: JSS 2A Classroom, Quantity: 8, Purpose: Classroom setup." Clicks "Issue".')
    step(doc, 3, 'System', 'Updates stock and logs the transaction',
         'Stock: 120 → 112. InventoryTransactions record: { type: "Issue", quantity: 8, issuedTo: "JSS 2A", recordedBy: Admin, date }.', is_system=True)
    step(doc, 4, 'Admin', 'Over time, more issues bring the stock to 18 — below the reorder level',
         'Item row changes from green "Good Stock" to red "Low Stock (18 left — reorder at 20)".')
    step(doc, 5, 'System', 'Sends low-stock alert to Admin',
         'Dashboard alert: "3 inventory items are below their reorder level. View Inventory."', is_system=True)
    step(doc, 6, 'Admin', 'Clicks "Restock" on the Student Chair row',
         'Form: Quantity Received (e.g. 50), Supplier (optional), Date. Clicks "Save Restock".')
    step(doc, 7, 'System', 'Updates stock',
         'Stock: 18 → 68. Transaction logged as Type = "Stock-In". Badge returns to green.', is_system=True)
    outcome(doc, [
        'Admin always knows what is in stock and where it was issued.',
        'Low-stock alerts prevent running out of critical items.',
        'Every movement in and out of stock is logged for accountability.',
    ])
    edge_cases(doc, [
        'Admin tries to issue more items than are in stock → "Insufficient stock. Only 18 Student Chairs available." Issue is blocked.',
        'Admin writes off damaged items → uses "Write-Off" transaction type. Stock decreases with a reason noted (e.g. "4 chairs damaged in flood"). Visible in audit log.',
    ])

    # ── Final note ────────────────────────────────────────────────────────────
    page_break(doc)
    section_header(doc, 'Implementation Notes', 'Before You Build — Key Principles')

    notes = [
        ('Build Phase 1 in order', 'Module 2 (SMS/Email) must come before Modules 11, 13, and 9 — they all fire SMS alerts. Class Notes (3) should be built before Video Lessons (4) since they share the same Learning hub UI.'),
        ('Every System step should be logged', 'All automatic actions — SMS sends, payment verifications, notification dispatches — must write to an audit log. This protects the school in disputes and helps support staff debug issues.'),
        ('Mobile-first for parents and students', 'Parents and students primarily access CASPAA on mobile. Every new flow should be tested at 375px width before being considered done.'),
        ('Error states are not optional', 'Every form that touches money (Module 6), communication (Modules 2, 11), or health (Module 13) must handle failure gracefully. A failed SMS or payment must never silently disappear.'),
        ('Existing messaging/chat system is NOT replaced', 'The Communication Diary (Module 11) is a structured record-keeping tool, not a replacement for the existing chat. Both coexist — parents can use chat for quick questions and the diary for formal teacher notes about their child.'),
    ]
    for title, body in notes:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
        add_run(p, title, bold=True, size=11, color=DARK_NAVY)
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Inches(0.2)
        p2.paragraph_format.space_after = Pt(8)
        add_run(p2, body, size=10, color=GRAY_TEXT)

    doc.save(r'C:\Users\USER\Desktop\CASPAA\CASPAA_Implementation_Roadmap.docx')
    print('Done: Implementation Roadmap (v2 — real-life flows) saved.')


if __name__ == '__main__':
    build()
