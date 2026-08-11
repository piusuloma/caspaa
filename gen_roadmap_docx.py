# -*- coding: utf-8 -*-
"""CASPAA product roadmap as Office Open XML (.docx).
A BUILD PLAN: per area, features listed in PRIORITY ORDER (1 = highest).
No phases. Nothing marked built. Lending is on hold."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"C:\Users\USER\Desktop\CASPAA\CASPAA_Roadmap.docx"

NAVY = RGBColor(0x10, 0x27, 0x30)
TEAL = RGBColor(0x0A, 0x6B, 0x50)
ACC  = RGBColor(0x0E, 0x9E, 0x76)
HOLD = RGBColor(0x8E, 0x4E, 0x5C)
GREY = RGBColor(0x5C, 0x6E, 0x74)
BODY = RGBColor(0x28, 0x38, 0x3E)
FILL_NAVY = "10333F"
FILL_SOFT = "EDF3F1"
FILL_HOLDS = "F3E7EA"

doc = Document()
for s in doc.sections:
    s.left_margin = s.right_margin = Inches(0.7)
    s.top_margin = s.bottom_margin = Inches(0.7)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"; normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.12


def shade(el, fill):
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), fill); el.append(sh)


def cell_fill(cell, fill):
    shade(cell._tc.get_or_add_tcPr(), fill)


def cell_borders(cell, color="BFC9C6"):
    tcPr = cell._tc.get_or_add_tcPr(); b = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "4"); e.set(qn("w:color"), color)
        b.append(e)
    tcPr.append(b)


def set_widths(table, widths):
    table.autofit = False
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w


def run(p, text, *, bold=False, italic=False, color=None, size=None, caps=False):
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    if color is not None: r.font.color.rgb = color
    if size is not None: r.font.size = Pt(size)
    if caps: r.font.all_caps = True
    return r


# ---------------- Title ----------------
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
run(p, "CASPAA — Product Roadmap", bold=True, size=22, color=NAVY)
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
run(p, "SCHOOL OPERATING SYSTEM  ·  BUILD PLAN, IN PRIORITY ORDER", bold=True, size=10, color=TEAL, caps=True)
p = doc.add_paragraph()
run(p, "As of 27 July 2026  ·  Owner: Product  ·  Status: Draft for stakeholder review", italic=True, color=GREY, size=9.5)

# Note callout
tbl = doc.add_table(rows=1, cols=1); c = tbl.rows[0].cells[0]
cell_borders(c, "E6D9A6"); cell_fill(c, "FFF7E0")
np = c.paragraphs[0]
run(np, "This is a prioritised build plan. ", bold=True)
run(np, "Within each area, features are listed in priority order — number 1 is built first, and so on down the list. Nothing here is built or live yet. The Lending engine is on hold this cycle by decision.")
doc.add_paragraph()

# ---------------- Priority lists by area ----------------
p = doc.add_paragraph(); run(p, "Build priority by area", bold=True, size=13, color=NAVY)

AREAS = [
    ("Students & Academics", "Enrolment, teaching, records", [
        "Student records & admissions", "Online admission applications", "Attendance register",
        "Results entry & approval", "Report cards & broadsheet", "Academic calendar & terms",
        "Class arms & timetable", "Scheme of work", "Assessments — CBT", "Assignments & inline marking",
        "Lesson plans & learning materials", "Library", "Alumni & graduation",
        "Bulk promotion / graduation", "Report-card templates", "Discipline & behaviour records",
        "Academic analytics & insights", "Expanded content library"]),
    ("Finance & Payments", "Fees, ledger, payroll, store", [
        "Fee structure", "Invoicing", "Itemised student ledger / statement", "Discounts breakdown",
        "Expenses", "School store (student purchases)", "Live Paystack payments (card/transfer/USSD)",
        "Bank reconciliation", "Real-time receipts", "Installment plans", "Financial reports",
        "Cost-centre P&L analytics", "Budgeting & forecasting"]),
    ("Staff & HR", "Workforce, payroll, appraisals", [
        "Staff profiles & roles", "Payroll & payslips", "Leave requests",
        "Appraisals", "Substitute coverage", "Staff attendance", "HR analytics"]),
    ("Communication & Engagement", "Parents, students, community", [
        "Announcements", "Notice board", "Calendar events", "Teacher-parent diary", "Consent forms",
        "Surveys & feedback", "House points", "Comms oversight",
        "Live WhatsApp / SMS / email delivery", "SMS campaigns", "In-app messaging / richer engagement"]),
    ("Operations & School Services", "Front desk, inventory, health, transport", [
        "Front desk & visitor log", "Inventory management", "Health & sickbay",
        "Transport routes & assignments", "Authorised pickup approval",
        "Real-time bus status push", "Inventory requests & stock alerts", "Asset & facilities management"]),
    ("Platform, Multi-branch & Operations Portal (COP)", "Groups, plans, admin, COP", [
        "Multi-branch school groups", "Feature entitlements & plan gating", "School onboarding (COP)",
        "School settings & branding", "Role dashboards", "Audit log", "Help & support (school-side)",
        "Live billing tiers & subscriptions", "Support desk & SLAs", "COP user management & RBAC",
        "System control (feature flags, quotas)", "Persist active branch",
        "Advanced BI & analytics", "Scheduled reports", "Revenue analytics"]),
    ("Foundations & Infrastructure", "Backend, auth, offline, mobile", [
        "Production backend (PostgreSQL + API)", "Role-aware unified login & RBAC",
        "Multi-tenant isolation", "Offline-first PWA", "Auth hardening & encryption",
        "Payments infrastructure (gateways, webhooks)", "Notifications infrastructure",
        "Native mobile app", "Offline sync hardening", "Scale to 10k+ schools"]),
]

table = doc.add_table(rows=1, cols=2)
hdr = table.rows[0].cells
for i, h in enumerate(["Area", "Features — highest priority first"]):
    cell_borders(hdr[i]); cell_fill(hdr[i], FILL_NAVY)
    run(hdr[i].paragraphs[0], h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=9.5)

for name, sub, feats in AREAS:
    cells = table.add_row().cells
    cell_borders(cells[0]); cell_fill(cells[0], FILL_SOFT)
    run(cells[0].paragraphs[0], name, bold=True, size=10, color=NAVY)
    run(cells[0].add_paragraph(), sub, size=8.5, color=GREY)
    cell_borders(cells[1])
    first = True
    for idx, feat in enumerate(feats, 1):
        p = cells[1].paragraphs[0] if first else cells[1].add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(2)
        run(p, f"{idx}. ", bold=True, color=ACC, size=9.5)
        run(p, feat, size=9.5, color=BODY)

set_widths(table, [Inches(1.9), Inches(6.4)])

# ---------------- Portals by role ----------------
doc.add_paragraph()
p = doc.add_paragraph(); run(p, "Portals by role", bold=True, size=13, color=NAVY)
p = doc.add_paragraph()
run(p, "One platform, a role-aware experience — each person lands in the portal built for them. Every module above is delivered to the roles that need it.", color=GREY)
ROLES = [
    ("School Owner / Proprietor (Admin)", "Everything — students, academic, staff & HR, finance, operations, communications, settings & branding, multi-branch groups, reports, audit."),
    ("Principal", "Students, academic, results approval, staff & HR, attendance oversight, calendar, house points, health, reports."),
    ("Finance Officer / Bursar", "Fee structure, invoicing, itemised ledger / statements, discounts, payments & reconciliation, expenses, payroll, school store, financial reports."),
    ("Teacher", "Attendance, results entry, assignments & marking, CBT / assessments, lesson plans & materials, timetable, diary, house points, my payslip, leave, appraisal."),
    ("Parent", "Per-child dashboards, fees & wallet / ledger, results & report cards, timetable, diary, consent, surveys, transport & pickup, health, house points."),
    ("Student", "Dashboard, learning & materials, assignments (submit / resubmit), CBT / assessments, my results, timetable, behaviour, house points, my wallet."),
    ("Super Admin (CASPAA Operations / COP)", "School onboarding & lifecycle, revenue, support desk, feature flags & plans, analytics, user management / RBAC, audit."),
]
rt = doc.add_table(rows=1, cols=2)
rh = rt.rows[0].cells
for i, h in enumerate(["Role", "Modules in their portal"]):
    cell_borders(rh[i]); cell_fill(rh[i], FILL_NAVY)
    run(rh[i].paragraphs[0], h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=9.5)
for name, mods in ROLES:
    c = rt.add_row().cells
    cell_borders(c[0]); cell_borders(c[1]); cell_fill(c[0], FILL_SOFT)
    run(c[0].paragraphs[0], name, bold=True, size=9.5, color=NAVY)
    run(c[1].paragraphs[0], mods, size=9.5, color=BODY)
set_widths(rt, [Inches(2.3), Inches(6.0)])

# ---------------- On hold band ----------------
doc.add_paragraph()
p = doc.add_paragraph(); run(p, "On hold — decided, not forgotten", bold=True, size=13, color=HOLD)
p = doc.add_paragraph()
run(p, "Embedded Finance — Lending. ", bold=True, color=HOLD)
run(p, "Not being built this cycle by decision. It stays documented so it can resume later; when it does, loan access will remain gated until the business greenlights it.", color=GREY)
ht = doc.add_table(rows=1, cols=3)
HOLD_ITEMS = ["Parent fee loans & credit scoring", "Approval & disbursement", "Repayment & portfolio risk (PAR)"]
for i, txt in enumerate(HOLD_ITEMS):
    c = ht.rows[0].cells[i]; cell_borders(c, "D8B7BE"); cell_fill(c, FILL_HOLDS)
    run(c.paragraphs[0], txt, bold=True, size=9.5, color=NAVY)
    run(c.add_paragraph(), "ON HOLD", bold=True, size=8, color=HOLD, caps=True)
set_widths(ht, [Inches(2.77), Inches(2.77), Inches(2.76)])

# ---------------- Future ideas (out of scope) ----------------
doc.add_paragraph()
p = doc.add_paragraph(); run(p, "Future ideas — out of scope this build", bold=True, size=13, color=NAVY)
p = doc.add_paragraph()
run(p, "Bulk disbursement  ·  NIBSS direct settlement  ·  Savings products  ·  Investment products  ·  Cross-border lending  ·  AI underwriting", color=GREY)

doc.add_paragraph()
f = doc.add_paragraph()
run(f, "CASPAA · AfriSprings Resources Ltd.  ·  A prioritised build plan — items are proposed work in priority order, not completed features or fixed dates.", italic=True, size=8.5, color=GREY)

doc.save(OUT)
print("Wrote", OUT)
