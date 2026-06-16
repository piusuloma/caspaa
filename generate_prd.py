from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colour palette ─────────────────────────────────────────────────────────
DARK_NAVY   = RGBColor(0x0a, 0x25, 0x40)
EMERALD     = RGBColor(0x0d, 0x6e, 0x3f)
LIGHT_MINT  = RGBColor(0xd1, 0xfa, 0xe5)
RED         = RGBColor(0xdc, 0x26, 0x26)
LIGHT_RED   = RGBColor(0xfe, 0xe2, 0xe2)
AMBER       = RGBColor(0xd9, 0x77, 0x06)
LIGHT_AMBER = RGBColor(0xfe, 0xf3, 0xc7)
BLUE        = RGBColor(0x1d, 0x4e, 0xd8)
LIGHT_BLUE  = RGBColor(0xdb, 0xea, 0xfe)
GRAY_BG     = RGBColor(0xf3, 0xf4, 0xf6)
GRAY_TEXT   = RGBColor(0x6b, 0x72, 0x80)
WHITE       = RGBColor(0xff, 0xff, 0xff)
MID_GRAY    = RGBColor(0x37, 0x41, 0x51)
PURPLE      = RGBColor(0x7c, 0x3a, 0xed)
LIGHT_PURPLE= RGBColor(0xed, 0xe9, 0xfe)
TEAL        = RGBColor(0x0f, 0x76, 0x6e)
LIGHT_TEAL  = RGBColor(0xcc, 0xfb, 0xf1)


# ── Helpers ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  str(rgb))
    for ex in tcPr.findall(qn('w:shd')):
        tcPr.remove(ex)
    tcPr.append(shd)


def set_col_widths(table, widths):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths):
                cell.width = widths[i]


def set_doc_margins(doc, top=1.0, bottom=1.0, left=1.1, right=1.1):
    for section in doc.sections:
        section.top_margin    = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin   = Inches(left)
        section.right_margin  = Inches(right)


def para(doc, text='', bold=False, size=11, color=DARK_NAVY,
         align=WD_ALIGN_PARAGRAPH.LEFT, sb=0, sa=6, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        r.font.color.rgb = color
    return p


def heading(doc, text, level=1, color=DARK_NAVY, sb=14, sa=4):
    sizes = {1: 22, 2: 15, 3: 12, 4: 11}
    return para(doc, text, bold=True, size=sizes.get(level, 11),
                color=color, sb=sb, sa=sa)


def label(doc, text, color=EMERALD):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = color
    return p


def bullet(doc, text, bold_prefix=None, color=MID_GRAY, size=10.5, indent=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.3 + indent * 0.2)
    if bold_prefix:
        rb = p.add_run(bold_prefix + ' ')
        rb.bold = True
        rb.font.size = Pt(size)
        rb.font.color.rgb = DARK_NAVY
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def hdr_row(table, cells_data, bg, text_color=WHITE, bold=True, size=9.5):
    row = table.add_row()
    for i, (cell, data) in enumerate(zip(row.cells, cells_data)):
        set_cell_bg(cell, bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        r = p.add_run(str(data))
        r.bold = bold
        r.font.size = Pt(size)
        r.font.color.rgb = text_color
    return row


def data_row(table, cells_data, bg=WHITE, text_color=MID_GRAY, size=9.5, bold_first=True):
    row = table.add_row()
    for i, (cell, data) in enumerate(zip(row.cells, cells_data)):
        set_cell_bg(cell, bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(str(data))
        r.bold = (i == 0 and bold_first)
        r.font.size = Pt(size)
        r.font.color.rgb = DARK_NAVY if (i == 0 and bold_first) else text_color
    return row


def page_break(doc):
    doc.add_page_break()


def flow_step(doc, number, text, color=EMERALD):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.2)
    rn = p.add_run(f'{number}.  ')
    rn.bold = True
    rn.font.size = Pt(10)
    rn.font.color.rgb = color
    rt = p.add_run(text)
    rt.font.size = Pt(10)
    rt.font.color.rgb = MID_GRAY
    return p


def actor_tag(doc, actor_label, color=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(f'  {actor_label}  ')
    r.bold = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = color
    return p


# ════════════════════════════════════════════════════════════════════════════
# BUILD PRD
# ════════════════════════════════════════════════════════════════════════════

def build_prd():
    doc = Document()
    set_doc_margins(doc)
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ── COVER ────────────────────────────────────────────────────────────────
    cover = doc.add_table(rows=1, cols=1)
    cover.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = cover.rows[0].cells[0]
    set_cell_bg(c, DARK_NAVY)
    from docx.oxml import OxmlElement as OE
    tr = cover.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trH = OE('w:trHeight')
    trH.set(qn('w:val'), '9000')
    trPr.append(trH)

    def cp(text, size, color, bold=False, sa=8, sb=0, align=WD_ALIGN_PARAGRAPH.CENTER):
        pp = c.add_paragraph()
        pp.alignment = align
        pp.paragraph_format.space_before = Pt(sb)
        pp.paragraph_format.space_after  = Pt(sa)
        rr = pp.add_run(text)
        rr.bold = bold
        rr.font.size = Pt(size)
        rr.font.color.rgb = color
        return pp

    cp('CASPAA', 40, WHITE, bold=True, sb=50, sa=4)
    cp('SCHOOL OPERATING SYSTEM  ·  AFRISPRINGS', 9, RGBColor(0x34, 0xd3, 0x99), sa=20)
    cp('Product Requirements Document', 26, WHITE, bold=True, sa=6)
    cp('User Stories & User Flows', 18, RGBColor(0xb0, 0xc4, 0xde), sa=24)
    cp('Version 2.0  ·  June 2026  ·  Confidential', 9, RGBColor(0x90, 0xa8, 0xc0), sa=4)
    cp('Prepared by: Product & Engineering  |  Audience: Founders, Engineers, Investors', 9,
       RGBColor(0x70, 0x90, 0xb0))

    page_break(doc)

    # ── TABLE OF CONTENTS ────────────────────────────────────────────────────
    label(doc, '00 — Contents')
    heading(doc, 'Table of Contents', 1)
    toc = [
        ('01', 'Product Overview',           'Vision, architecture, value proposition'),
        ('02', 'User Roles',                 'Seven distinct roles and their access levels'),
        ('03', 'Module Index',               '20 modules mapped to their key actors'),
        ('04', 'User Stories — Admin',       'School Admin, Principal, Finance Officer'),
        ('05', 'User Stories — Teachers',    'All teacher-facing stories'),
        ('06', 'User Stories — Parents',     'All parent-facing stories'),
        ('07', 'User Stories — Students',    'All student-facing stories'),
        ('08', 'User Stories — Super Admin', 'Platform operator stories'),
        ('09', 'User Flows',                 '18 end-to-end flows with step-by-step sequences'),
        ('10', 'Permissions Matrix',         'All roles × all features'),
        ('11', 'Data Model',                 '56+ collections grouped by domain'),
    ]
    t = doc.add_table(rows=0, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_row(t, ['Section', 'Title', 'Contents'], DARK_NAVY)
    for i, (sec, title, desc) in enumerate(toc):
        bg = GRAY_BG if i % 2 else WHITE
        data_row(t, [sec, title, desc], bg)
    set_col_widths(t, [Inches(0.7), Inches(2.2), Inches(3.6)])
    page_break(doc)

    # ── 01 PRODUCT OVERVIEW ──────────────────────────────────────────────────
    label(doc, '01 — Product Overview')
    heading(doc, 'CASPAA — School Operating System')
    para(doc,
         'CASPAA is a fully integrated School Operating System for African schools. It combines a School ERP '
         '(student information, academic management, HR, operations) with embedded financial services '
         '(fee collection, lending, payroll) into a single platform accessible to all school stakeholders — '
         'administrators, teachers, parents, and students.',
         size=11, color=GRAY_TEXT, sa=10)

    heading(doc, 'Core Value Proposition', 2, sb=10)
    vp = [
        ('For Schools',    'Replace spreadsheets, paper registers, and disconnected tools with one integrated system.'),
        ('For Parents',    'Pay fees, track children, communicate with school — from any phone, anytime.'),
        ('For Teachers',   'Mark attendance, enter results, and communicate with parents without paperwork.'),
        ('For CASPAA',     'SaaS subscription revenue + transaction revenue from embedded payments and lending.'),
    ]
    for role, desc in vp:
        bullet(doc, desc, bold_prefix=role + ':', color=MID_GRAY)

    heading(doc, 'Technical Architecture', 2, sb=10)
    arch = [
        ('Frontend',      'Vanilla JS, Tailwind CSS, Chart.js — single-page app, no build step required'),
        ('Storage',       'LocalStorage mock DB (prototype) → Cloud DB (production)'),
        ('Auth',          'Multi-role session with role-based access control and 2FA for sensitive roles'),
        ('Payments',      'Paystack integration — card, bank transfer, USSD'),
        ('Notifications', 'In-app push notifications + WhatsApp deep-links'),
        ('AI',            'Claude API for report card comment generation and narration matching'),
    ]
    for k, v in arch:
        bullet(doc, v, bold_prefix=k + ':', color=MID_GRAY)
    page_break(doc)

    # ── 02 USER ROLES ────────────────────────────────────────────────────────
    label(doc, '02 — User Roles')
    heading(doc, 'User Roles & Access Levels')

    roles_data = [
        ('Super Admin',     'CASPAA Staff',        'Platform-wide',   'School management, revenue, support, audit'),
        ('School Admin',    'Proprietor / Bursar', 'School-wide',     'All modules — full access across every feature'),
        ('Principal',       'Head Teacher',        'School-wide',     'Academic + HR management, no finance module'),
        ('Finance Officer', 'Bursar / Accounts',   'Finance only',    'Fees, payments, expenses, payroll, lending'),
        ('Teacher',         'Class Teacher',       'Class-level',     'Teaching, attendance, results, communication'),
        ('Parent',          'Guardian',            'Own children',    'Fees, academics, transport, communication'),
        ('Student',         'Enrolled Student',    'Own record only', 'Learning, results, timetable, house points'),
    ]
    t = doc.add_table(rows=0, cols=4)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_row(t, ['Role', 'Who They Are', 'Scope', 'Primary Concern'], DARK_NAVY)
    colors = [LIGHT_MINT, LIGHT_BLUE, LIGHT_TEAL, LIGHT_AMBER, LIGHT_PURPLE, LIGHT_RED, GRAY_BG]
    for i, (role, who, scope, concern) in enumerate(roles_data):
        row = t.add_row()
        set_cell_bg(row.cells[0], colors[i])
        for j in range(1, 4):
            set_cell_bg(row.cells[j], GRAY_BG if i % 2 else WHITE)
        for j, val in enumerate([role, who, scope, concern]):
            row.cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = row.cells[j].paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            r = p.add_run(val)
            r.bold = (j == 0)
            r.font.size = Pt(10)
            r.font.color.rgb = DARK_NAVY
    set_col_widths(t, [Inches(1.3), Inches(1.4), Inches(1.3), Inches(2.5)])
    page_break(doc)

    # ── 03 MODULE INDEX ──────────────────────────────────────────────────────
    label(doc, '03 — Module Index')
    heading(doc, '20 Modules — At a Glance')
    modules = [
        ('1',  'Students & Admissions',      'Admin, Principal'),
        ('2',  'Academic Management',        'Admin, Principal, Teacher'),
        ('3',  'Attendance',                 'Teacher, Admin'),
        ('4',  'Results & Reporting',        'Teacher, Admin, Student, Parent'),
        ('5',  'Staff & HR',                 'Admin, Principal'),
        ('6',  'Finance & Fees',             'Finance Officer, Admin, Parent'),
        ('7',  'Lending Engine',             'Finance Officer, Admin, Parent'),
        ('8',  'Payments',                   'Parent, Finance Officer'),
        ('9',  'Communications',             'All roles'),
        ('10', 'Calendar & Notice Board',    'Admin, Principal — all roles view'),
        ('11', 'House Points',               'Admin, Teacher, Student, Parent'),
        ('12', 'Transport & Pickup',         'Admin, Parent'),
        ('13', 'Health & Sickbay',           'Admin, Parent'),
        ('14', 'Alumni',                     'Admin'),
        ('15', 'Surveys & Feedback',         'Admin, Parent'),
        ('16', 'Communication Diary',        'Teacher, Parent'),
        ('17', 'Inventory',                  'Admin'),
        ('18', 'Formative Assessments/CBT',  'Teacher, Student'),
        ('19', 'Payslips & Payroll',         'Admin, Finance, Teacher'),
        ('20', 'Super Admin Portal',         'CASPAA Staff only'),
    ]
    t = doc.add_table(rows=0, cols=3)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_row(t, ['#', 'Module', 'Key Actors'], DARK_NAVY)
    for i, (num, mod, actors) in enumerate(modules):
        bg = GRAY_BG if i % 2 else WHITE
        data_row(t, [num, mod, actors], bg)
    set_col_widths(t, [Inches(0.4), Inches(2.8), Inches(3.3)])
    page_break(doc)

    # ── 04 USER STORIES — ADMIN ──────────────────────────────────────────────
    label(doc, '04 — User Stories')
    heading(doc, 'School Admin / Proprietor')

    story_sections = [
        ('Students & Admissions', EMERALD, [
            ('Add a new student', 'the student has a complete profile and can access all school services from day one'),
            ('Bulk upload students via CSV', 'I can migrate from spreadsheets without entering records one by one'),
            ('View all admission applications', 'I can review and accept or reject them online'),
            ('Promote students to the next class in bulk', "I don't have to update each student record individually at year end"),
            ('Graduate SS3 students to alumni individually or in bulk', 'their records are preserved and the active student list stays clean'),
            ('Suspend or withdraw a student with a reason', 'the system always reflects the real enrollment status'),
            ('Update alumni records with post-graduation info (university, course, contact)', 'the school maintains a living alumni network'),
            ('Print a School Leaving Certificate for any alumnus', 'leavers receive official documentation from the school'),
            ('Re-admit an alumnus as an active student', 'returning students can be reinstated without losing their history'),
        ]),
        ('Academic Management', BLUE, [
            ('Set up the academic calendar with terms, sessions, and holidays', 'all scheduling is anchored to real school dates'),
            ('Create class arms (e.g. JSS1A, JSS1B)', 'large year groups can be split into manageable classes'),
            ('Build the school timetable', 'teachers and students know what happens when'),
            ('Create a scheme of work per subject per term', 'teaching is planned and documented in advance'),
            ('View a results broadsheet sorted by student average with proper rank positions', 'I can see academic standing at a glance'),
            ('Approve or reject teacher-submitted results with a comment', 'results are verified before students and parents can see them'),
            ('Manage the school library (add/remove books, track loans)', 'library usage is recorded and books are accounted for'),
        ]),
        ('Staff & HR', TEAL, [
            ('Add and manage staff profiles with roles, qualifications, and class assignments', 'workforce data is centralised'),
            ('Run monthly payroll through a structured stepper (generate → review → adjustments → confirm → publish)', 'salary processing is auditable'),
            ('Approve or reject staff leave requests', 'staff absence is tracked and formally authorised'),
            ('Record staff appraisals with scores and comments', 'performance management is documented'),
            ('Track substitute teacher coverage when a teacher is absent', 'no class goes unattended'),
        ]),
        ('Operations & Transport', AMBER, [
            ('Manage the school physical inventory (books, equipment, uniforms)', 'stock levels and asset values are tracked at all times'),
            ('Log visitor check-ins and check-outs', 'there is a complete record of everyone who enters the school'),
            ('Record sickbay visits with outcome and parent notification', 'student health incidents are documented and parents are always informed'),
            ('Manage bus routes and assign students to routes with their boarding stop', 'transport logistics are organised and parents know which bus their child is on'),
            ('Mark daily bus status (Waiting / Departed / Arrived / Delayed)', 'parents get real-time updates without having to call the school'),
            ('Approve or deny pickup authorisation requests submitted by parents', 'only authorised persons can collect students — child safety is maintained'),
        ]),
        ('Communications', PURPLE, [
            ('Send announcements to parents, teachers, or everyone', 'school-wide information reaches the right audience quickly'),
            ('Post notices to the Notice Board', 'important, non-date-specific information stays visible to the community'),
            ('Send bulk SMS/email campaigns to parents or staff', 'I can reach everyone instantly for urgent matters'),
            ('Create digital consent forms for events or permissions', 'I can collect parental consent electronically without paper'),
            ('Run satisfaction surveys for parents and view aggregated responses', 'I can improve school services based on real data'),
        ]),
    ]

    for section_title, sec_color, stories in story_sections:
        heading(doc, section_title, 2, color=sec_color, sb=12)
        for action, benefit in stories:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.left_indent  = Inches(0.3)
            rb = p.add_run('As a school admin, I want to ')
            rb.font.size = Pt(10.5)
            rb.font.color.rgb = GRAY_TEXT
            ra = p.add_run(action)
            ra.bold = True
            ra.font.size = Pt(10.5)
            ra.font.color.rgb = DARK_NAVY
            rs = p.add_run(', so that ')
            rs.font.size = Pt(10.5)
            rs.font.color.rgb = GRAY_TEXT
            rb2 = p.add_run(benefit + '.')
            rb2.italic = True
            rb2.font.size = Pt(10.5)
            rb2.font.color.rgb = sec_color

    page_break(doc)

    # Finance Officer
    heading(doc, 'Finance Officer', 1, sb=4)
    fin_stories = [
        ('Create and manage the fee structure per class and term', 'invoices are generated with the correct amounts'),
        ('Generate and send invoices to all parents at the start of term', 'fee collection can begin immediately'),
        ('View and export the payment ledger', 'I can see which invoices are paid, outstanding, or overdue at any moment'),
        ('Reconcile bank transfer payments against invoices using narration matching', 'manual payment matching is fast and accurate'),
        ('Record school expenses with category, amount, and notes', 'the P&L report is accurate'),
        ('Run and manage monthly payroll', 'staff salaries are processed on time'),
        ('Review and decide on parent loan applications', 'the lending programme runs in a controlled, auditable way'),
        ('Export financial reports (P&L, fee collection, payment ledger)', 'I can share data with the proprietor or external accountant'),
    ]
    for action, benefit in fin_stories:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.left_indent  = Inches(0.3)
        rb = p.add_run('As a finance officer, I want to ')
        rb.font.size = Pt(10.5); rb.font.color.rgb = GRAY_TEXT
        ra = p.add_run(action)
        ra.bold = True; ra.font.size = Pt(10.5); ra.font.color.rgb = DARK_NAVY
        rs = p.add_run(', so that ')
        rs.font.size = Pt(10.5); rs.font.color.rgb = GRAY_TEXT
        rb2 = p.add_run(benefit + '.')
        rb2.italic = True; rb2.font.size = Pt(10.5); rb2.font.color.rgb = EMERALD

    # Principal
    heading(doc, 'Principal', 1, sb=14)
    prin_stories = [
        ('View the academic dashboard with attendance rates, average results, and discipline summary',
         'I can monitor school performance at a glance'),
        ('Approve or reject teacher results before they are released to students',
         'I sign off on all academic outcomes'),
        ('Manage staff (add, edit, leave, appraisals)',
         'HR is handled from one place'),
        ('View and manage the timetable',
         'academic scheduling is always correct'),
        ('Post calendar events and notices to the school community',
         'the school community is kept informed'),
        ('Manage house points and record inter-house competitions',
         'the co-curricular programme is properly tracked'),
    ]
    for action, benefit in prin_stories:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.left_indent  = Inches(0.3)
        rb = p.add_run('As a principal, I want to ')
        rb.font.size = Pt(10.5); rb.font.color.rgb = GRAY_TEXT
        ra = p.add_run(action)
        ra.bold = True; ra.font.size = Pt(10.5); ra.font.color.rgb = DARK_NAVY
        rs = p.add_run(', so that ')
        rs.font.size = Pt(10.5); rs.font.color.rgb = GRAY_TEXT
        rb2 = p.add_run(benefit + '.')
        rb2.italic = True; rb2.font.size = Pt(10.5); rb2.font.color.rgb = TEAL

    page_break(doc)

    # ── 05 TEACHER STORIES ────────────────────────────────────────────────────
    label(doc, '05 — User Stories')
    heading(doc, 'Teacher')
    teacher_stories = [
        ('Take class attendance daily', 'the register is always up to date and parents of absent students are notified automatically', EMERALD),
        ('Enter exam and test results for my subject', "each student's academic performance is recorded against the right subject", EMERALD),
        ('Submit results for admin approval', 'results go through a proper verification process before being released', EMERALD),
        ('Create and assign homework', 'students have structured work to do outside class', BLUE),
        ('Mark submitted assignments and give grades and feedback', 'students know how they performed and what to improve', BLUE),
        ('Create and publish formative tests and CBT exams', 'students are assessed regularly beyond just end-of-term exams', BLUE),
        ('Write lesson plans and class notes and attach materials', 'my teaching is planned and resources are accessible to students', TEAL),
        ('Award house points to students for positive behaviour and effort', 'the house system rewards and motivates students', TEAL),
        ('Write diary notes about students that parents can read and reply to', 'parent-teacher communication is ongoing and documented per student', AMBER),
        ('Message parents directly', 'I can communicate about individual students when needed', AMBER),
        ('View my timetable', 'I know my schedule for every day of the week', PURPLE),
        ('View my payslip', 'I can verify my salary and see a breakdown of components', PURPLE),
        ('Apply for leave and see the outcome', 'my absence is formally requested and approved', PURPLE),
        ('Submit an appraisal self-assessment', 'my performance review is on record', PURPLE),
        ('Indicate which bus stop students board at', 'transport assignment is accurate for each student on my class list', TEAL),
    ]
    for action, benefit, color in teacher_stories:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.left_indent  = Inches(0.3)
        rb = p.add_run('As a teacher, I want to ')
        rb.font.size = Pt(10.5); rb.font.color.rgb = GRAY_TEXT
        ra = p.add_run(action)
        ra.bold = True; ra.font.size = Pt(10.5); ra.font.color.rgb = DARK_NAVY
        rs = p.add_run(', so that ')
        rs.font.size = Pt(10.5); rs.font.color.rgb = GRAY_TEXT
        rb2 = p.add_run(benefit + '.')
        rb2.italic = True; rb2.font.size = Pt(10.5); rb2.font.color.rgb = color

    page_break(doc)

    # ── 06 PARENT STORIES ─────────────────────────────────────────────────────
    label(doc, '06 — User Stories')
    heading(doc, 'Parent / Guardian')
    parent_stories = [
        ("See a dashboard for each of my children showing attendance, results, fees, and behaviour", "I can monitor my children's progress without going to the school", BLUE),
        ('Pay school fees online using card, bank transfer, or USSD', "I don't have to queue at the school office", EMERALD),
        ('View invoices and payment receipts at any time', 'I have permanent records of all payments made', EMERALD),
        ('Apply for a loan to cover school fees', 'I can spread the cost over manageable installments', EMERALD),
        ('View my child\'s exam results only after the school has approved them', 'I receive verified, official information', BLUE),
        ("Print a report card or transcript for my child", 'I have an official record of academic performance I can keep', BLUE),
        ("See my child's attendance record", 'I know exactly when they were present or absent', BLUE),
        ("Read diary notes from teachers and reply to them", "I stay informed about my child's day-to-day progress at school", AMBER),
        ('Add authorised pickup persons for my child and submit them for school approval', 'trusted adults can collect my child safely with school knowledge', AMBER),
        ("See the bus status for my child's route in real time", 'I know when to expect the bus without having to call the school', TEAL),
        ('Respond to digital consent forms', 'I can give or withhold permission for school events electronically', TEAL),
        ('Complete school surveys before the deadline', 'my feedback reaches the school in time to make a difference', TEAL),
        ('See school announcements and the notice board', 'I am kept informed of important school news', PURPLE),
        ('View the school calendar', 'I know about upcoming events, holidays, and exam dates well in advance', PURPLE),
        ("See my child's sickbay visit history", 'I am aware of any medical incidents at school', PURPLE),
        ("See my child's house points standing and their house's position", 'I can encourage participation in the house system', PURPLE),
    ]
    for action, benefit, color in parent_stories:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.left_indent  = Inches(0.3)
        rb = p.add_run('As a parent, I want to ')
        rb.font.size = Pt(10.5); rb.font.color.rgb = GRAY_TEXT
        ra = p.add_run(action)
        ra.bold = True; ra.font.size = Pt(10.5); ra.font.color.rgb = DARK_NAVY
        rs = p.add_run(', so that ')
        rs.font.size = Pt(10.5); rs.font.color.rgb = GRAY_TEXT
        rb2 = p.add_run(benefit + '.')
        rb2.italic = True; rb2.font.size = Pt(10.5); rb2.font.color.rgb = color

    page_break(doc)

    # ── 07 STUDENT STORIES ────────────────────────────────────────────────────
    label(doc, '07 — User Stories')
    heading(doc, 'Student')
    student_stories = [
        ('See my dashboard with attendance, upcoming assignments, results, and house points', 'I always know where I stand academically and in the house competition', BLUE),
        ('Access lesson notes and videos uploaded by my teacher', 'I can study and revise at home with the right materials', EMERALD),
        ('Submit assignments online before the deadline', "I don't have to carry physical work and I have a submission record", EMERALD),
        ('Take CBT tests and formative assessments on the platform', 'my performance is recorded immediately and I get instant feedback', TEAL),
        ('View my approved results with grades and position', 'I can see my grades as soon as the teacher and admin have signed off', TEAL),
        ('See my timetable', 'I know what subjects I have each day and can prepare', AMBER),
        ('See the school calendar', 'I am aware of upcoming events, exam dates, and school activities', AMBER),
        ('See the house leaderboard and my personal house points contribution', 'I know how my house is performing and how I am helping', PURPLE),
        ('See overdue assignments clearly marked', "I know when I've missed a deadline and can take responsibility", RED),
        ('See a banner when my results are pending review', "I understand why results aren't showing yet and when to expect them", RED),
    ]
    for action, benefit, color in student_stories:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.left_indent  = Inches(0.3)
        rb = p.add_run('As a student, I want to ')
        rb.font.size = Pt(10.5); rb.font.color.rgb = GRAY_TEXT
        ra = p.add_run(action)
        ra.bold = True; ra.font.size = Pt(10.5); ra.font.color.rgb = DARK_NAVY
        rs = p.add_run(', so that ')
        rs.font.size = Pt(10.5); rs.font.color.rgb = GRAY_TEXT
        rb2 = p.add_run(benefit + '.')
        rb2.italic = True; rb2.font.size = Pt(10.5); rb2.font.color.rgb = color

    # ── 08 SUPER ADMIN STORIES ────────────────────────────────────────────────
    heading(doc, 'Super Admin (CASPAA Staff)', 1, sb=18)
    sa_stories = [
        ('Onboard new schools with full configuration', 'a school is operational from day one without engineering involvement', EMERALD),
        ('View platform revenue broken down by subscriptions, transaction fees, and lending margins', "I can track CASPAA's financial health in real time", EMERALD),
        ('Monitor all active loans across all schools', 'credit risk is visible at the platform level and I can act early', BLUE),
        ('Manage support tickets from schools', 'issues are resolved quickly and tracked to closure', BLUE),
        ('View platform analytics (active schools, MAU, top features)', 'product decisions are based on real usage data', TEAL),
        ('Audit all sensitive actions across the platform', 'compliance and security are maintained for all schools', AMBER),
        ('Manage feature flags per school', 'I can roll out new features gradually without a code deployment', PURPLE),
    ]
    for action, benefit, color in sa_stories:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.left_indent  = Inches(0.3)
        rb = p.add_run('As a super admin, I want to ')
        rb.font.size = Pt(10.5); rb.font.color.rgb = GRAY_TEXT
        ra = p.add_run(action)
        ra.bold = True; ra.font.size = Pt(10.5); ra.font.color.rgb = DARK_NAVY
        rs = p.add_run(', so that ')
        rs.font.size = Pt(10.5); rs.font.color.rgb = GRAY_TEXT
        rb2 = p.add_run(benefit + '.')
        rb2.italic = True; rb2.font.size = Pt(10.5); rb2.font.color.rgb = color

    page_break(doc)

    # ── 09 USER FLOWS ─────────────────────────────────────────────────────────
    label(doc, '09 — User Flows')
    heading(doc, 'End-to-End User Flows')
    para(doc,
         'The following flows describe the step-by-step sequence of actions for each key process in CASPAA. '
         'Each flow includes the actor, trigger, and every screen/action in order.',
         size=10.5, color=GRAY_TEXT, sa=10)

    def flow_block(doc, flow_num, title, actors, trigger, phases):
        """
        phases: list of (phase_label, color, steps_list)
        where steps_list: list of (step_text,)  or  (actor_label, step_text)
        """
        # Flow header
        hdr_t = doc.add_table(rows=1, cols=1)
        hdr_t.alignment = WD_TABLE_ALIGNMENT.CENTER
        hc = hdr_t.rows[0].cells[0]
        set_cell_bg(hc, DARK_NAVY)
        hp = hc.paragraphs[0]
        hp.paragraph_format.space_before = Pt(5)
        hp.paragraph_format.space_after  = Pt(5)
        r1 = hp.add_run(f'Flow {flow_num:02d}  —  ')
        r1.bold = True; r1.font.size = Pt(12); r1.font.color.rgb = RGBColor(0x34, 0xd3, 0x99)
        r2 = hp.add_run(title)
        r2.bold = True; r2.font.size = Pt(13); r2.font.color.rgb = WHITE

        # Meta
        meta_t = doc.add_table(rows=1, cols=3)
        meta_t.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_cell_bg(meta_t.rows[0].cells[0], LIGHT_BLUE)
        set_cell_bg(meta_t.rows[0].cells[1], LIGHT_MINT)
        set_cell_bg(meta_t.rows[0].cells[2], LIGHT_AMBER)
        for ci, (lbl, val) in enumerate([('Actors', actors), ('Trigger', trigger), ('Complexity', phases[0][0] if phases else '')]):
            cell = meta_t.rows[0].cells[ci]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            rl = p.add_run(lbl + ': ')
            rl.bold = True; rl.font.size = Pt(8.5); rl.font.color.rgb = GRAY_TEXT
            rv = p.add_run(val)
            rv.font.size = Pt(9); rv.font.color.rgb = DARK_NAVY
        set_col_widths(meta_t, [Inches(2.0), Inches(3.0), Inches(1.5)])

        # Steps
        step_n = 1
        for phase_label, phase_color, steps in phases:
            if phase_label and phase_label not in ('—', ''):
                pp = doc.add_paragraph()
                pp.paragraph_format.space_before = Pt(8)
                pp.paragraph_format.space_after  = Pt(3)
                pp.paragraph_format.left_indent  = Inches(0.1)
                r = pp.add_run(phase_label)
                r.bold = True; r.font.size = Pt(10); r.font.color.rgb = phase_color
            for step in steps:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after  = Pt(2)
                p.paragraph_format.left_indent  = Inches(0.25)
                rn = p.add_run(f'{step_n}. ')
                rn.bold = True; rn.font.size = Pt(10); rn.font.color.rgb = phase_color
                rt = p.add_run(step)
                rt.font.size = Pt(10); rt.font.color.rgb = MID_GRAY
                step_n += 1
        doc.add_paragraph()

    # Define all flows
    flows = [
        (1, 'Student Enrollment',
         'School Admin', 'New student joins the school',
         [
             ('—', EMERALD, [
                 'Admin → Students tab → Add Student button',
                 'Fills form: name, date of birth, gender, class, admission number, photo, parent link',
                 'System validates: checks for duplicate admission number',
                 'On save: student record created, parent account linked and notified',
                 'Student now appears in class roster, attendance register, and fee invoice system',
                 'BULK UPLOAD: Admin uploads CSV → system previews all rows → Admin confirms import',
             ])
         ]),

        (2, 'Fee Invoice & Collection',
         'Finance Officer / Admin → Parent', 'Start of each new term',
         [
             ('Admin side', BLUE, [
                 'Finance → Fee Structure → Set fees per class for current term',
                 'Finance → Invoices → Generate Invoices (bulk, for all active students)',
                 'Invoices created with status: Unpaid',
                 'Parents notified in-app: "Your fee invoice for Term X is ready"',
             ]),
             ('Parent side', EMERALD, [
                 'Parent → Fees → Views invoice (itemised breakdown)',
                 'Parent clicks "Pay Now"',
                 'Payment modal opens — selects method: Card / Bank Transfer / USSD',
                 'Card: enters details → OTP → payment confirmed, invoice updated to Paid',
                 'Bank Transfer: copies account number → pays from banking app → system auto-reconciles',
                 'Parent can download or print receipt at any time',
             ]),
             ('Reconciliation', AMBER, [
                 'Finance → Reconciliation → unreconciled bank transfers appear',
                 'System auto-matches by narration (student name or admission number in reference)',
                 'Finance confirms the match or manually selects the invoice',
                 'Transaction marked reconciled, invoice marked Paid',
             ]),
         ]),

        (3, 'Daily Attendance',
         'Teacher', 'Start of each school day',
         [
             ('—', TEAL, [
                 "Teacher → Attendance → Select class → Today's date auto-fills",
                 'Student roster loads — all students shown as Present by default',
                 'Teacher taps a name to toggle: Present / Absent / Late',
                 'For late arrivals: system shows Late badge if clock-in is after the threshold time',
                 'Teacher clicks "Submit Attendance"',
                 'System saves record with schoolId, classId, date, and teacher ID',
                 'For absent students: parents notified in-app ("Your child was marked absent today")',
                 'Admin can view attendance summary per class or whole school from dashboard',
             ])
         ]),

        (4, 'Results Submission & Approval',
         'Teacher → Admin → Student / Parent', 'End of term / after exams',
         [
             ('Teacher submits', BLUE, [
                 'Teacher → Enter Results → Select subject + class + term',
                 'For each student: enter CA score and exam score',
                 'System computes total and grade automatically',
                 'Teacher reviews, then clicks "Submit for Approval"',
                 'Results saved with approved: false — students cannot see them yet',
             ]),
             ('Admin approves', EMERALD, [
                 'Admin → Academic → Results → sees pending results (badge count shown)',
                 'Admin clicks class/subject → reviews broadsheet (sorted by average, with proper rank)',
                 'Admin adds report comment per student (required before approval is possible)',
                 'Admin clicks "Approve" → system sets approved: true',
                 'Parents and students notified: "Results for [Subject] are now available"',
             ]),
             ('Student / Parent views', TEAL, [
                 'Student → My Results → sees only approved results',
                 'Parent → My Children → select child → Results tab → approved results',
                 'Parent can print Report Card (shows comment, grade, position in class)',
                 'If results still pending: amber banner shown — "X result(s) awaiting approval"',
             ]),
         ]),

        (5, 'Loan Application',
         'Parent → Finance Officer', 'Parent cannot pay full fees upfront',
         [
             ('Parent applies', EMERALD, [
                 'Parent → Loans → Apply for Loan',
                 'Fills application: amount requested, purpose, preferred repayment period',
                 'System runs live credit score: analyses payment history + income estimate + tenure',
                 'Score shown to parent (Excellent / Good / Fair / Poor) with breakdown',
                 'Application submitted — Finance Officer notified',
             ]),
             ('Finance decides', BLUE, [
                 'Finance → Lending → Pending applications',
                 'Reviews application + full credit score breakdown (5-factor assessment)',
                 'Approves with repayment schedule, or rejects with reason',
                 'If approved: disbursement recorded, parent notified with installment schedule',
             ]),
             ('Repayment', AMBER, [
                 'Parent → Loans → sees outstanding balance + installment due dates',
                 'Parent makes repayment via payment modal (same flow as fee payment)',
                 'Installment marked paid, outstanding balance reduces',
                 'Loan closes automatically when all installments are paid',
             ]),
         ]),

        (6, 'House Points Award',
         'Teacher → Admin → Student / Parent', 'Ongoing throughout each term',
         [
             ('Individual award', EMERALD, [
                 'Teacher → House Points → Award Points button',
                 'Selects student, enters points amount, toggles Award (green) or Deduct (red)',
                 'Enters reason (Good conduct, Academic excellence, etc.)',
                 "Saves → student's house total increases or decreases",
                 'Student and parent notified: "[Name] was awarded 5 house points for Good Conduct"',
             ]),
             ('Inter-house competition', BLUE, [
                 'Admin → House Points → Competitions → Record Competition',
                 'Enters event name (e.g. Sports Day) and event type',
                 'For each participating house: selects house and their finishing position',
                 'System assigns competition points: 1st=50, 2nd=35, 3rd=20, 4th=10',
                 'Leaderboard recalculates: individual merit points + competition points combined',
                 'All house members notified of the result',
             ]),
             ('Viewing the leaderboard', TEAL, [
                 'All roles → House Points → Leaderboard tab',
                 'Houses sorted by total combined points',
                 'Breakdown visible: merit pts · competition pts per house',
                 'Student sees their personal points contribution and their house rank',
             ]),
         ]),

        (7, 'Authorized Pickup Management',
         'Parent → Admin', 'Parent wants someone else to collect their child',
         [
             ('Parent submits request', EMERALD, [
                 'Parent → Transport → Authorized Pickup Persons → Add Person',
                 'Fills: full name, relationship to child, phone number',
                 'Submits request — status set to Pending',
                 'School admin notified: "New pickup authorization request for [child name]"',
             ]),
             ('Admin reviews and decides', BLUE, [
                 'Admin → Transport → Pickup Authorizations tab',
                 'Sees all pending requests highlighted in amber',
                 'Reviews: name, relationship, phone number, and which student',
                 'Clicks Approve or Deny',
                 'Parent notified of the outcome via in-app notification',
                 'If approved: person appears in the approved pickup list for that child',
             ]),
             ('Daily use', TEAL, [
                 'When person arrives, admin verifies against the approved list',
                 'Admin can revoke authorization at any time — parent is notified',
             ]),
         ]),

        (8, 'Bus Status Updates',
         'Admin / Transport Officer → Parent', 'Each school day — morning and afternoon runs',
         [
             ('—', AMBER, [
                 'Admin → Transport → Bus Status tab',
                 'Sees all active routes with current status (default: Waiting at School)',
                 'When bus departs: Admin clicks "Departed — En Route" for that route',
                 'All parents of students on that route receive push notification instantly',
                 'If delay: Admin clicks "Delayed" and optionally adds a note ("Stuck in traffic at Lekki bridge")',
                 'When bus arrives at destination: Admin clicks "Arrived" → parents notified',
                 'Parent → Transport → sees live status pill on their child\'s route card',
             ])
         ]),

        (9, 'Student Graduation & Alumni',
         'School Admin', 'End of final year (SS3 or JSS3)',
         [
             ('Individual graduation', EMERALD, [
                 'Admin → Students → open student profile → Actions → Graduate to Alumni',
                 'Graduation modal opens with pre-filled current class and year',
                 'Admin enters: graduation year, examination type (WAEC/NECO/BECE), index number, awards, certificate issued status',
                 'Confirms → student status changes to "alumni"',
                 'Student no longer appears in active registers, fee invoices, or attendance',
             ]),
             ('Bulk graduation', BLUE, [
                 'Admin → Academic → Bulk Promotion',
                 'Selects class (e.g. SS3A) → destination: "Graduate to Alumni"',
                 'All students in the class are graduated at once with the same graduation year',
             ]),
             ('Alumni record management', TEAL, [
                 'Admin → Alumni → search for graduate by name or admission number',
                 'Clicks "Update Info" → fills post-graduation details: university, course, personal email, phone',
                 'Clicks "Certificate" → prints School Leaving Certificate in a new print window',
                 'Stats shown: total alumni · this year\'s graduates · records with post-grad info filled',
             ]),
             ('Re-admission', AMBER, [
                 'Admin → Alumni → find alumnus → "Re-admit" button',
                 'Confirms → student status returns to Active',
                 'All alumni data is preserved — student rejoins the active register',
             ]),
         ]),

        (10, 'Teacher–Parent Communication Diary',
         'Teacher → Parent (and back)', 'Ongoing — teacher writes about a student',
         [
             ('Teacher writes', EMERALD, [
                 'Teacher → Diary → select student → Write Note',
                 'Chooses category: Homework / Academic / Behaviour / Health / General',
                 'Writes note: e.g. "Dawit struggled with fractions today — extra practice recommended"',
                 'Saves → parent notified: "Your child\'s teacher left a note"',
             ]),
             ('Parent reads and replies', BLUE, [
                 'Parent → Diary → sees notes, filterable by category',
                 'Reads note — automatically marked as read when opened',
                 'Taps Reply → types response',
                 'Teacher notified: "A parent replied to your note about [student name]"',
             ]),
             ('Teacher follows up', TEAL, [
                 'Teacher → Diary → sees parent reply in the thread',
                 'Can continue the conversation or consider it resolved',
             ]),
         ]),

        (11, 'Digital Consent Form',
         'Admin → Parent', 'School event requiring parental permission',
         [
             ('—', PURPLE, [
                 'Admin → Communications → Digital Consent → Create Form',
                 'Enters: title, description, deadline, target audience (all parents or specific class)',
                 'Publishes form → parents notified in-app',
                 'Parent → Consent → sees open forms with deadline countdown',
                 'Parent reads form, selects Agree or Decline',
                 'If declining: optionally notes a reason',
                 'Admin → Consent → sees real-time response count and breakdown',
                 'Admin can export consent list for trip registers or records',
                 'After deadline: form auto-closes, final tally is permanently visible',
             ])
         ]),

        (12, 'Formative Assessment / CBT Exam',
         'Teacher → Student', 'Teacher wants to assess students',
         [
             ('Teacher creates test', BLUE, [
                 'Teacher → Assessments → Create Test',
                 'Enters: title, subject, duration, pass mark, due date',
                 'Adds questions: MCQ (A/B/C/D), True/False, Short Answer with marks per question',
                 'Sets correct answers for auto-graded questions',
                 'Publishes test → all students in class notified',
             ]),
             ('Student takes test', EMERALD, [
                 'Student → Assessments → sees Pending tests (overdue shown with Overdue badge)',
                 'Opens test → sees duration reminder at the top',
                 'Answers all questions and clicks Submit',
                 'Confirmation dialog shown: "Submit? You cannot edit answers after submission"',
                 'Confirms → submission saved, MCQ/True-False scored immediately',
                 'Teacher notified: "[Name] submitted [Test Name]"',
             ]),
             ('Teacher grades and releases', TEAL, [
                 'Teacher → Assessments → sees submission list per test',
                 'For short-answer questions: enters score and written feedback',
                 'Final score released to student once all grading is complete',
                 'Student sees their score in the completed tests section',
             ]),
         ]),

        (13, 'Payroll Processing',
         'Finance Officer / Admin', 'Monthly — on payroll run date',
         [
             ('4-step payroll stepper', BLUE, [
                 'STEP 1 — Generate: Finance → Payroll → Start New Payroll Run; system loads all active staff with salary components',
                 'STEP 2 — Review: Finance reviews auto-calculated gross, deductions, and net pay for every staff member',
                 'STEP 3 — Adjustments: Finance adds one-off items: bonuses, salary advance deductions, corrections',
                 'STEP 4 — Approve & Publish: Finance/Admin clicks Approve; payslip records created for all staff',
                 'Each staff member notified: "Your payslip for [Month] is available"',
                 'Teacher → My Payslip → views itemised breakdown, downloads PDF',
             ])
         ]),

        (14, 'Calendar Event & Notice Board',
         'Admin / Principal — all roles view', 'Upcoming school event or information to share',
         [
             ('Calendar event (date-specific)', BLUE, [
                 'Admin → Calendar → Add Event button',
                 'Enters: title, start/end date, type (Holiday / Academic / Sports / Exam / Meeting), audience',
                 'Event appears on the calendar grid for all relevant users',
                 'Admin can edit or delete events; role-gated — only admin/principal can modify',
             ]),
             ('Notice Board (non-date pinned info)', EMERALD, [
                 'Admin → Calendar → Notice Board tab → Post Notice',
                 'Enters: title, message body, audience (Everyone / Parents / Teachers / Students)',
                 'Notice appears as a pinned card, sorted newest first',
                 'All users in the audience see it under Calendar → Notice Board tab',
                 'Admin can delete notices when they are no longer relevant',
             ]),
         ]),

        (15, 'Survey / Feedback Collection',
         'Admin → Parent', 'End of term or ongoing quality monitoring',
         [
             ('—', TEAL, [
                 'Admin → Surveys → Create Survey',
                 'Enters: title, description, deadline; adds questions (rating / multiple choice / open text)',
                 'Publishes → all parents notified: "New survey: [Title]"',
                 'Parent → Surveys → sees open surveys with time remaining to deadline',
                 'Parent opens survey, answers all questions, and submits',
                 'If deadline has passed: "This survey has closed" message shown',
                 'Admin → Surveys → clicks "View Results (N)" to see aggregated responses',
                 'Admin sees: rating averages, option breakdowns, all open-text answers',
             ])
         ]),

        (16, 'Sickbay Visit Recording',
         'Admin / Health Staff → Parent', 'Student reports ill during school hours',
         [
             ('—', RED, [
                 'Admin → Health/Sickbay → Log Visit',
                 'Selects student, enters complaint, any vital signs, and treatment given',
                 'Selects outcome: Treated & Returned / Resting in Sickbay / Sent Home / Referred to Hospital',
                 'For "Sent Home" or "Referred to Hospital": parent notification is mandatory and fires automatically',
                 'Parent receives: "[Child name] has been sent home from school. Please contact the school."',
                 'Admin → Health → Student record → full visit history (newest first)',
                 'Parent → Health → sees their child\'s complete sickbay visit history',
             ])
         ]),

        (17, 'Staff Leave Request',
         'Teacher → Admin', 'Teacher needs authorised time off',
         [
             ('Teacher applies', BLUE, [
                 'Teacher → Payslip (or HR section) → Leave Requests → Apply for Leave',
                 'Fills: leave type (Annual / Sick / Emergency / Study), start date, end date, reason',
                 'Submits → Admin and Principal notified',
             ]),
             ('Admin decides', EMERALD, [
                 'Admin → Staff & HR → Leave Requests → sees pending requests',
                 'Reviews: dates, leave type, reason, and remaining leave balance',
                 'Approves or denies with optional comment',
                 'Teacher notified of the outcome',
                 'If approved: leave dates reflected in staff attendance record',
                 'Admin arranges substitute teacher coverage for affected classes',
             ]),
         ]),

        (18, 'Inventory Management',
         'Admin', 'Receiving stock, issuing items, stock checks',
         [
             ('—', AMBER, [
                 'ADDING STOCK: Admin → Operations → Inventory',
                 'New item: Admin clicks Add Item → fills name, category, initial quantity, unit cost, minimum stock level, supplier',
                 'Existing item stock replenishment: Admin selects item → Stock In → quantity and reason',
                 'History entry created automatically for every movement',
                 'ISSUING: Admin selects item → Issue → enters quantity, recipient, and reason',
                 'Stock level decreases; if below minimum, a Low Stock badge appears on the item card',
                 'WRITE-OFF: Admin selects item → Write Off → enters quantity and reason (damaged/lost)',
                 'HISTORY: Admin → item → View History → full log of all movements (newest first)',
             ])
         ]),
    ]

    for f_num, f_title, f_actors, f_trigger, f_phases in flows:
        flow_block(doc, f_num, f_title, f_actors, f_trigger, f_phases)
        if f_num in (2, 6, 11, 14):
            page_break(doc)

    page_break(doc)

    # ── 10 PERMISSIONS MATRIX ─────────────────────────────────────────────────
    label(doc, '10 — Permissions Matrix')
    heading(doc, 'Roles × Features Matrix')
    para(doc, '✓ Full access  ·  ○ Own data only  ·  V View only  ·  — No access',
         size=9.5, color=GRAY_TEXT, sa=8)

    matrix = [
        # Feature, SuperAdmin, SchoolAdmin, Principal, Finance, Teacher, Parent, Student
        ('__HDR__', 'Academic Management', '', '', '', '', '', ''),
        ('Student records',        '✓', '✓', '✓', '—', 'V Class', 'O Child', 'O Own'),
        ('Results (enter)',         '—', '—', '—', '—', '✓',       '—',       '—'),
        ('Results (approve)',       '—', '✓', '✓', '—', '—',       '—',       '—'),
        ('Results (view)',          '✓', '✓', '✓', '—', '✓',       'Approved', 'Approved'),
        ('Attendance (mark)',       '—', '—', '—', '—', '✓',       '—',       '—'),
        ('Attendance (view)',       '✓', '✓', '✓', '—', 'O Class', 'O Child', 'O Own'),
        ('Timetable',               '✓', '✓', '✓', '—', 'O Own',   '—',       'V Own'),
        ('Assignments',             '✓', '✓', '✓', '—', '✓ Own',   'V Child', 'Submit'),
        ('CBT / Assessments',       '✓', '✓', '✓', '—', '✓ Create','—',       'Take'),
        ('__HDR__', 'Finance', '', '', '', '', '', ''),
        ('Fee structure',           '—', '✓', '—', '✓', '—',       '—',       '—'),
        ('Invoices',                '—', '✓', '—', '✓', '—',       'O Own',   '—'),
        ('Payments',                '—', '✓', '—', '✓', '—',       'O Own',   '—'),
        ('Loans',                   '—', '✓', '—', '✓', '—',       'O Own',   '—'),
        ('Payroll',                 '—', '✓', '—', '✓', '—',       '—',       '—'),
        ('Payslip',                 '—', '—', '—', 'All staff','O Own','—',    '—'),
        ('__HDR__', 'Communications & Community', '', '', '', '', '', ''),
        ('Announcements (post)',     '—', '✓', '✓', '—', '—',       '—',       '—'),
        ('Notice Board (post)',      '—', '✓', '✓', '—', '—',       '—',       '—'),
        ('Diary (write)',            '—', '—', '—', '—', '✓',       '—',       '—'),
        ('Diary (read/reply)',       '—', '—', '—', '—', '—',       'O Child', '—'),
        ('Surveys (create)',         '—', '✓', '✓', '—', '—',       '—',       '—'),
        ('Surveys (respond)',        '—', '—', '—', '—', '—',       '✓',       '—'),
        ('Calendar (post)',          '—', '✓', '✓', '—', '—',       '—',       '—'),
        ('Calendar (view)',          '—', '✓', '✓', '—', '✓',       '✓',       '✓'),
        ('__HDR__', 'Operations', '', '', '', '', '', ''),
        ('House Points (award)',     '—', '✓', '✓', '—', '✓',       '—',       '—'),
        ('House Points (view)',      '✓', '✓', '✓', '—', '✓',       'O Child', '✓'),
        ('Transport (manage)',       '—', '✓', '—', '—', '—',       '—',       '—'),
        ('Pickup (authorise)',       '—', '✓', '—', '—', '—',       'Submit',  '—'),
        ('Sickbay (log)',            '—', '✓', '✓', '—', '—',       '—',       '—'),
        ('Sickbay (view history)',   '—', '✓', '✓', '—', '—',       'O Child', '—'),
        ('Inventory',               '—', '✓', '✓', '—', '—',       '—',       '—'),
        ('Alumni',                  '—', '✓', '✓', '—', '—',       '—',       '—'),
        ('__HDR__', 'Platform (Super Admin only)', '', '', '', '', '', ''),
        ('All schools data',        '✓', '—', '—', '—', '—',       '—',       '—'),
        ('Platform revenue',        '✓', '—', '—', '—', '—',       '—',       '—'),
        ('Feature flags',           '✓', '—', '—', '—', '—',       '—',       '—'),
        ('Audit log',               '✓', 'Own school', '—', '—','—','—',       '—'),
    ]

    mt = doc.add_table(rows=0, cols=8)
    mt.style = 'Table Grid'
    mt.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_row(mt, ['Feature', 'Super Admin', 'Sch. Admin', 'Principal', 'Finance', 'Teacher', 'Parent', 'Student'],
            DARK_NAVY, size=8.5)

    role_colors = {
        '✓': EMERALD, 'V': BLUE, 'O': TEAL, '—': GRAY_TEXT,
        'Approved': AMBER, 'Submit': AMBER, 'All staff': EMERALD,
    }

    for i, row_data in enumerate(matrix):
        if row_data[0] == '__HDR__':
            r = mt.add_row()
            for ci in range(8):
                set_cell_bg(r.cells[ci], RGBColor(0x1a, 0x4a, 0x7a))
            r.cells[0].merge(r.cells[1])
            r.cells[0].merge(r.cells[2])
            r.cells[0].merge(r.cells[3])
            r.cells[0].merge(r.cells[4])
            r.cells[0].merge(r.cells[5])
            r.cells[0].merge(r.cells[6])
            r.cells[0].merge(r.cells[7])
            p = r.cells[0].paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            rr = p.add_run(row_data[1])
            rr.bold = True; rr.font.size = Pt(9); rr.font.color.rgb = WHITE
        else:
            bg = GRAY_BG if i % 2 == 0 else WHITE
            dr = mt.add_row()
            for ci, val in enumerate(row_data):
                set_cell_bg(dr.cells[ci], bg)
                dr.cells[ci].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p = dr.cells[ci].paragraphs[0]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after  = Pt(2)
                rr = p.add_run(str(val))
                rr.bold = (ci == 0)
                rr.font.size = Pt(8.5)
                if ci == 0:
                    rr.font.color.rgb = DARK_NAVY
                else:
                    first = str(val)[0] if val else '—'
                    rr.font.color.rgb = role_colors.get(first, GRAY_TEXT)

    set_col_widths(mt, [Inches(1.7), Inches(0.75), Inches(0.75), Inches(0.72), Inches(0.68), Inches(0.72), Inches(0.72), Inches(0.72)])

    page_break(doc)

    # ── 11 DATA MODEL ─────────────────────────────────────────────────────────
    label(doc, '11 — Data Model')
    heading(doc, 'Database Collections (56+)')
    para(doc, 'All tables are scoped by schoolId. Super Admin queries span all schools; all other roles are filtered to their schoolId.',
         size=10, color=GRAY_TEXT, sa=10)

    data_model = [
        ('People',          ['schools', 'teachers', 'parents', 'students']),
        ('Academic',        ['classes', 'subjects', 'arms', 'academicSessions', 'academicTerms', 'schemesOfWork',
                             'timetable', 'lessonPlans', 'learningMaterials', 'cbtExams', 'cbtSubmissions',
                             'formativeTests', 'formativeSubmissions', 'reportComments']),
        ('Records',         ['attendance', 'results', 'assignments', 'discipline', 'appraisals', 'appraisalCycles']),
        ('Finance',         ['feeStructures', 'invoices', 'transactions', 'expenses', 'loans', 'disbursements',
                             'payrollRuns', 'payslips', 'salaryAdvances', 'budgets', 'remittances', 'schoolInvoices']),
        ('Communications',  ['conversations', 'announcements', 'notifications', 'diaryEntries',
                             'consentForms', 'consentResponses', 'feedbackForms', 'feedbackResponses', 'smsCampaigns']),
        ('Operations',      ['inventory', 'inventoryRequests', 'busRoutes', 'busAssignments', 'authorizedPickups',
                             'busStatus', 'sickbayVisits', 'visitorLog', 'leaveRequests', 'staffAttendance',
                             'substituteCoverage']),
        ('Engagement',      ['houses', 'housePoints', 'houseEvents', 'schoolEvents', 'activities',
                             'studentActivities', 'libraryBooks', 'libraryLoans', 'admissionApplications']),
        ('Platform',        ['auditLog', 'loginSessions', 'supportTickets', 'helpArticles', 'usageEvents',
                             'errorLogs', 'systemMetrics', 'platformTeam', 'schoolRoles']),
        ('Config',          ['settings', 'academicCalendar']),
    ]

    dt = doc.add_table(rows=0, cols=2)
    dt.style = 'Table Grid'
    dt.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_row(dt, ['Domain', 'Collections'], DARK_NAVY)
    dm_colors = [LIGHT_MINT, LIGHT_BLUE, LIGHT_TEAL, LIGHT_AMBER, LIGHT_PURPLE, LIGHT_RED, GRAY_BG, LIGHT_BLUE, GRAY_BG]
    for i, (domain, tables) in enumerate(data_model):
        row = dt.add_row()
        set_cell_bg(row.cells[0], dm_colors[i % len(dm_colors)])
        set_cell_bg(row.cells[1], GRAY_BG if i % 2 else WHITE)
        row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p0 = row.cells[0].paragraphs[0]
        p0.paragraph_format.space_before = Pt(5)
        p0.paragraph_format.space_after  = Pt(5)
        r0 = p0.add_run(domain)
        r0.bold = True; r0.font.size = Pt(10); r0.font.color.rgb = DARK_NAVY
        p1 = row.cells[1].paragraphs[0]
        p1.paragraph_format.space_before = Pt(5)
        p1.paragraph_format.space_after  = Pt(5)
        r1 = p1.add_run('  ·  '.join(tables))
        r1.font.size = Pt(9); r1.font.color.rgb = MID_GRAY
    set_col_widths(dt, [Inches(1.3), Inches(5.2)])

    # Footer note
    para(doc, '\nDocument generated from the CASPAA prototype codebase · June 2026 · AfriSprings Confidential',
         size=8.5, color=GRAY_TEXT, align=WD_ALIGN_PARAGRAPH.CENTER, sa=0, sb=20, italic=True)

    out = r'C:\Users\USER\Desktop\CASPAA\CASPAA_PRD_Updated.docx'
    doc.save(out)
    print(f'Done: PRD saved to {out}')


if __name__ == '__main__':
    build_prd()
